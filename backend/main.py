from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse, RedirectResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional, Dict, Any
import os
import io
import pypdf
import docx
import json
import json_repair
import base64
import uuid
import asyncio
import httpx
import random
import subprocess
import fal_client
from pathlib import Path
from datetime import datetime
from dotenv import load_dotenv
import re


def extract_ig_handle(url: str) -> str:
    """Extracts clean Instagram handle from a URL."""
    try:
        return url.split('instagram.com/')[1].strip('/').split('?')[0].split('/')[0]
    except Exception:
        return ""


# Load .env from backend directory
load_dotenv(Path(__file__).parent / ".env")

from .ai_service import AIService
from .storage_service import StorageService, CLIENTS_DIR, TASKS_FILE, _get_sb
try:
    from .notion_service import notion_service, NOTION_COPY_VAULT_DB_ID
    _notion_available = True
except Exception as _ne:
    print(f"⚠️  Notion service non disponibile in main.py: {_ne}")
    notion_service = None
    NOTION_COPY_VAULT_DB_ID = None
    _notion_available = False
from .strategic_context_loader import get_strategic_context_for_generator
from .knowledge_loader import get_awareness_context
from .smart_lists_service import smart_lists_service
from .aria_agent import get_aria_agent
from .aria_memory import aria_memory

app = FastAPI(title="Antigravity Script Manager")

# CORS configuration — Triggering redeploy for task data sync
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

ai_service = AIService()
storage_service = StorageService()

# Job registry for long-running background analysis tasks
# { job_id: { "status": "running"|"done"|"error", "result": ..., "error": ... } }
_analysis_jobs: Dict[str, Dict[str, Any]] = {}

# Health check endpoint for Render
@app.get("/")
async def root():
    return {"status": "ok", "service": "Antigravity Backend", "version": "1.0"}

@app.get("/health")
async def health():
    return {"status": "healthy"}

@app.on_event("startup")
async def on_startup():
    import asyncio
    # Run sync in background so Railway health check passes immediately
    asyncio.create_task(asyncio.to_thread(storage_service.sync_from_supabase))

class StructuredLink(BaseModel):
    url: str
    description: Optional[str] = ""
    label: Optional[str] = ""

class Competitor(BaseModel):
    name: str
    links: List[StructuredLink] = []

class ClientCreate(BaseModel):
    name: str
    industry: Optional[str] = ""
    links: List[StructuredLink] = []
    competitors: List[Competitor] = []

class FeedbackRequest(BaseModel):
    script_id: Optional[str] = "script_1"
    feedback: str
    angle_title: str

class AngleRequest(BaseModel):
    user_prompt: Optional[str] = ""
    funnel_stage: Optional[str] = ""

class BrandIdentityUpdate(BaseModel):
    tone: Optional[str] = None
    visuals: Optional[str] = None
    colors: Optional[List[str]] = None
    buyer_personas: Optional[List[Dict[str, Any]]] = None

class ResearchUpdate(BaseModel):
    content: str

class ResearchRequest(BaseModel):
    user_prompt: Optional[str] = ""

class KnowledgeRequest(BaseModel):
    name: str
    instructions: str
    kb_type: str
    funnel_stage: Optional[str] = None

@app.post("/clients")
async def create_client(client: ClientCreate):
    client_id = storage_service.create_client(client.name, client.industry, client.links, client.competitors)
    return {"client_id": client_id}

@app.get("/clients")
async def list_clients():
    return storage_service.list_clients()

@app.get("/clients/{client_id}")
async def get_client(client_id: str):
    try:
        return storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")

@app.delete("/clients/{client_id}")
async def delete_client(client_id: str):
    try:
        storage_service.delete_client(client_id)
        return {"message": "Client deleted"}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")

@app.get("/clients/{client_id}/notion-sync")
async def sync_client_notion(client_id: str):
    try:
        metadata = storage_service.get_metadata(client_id)
        success = await notion_service.sync_client_metadata(client_id, metadata)
        if not success:
            raise HTTPException(status_code=500, detail="Failed to sync with Notion")
        return {"status": "success"}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")

@app.get("/api/vault/thumbnails/{filename}")
async def get_vault_thumbnail(filename: str):
    thumb_path = notion_service.thumbnails_dir / filename
    if thumb_path.exists():
        return FileResponse(thumb_path)
    raise HTTPException(status_code=404, detail="Thumbnail not found")

@app.patch("/clients/{client_id}/brand")
async def update_brand_identity(client_id: str, identity: BrandIdentityUpdate):
    try:
        metadata = storage_service.get_metadata(client_id)
        if identity.tone is not None:
            metadata["brand_identity"]["tone"] = identity.tone
            metadata["preferences"]["tone"] = identity.tone # Sync legacy
        if identity.visuals is not None:
            metadata["brand_identity"]["visuals"] = identity.visuals
        if identity.colors is not None:
            metadata["brand_identity"]["colors"] = identity.colors
        if identity.buyer_personas is not None:
            old_personas = metadata.get("brand_identity", {}).get("buyer_personas", [])
            new_personas = identity.buyer_personas
            
            print(f"PATCH Brand: Updating buyer_personas for client {client_id}")
            print(f"  - Old count: {len(old_personas)}")
            print(f"  - New count: {len(new_personas)}")
            
            # Safeguard: if new count is 0 but old count was large, it might be a race condition/error
            if len(new_personas) == 0 and len(old_personas) > 1:
                print(f"  WARNING: Attempting to clear {len(old_personas)} personas. Skipping update to prevent data loss.")
            else:
                metadata["brand_identity"]["buyer_personas"] = new_personas
        
        storage_service.save_metadata(client_id, metadata)
        return metadata["brand_identity"]
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")

@app.post("/clients/{client_id}/logo")
async def upload_logo(client_id: str, file: UploadFile = File(...)):
    try:
        metadata = storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")
    
    # Save logo file
    logo_dir = CLIENTS_DIR / client_id / "brand"
    logo_dir.mkdir(parents=True, exist_ok=True)
    
    # Remove old logo
    for old_file in logo_dir.glob("logo.*"):
        old_file.unlink()
    
    ext = os.path.splitext(file.filename)[1] or ".png"
    logo_filename = f"logo{ext}"
    logo_path = logo_dir / logo_filename
    
    content = await file.read()
    with open(logo_path, "wb") as f:
        f.write(content)

    metadata["brand_identity"]["logo"] = logo_filename
    storage_service.save_metadata(client_id, metadata)
    storage_service.save_logo_to_supabase(client_id, logo_filename, content, ext)
    return {"logo": logo_filename}

@app.get("/clients/{client_id}/logo")
async def get_logo(client_id: str):
    try:
        metadata = storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")

    logo_filename = metadata.get("brand_identity", {}).get("logo", "")
    if not logo_filename:
        raise HTTPException(status_code=404, detail="No logo uploaded")

    logo_path = CLIENTS_DIR / client_id / "brand" / logo_filename
    if not logo_path.exists():
        storage_service.ensure_local_logo(client_id)
    if not logo_path.exists():
        raise HTTPException(status_code=404, detail="Logo file not found")

    return FileResponse(logo_path)

@app.post("/clients/{client_id}/extract-colors")
async def extract_colors(client_id: str):
    try:
        metadata = storage_service.get_metadata(client_id)
        logo_filename = metadata.get("brand_identity", {}).get("logo", "")
        if not logo_filename:
            raise HTTPException(status_code=400, detail="Nessun logo caricato per questo cliente")
        
        logo_path = CLIENTS_DIR / client_id / "brand" / logo_filename
        if not logo_path.exists():
            storage_service.ensure_local_logo(client_id)
        if not logo_path.exists():
            raise HTTPException(status_code=404, detail="File logo non trovato")

        script_path = os.path.join(os.getcwd(), "execution", "extract_dominant_colors.py")
        # Explicitly use python3 to ensure Pillow is found
        result = subprocess.run(["python3", script_path, str(logo_path), "5"], capture_output=True, text=True)
        
        if result.returncode != 0:
            print(f"Extraction script error: {result.stderr}")
            raise HTTPException(status_code=500, detail=f"Errore script: {result.stderr}")
            
        colors = json.loads(result.stdout)
        return {"colors": colors}
    except Exception as e:
        print(f"Extraction error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/clients/{client_id}/extract-industry")
async def extract_industry(client_id: str):
    try:
        metadata = storage_service.get_metadata(client_id)
        # If already in metadata (e.g. from research), return it
        if metadata.get("industry"):
            return {"industry": metadata["industry"]}

        research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
        if not research_path.exists():
            raise HTTPException(status_code=404, detail="Analisi di mercato non trovata. Avvia prima la ricerca.")
            
        with open(research_path, "r") as f:
            content = f.read()
            
        # 1. Regex Match for explicit phrases
        match = re.search(r"specializzat[ao] in\s+([^\.]+)", content, re.IGNORECASE)
        if match:
            industry = match.group(1).strip()
            industry = industry[:60]
            industry = re.sub(r",| e | per | come | ad | in | con | \d+%", "", industry).strip().title()
        else:
            # 2. AI Fallback: Summarize the industry from the research text
            print(f"DEBUG: Industry regex failed. Using AI fallback.")
            prompt = f"Basandoti su questo testo di ricerca, estrai SOLO il settore/categoria del business (max 3-5 parole). Esempio: 'Centro Estetico Avanzato', 'Ristorante Gourmet Irpino'.\n\nTesto: {content[:1000]}"
            industry = await ai_service._call_ai("anthropic/claude-3-7-sonnet", [{"role": "user", "content": prompt}], temperature=0)
            industry = industry.strip().strip('"').strip("'")

        metadata["industry"] = industry
        storage_service.save_metadata(client_id, metadata)
        
        return {"industry": industry}
    except Exception as e:
        print(f"Industry extraction error: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/clients/{client_id}/logo")
async def delete_logo(client_id: str):
    try:
        metadata = storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")
    
    logo_filename = metadata.get("brand_identity", {}).get("logo", "")
    if logo_filename:
        logo_path = CLIENTS_DIR / client_id / "brand" / logo_filename
        if logo_path.exists():
            logo_path.unlink()
        storage_service.delete_logo_from_supabase(client_id, logo_filename)

    metadata["brand_identity"]["logo"] = ""
    storage_service.save_metadata(client_id, metadata)
    return {"message": "Logo deleted"}

@app.post("/clients/{client_id}/upload")
async def upload_file(client_id: str, file: UploadFile = File(...)):
    content = await file.read()
    storage_service.save_file(client_id, file.filename, content)
    return {"message": f"File {file.filename} uploaded successfully"}

@app.post("/clients/{client_id}/files")
async def upload_multiple_files(client_id: str, files: List[UploadFile] = File(...)):
    uploaded = []
    for file in files:
        content = await file.read()
        storage_service.save_file(client_id, file.filename, content)
        uploaded.append(file.filename)
    return {"message": f"Uploaded {len(uploaded)} files", "files": uploaded}

@app.get("/clients/{client_id}/files")
async def list_files(client_id: str):
    return storage_service.list_files(client_id)

@app.delete("/clients/{client_id}/files/{filename}")
async def delete_file(client_id: str, filename: str):
    try:
        storage_service.delete_file(client_id, filename)
        return {"message": f"File {filename} deleted"}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="File not found")

@app.get("/clients/{client_id}/research")
async def get_research(client_id: str):
    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    if not research_path.exists():
        return {"content": ""}
    with open(research_path, "r") as f:
        return {"content": f.read()}

@app.put("/clients/{client_id}/research")
async def update_research(client_id: str, update: ResearchUpdate):
    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    research_path.parent.mkdir(parents=True, exist_ok=True)
    with open(research_path, "w") as f:
        f.write(update.content)
    storage_service.sync_research(client_id, update.content)
    return {"message": "Research updated"}

class SWOTUpdate(BaseModel):
    strengths: Optional[str] = None
    weaknesses: Optional[str] = None
    opportunities: Optional[str] = None
    threats: Optional[str] = None

class LinksUpdate(BaseModel):
    links: List[StructuredLink]

class CompetitorsUpdate(BaseModel):
    competitors: List[Competitor]

@app.post("/clients/{client_id}/research")
async def perform_research(client_id: str, request: ResearchRequest = ResearchRequest()):
    try:
        metadata = storage_service.get_metadata(client_id)
        
        print(f"DEBUG: Starting research for {client_id}")
        
        # --- AUTOMATED DEEP RESEARCH (META / INSTAGRAM INTEGRATION) ---
        ig_data_context = ""
        ig_token = os.getenv("META_ACCESS_TOKEN")

        # Collect Instagram handles from client links AND all competitor links
        client_ig_handles = [
            (extract_ig_handle(l.get('url', '')), "cliente")
            for l in metadata.get('links', [])
            if 'instagram.com' in l.get('url', '').lower()
        ]
        competitor_ig_handles = [
            (extract_ig_handle(cl.get('url', '')), f"competitor: {comp.get('name','')}")
            for comp in metadata.get('competitors', [])
            for cl in comp.get('links', [])
            if 'instagram.com' in cl.get('url', '').lower()
        ]
        all_handles = [(h, label) for h, label in client_ig_handles + competitor_ig_handles if h]

        if all_handles and ig_token:
            print(f"DEBUG: IG research for {len(all_handles)} accounts: {[h for h,_ in all_handles]}")

            # Find OUR IG Business Account ID once (reused for all Business Discovery calls)
            ig_user_id = None
            try:
                async with httpx.AsyncClient(timeout=20.0) as hc:
                    pages = (await hc.get(
                        "https://graph.facebook.com/v19.0/me/accounts",
                        params={"access_token": ig_token}
                    )).json().get("data", [])
                    for page in pages:
                        ig_resp = (await hc.get(
                            f"https://graph.facebook.com/v19.0/{page['id']}",
                            params={"fields": "instagram_business_account", "access_token": ig_token}
                        )).json()
                        if "instagram_business_account" in ig_resp:
                            ig_user_id = ig_resp["instagram_business_account"]["id"]
                            break
            except Exception as e:
                print(f"DEBUG: Failed to find IG user ID: {e}")

            async def fetch_ig_account(handle: str, label: str) -> dict:
                """Fetches IG account data via Business Discovery. Returns dict with all data."""
                result = {"handle": handle, "label": label, "followers": 0, "comments_collected": 0, "text": ""}
                if not ig_user_id:
                    return result
                try:
                    async with httpx.AsyncClient(timeout=30.0) as hc:
                        disc = (await hc.get(
                            f"https://graph.facebook.com/v19.0/{ig_user_id}",
                            params={
                                "fields": f"business_discovery.username({handle}){{followers_count,media_count,biography,media{{id,caption,like_count,comments_count,timestamp,media_type}}}}",
                                "access_token": ig_token,
                            }
                        )).json()

                    if "error" in disc:
                        print(f"DEBUG: Business Discovery error for @{handle}: {disc['error'].get('message')}")
                        return result

                    ig_data = disc.get("business_discovery", {})
                    followers = ig_data.get("followers_count", 0)
                    media_count = ig_data.get("media_count", 0)
                    bio = ig_data.get("biography", "")
                    posts = ig_data.get("media", {}).get("data", [])

                    total_eng = sum(p.get("like_count", 0) + p.get("comments_count", 0) for p in posts)
                    avg_eng_rate = round((total_eng / len(posts) / followers * 100), 2) if posts and followers else 0

                    block = (
                        f"\n{'='*60}\n"
                        f"📊 ACCOUNT: @{handle} [{label.upper()}]\n"
                        f"Follower: {followers:,} | Post totali: {media_count} | "
                        f"Engagement Rate medio: {avg_eng_rate}%\n"
                        f"Bio: {bio}\n\nTOP POST (engagement più alto):\n"
                    )

                    # Sort posts by engagement to show the best ones first
                    sorted_posts = sorted(posts, key=lambda p: p.get("like_count", 0) + p.get("comments_count", 0), reverse=True)
                    for p in sorted_posts[:6]:
                        caption_preview = (p.get("caption", "") or "")[:150]
                        block += (
                            f"  [{p.get('timestamp','')[:10]}] {p.get('media_type','?')} | "
                            f"❤️ {p.get('like_count',0):,} likes | 💬 {p.get('comments_count',0)} commenti\n"
                            f"  Caption: {caption_preview}\n"
                        )

                    # Fetch comment text from top 5 posts in parallel
                    async def fetch_comments(p: dict) -> tuple:
                        m_id = p.get("id")
                        if not m_id:
                            return p, []
                        try:
                            async with httpx.AsyncClient(timeout=15.0) as hc:
                                data = (await hc.get(
                                    f"https://graph.facebook.com/v19.0/{m_id}/comments",
                                    params={"fields": "text,like_count", "access_token": ig_token, "limit": 30}
                                )).json().get("data", [])
                            return p, data
                        except Exception:
                            return p, []

                    comments_results = await asyncio.gather(*[fetch_comments(p) for p in sorted_posts[:5]])
                    comments_collected = 0
                    for post, comm_data in comments_results:
                        if comm_data:
                            block += f"\n  💬 COMMENTI POST (❤️{post.get('like_count',0)} likes):\n"
                            for c in comm_data[:25]:
                                text = c.get("text", "").strip()
                                if text:
                                    block += f"    - \"{text}\"\n"
                                    comments_collected += 1

                    result["followers"] = followers
                    result["comments_collected"] = comments_collected
                    result["text"] = block
                    return result

                except Exception as e:
                    print(f"DEBUG: fetch_ig_account error for @{handle}: {e}")
                    return result

            # Fetch all known accounts in parallel — collect results to reuse for threshold check
            accounts_results = await asyncio.gather(*[fetch_ig_account(h, l) for h, l in all_handles])
            total_followers_check = 0
            total_comments_check = 0
            for account_data in accounts_results:
                ig_data_context += account_data["text"]
                total_followers_check += account_data["followers"]
                total_comments_check += account_data["comments_collected"]

            # --- FALLBACK: se i dati sono scarsi, cerca account leader nello stesso settore ---
            FOLLOWER_THRESHOLD = 8000
            COMMENT_THRESHOLD = 20
            if total_followers_check < FOLLOWER_THRESHOLD or total_comments_check < COMMENT_THRESHOLD:
                print(f"DEBUG: Insufficient data (followers={total_followers_check}, comments={total_comments_check}). Searching for niche leaders...")
                industry = metadata.get("industry", "") or metadata.get("name", "")
                try:
                    niche_prompt = (
                        f"Settore del cliente: {industry}. Cliente: {metadata.get('name','')}.\n"
                        f"Dammi esattamente 5 handle Instagram (@username) di account pubblici ITALIANI "
                        f"con molti follower (minimo 20k) in questo stesso settore/nicchia. "
                        f"Devono essere account reali, attivi, con molti commenti sui post. "
                        f"Rispondi SOLO con una lista JSON di handle senza @ es: [\"account1\",\"account2\",...]. "
                        f"Zero spiegazioni, solo il JSON."
                    )
                    handles_str = await ai_service._call_ai(
                        "perplexity/sonar",
                        [{"role": "user", "content": niche_prompt}]
                    )
                    niche_handles = json_repair.loads(handles_str)
                    if isinstance(niche_handles, list):
                        ig_data_context += f"\n\n{'='*60}\n⚡ ACCOUNT LEADER DI SETTORE (ricerca allargata — stessa nicchia, più follower):\n"
                        already_done = {h for h, _ in all_handles}
                        new_handles = [nh.strip().lstrip("@") for nh in niche_handles[:5] if nh.strip().lstrip("@") and nh.strip().lstrip("@") not in already_done]
                        print(f"DEBUG: Fetching {len(new_handles)} niche leaders in parallel: {new_handles}")
                        leader_results = await asyncio.gather(*[fetch_ig_account(nh, "leader di settore") for nh in new_handles])
                        for leader_data in leader_results:
                            ig_data_context += leader_data["text"]
                except Exception as e:
                    print(f"DEBUG: Niche leader search failed: {e}")

        raw_content = storage_service.get_raw_data_content(client_id)

        # AIService now returns a DICT directly — pass social data separately
        research_data = await ai_service.perform_market_research(
            metadata, raw_content, request.user_prompt, ig_data_context
        )
        
    except Exception as e:
        error_msg = str(e)
        if "402 Payment Required" in error_msg:
            return {"error": "Crediti API esauriti (402 Payment Required). Ricarica il tuo account OpenRouter per continuare."}
        return {"error": f"Errore di connessione API AI: {error_msg}"}
    
    try:
        # Distribute data into metadata
        metadata["strategy"] = research_data.get("strategy", metadata.get("strategy", ""))
        metadata["industry"] = research_data.get("industry", metadata.get("industry", ""))
        
        # Merge personas
        personas = research_data.get("buyer_personas", [])
        if isinstance(personas, list):
            metadata["brand_identity"]["buyer_personas"] = personas
        
        # Handle tone
        tone_label = research_data.get("suggested_tone", "naturale")
        tone_details = research_data.get("tone_description", "")
        metadata["brand_identity"]["tone"] = f"{tone_label}: {tone_details}" if tone_details else tone_label
        metadata["preferences"]["tone"] = metadata["brand_identity"]["tone"] # Sync
        
        if "target_vocabulary" in research_data:
            metadata["preferences"]["target_vocabulary"] = research_data["target_vocabulary"]

        if "audience_pain_points" in research_data:
            metadata["preferences"]["audience_pain_points"] = research_data["audience_pain_points"]

        if "top_content_patterns" in research_data:
            metadata["preferences"]["top_content_patterns"] = research_data["top_content_patterns"]

        if "key_products" in research_data and isinstance(research_data["key_products"], list):
            metadata["key_products"] = research_data["key_products"]

        storage_service.save_metadata(client_id, metadata)

        # Save the research text as markdown file — append extra fields for future AI use
        research_text = research_data.get("research_text", "Errore generazione testo")
        pain_points = research_data.get("audience_pain_points", [])
        content_patterns = research_data.get("top_content_patterns", [])
        vocab = research_data.get("target_vocabulary", [])

        extra_sections = ""
        if pain_points:
            extra_sections += "\n\n---\n## DOLORI REALI DEL TARGET (da commenti/dati reali)\n" + "\n".join(f"- {p}" for p in pain_points)
        if content_patterns:
            extra_sections += "\n\n---\n## PATTERN CONTENUTO AD ALTO ENGAGEMENT\n" + "\n".join(f"- {p}" for p in content_patterns)
        if vocab:
            extra_sections += "\n\n---\n## VOCABOLARIO REALE DEL TARGET\n" + "\n".join(f"- {v}" for v in vocab)

        full_research_text = research_text + extra_sections

        research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
        research_path.parent.mkdir(parents=True, exist_ok=True)
        with open(research_path, "w") as f:
            f.write(full_research_text)
        storage_service.sync_research(client_id, full_research_text)

        return {"research": full_research_text, "data": research_data}
        
    except Exception as e:
        print(f"RESEARCH HANDLER ERROR: {str(e)}")
        return {"error": f"Errore Elaborazione Dati: {str(e)}"}

@app.patch("/clients/{client_id}/swot")
async def update_swot(client_id: str, update: SWOTUpdate):
    metadata = storage_service.get_metadata(client_id)
    if "swot" not in metadata:
        metadata["swot"] = {}
    for key, value in update.dict(exclude_unset=True).items():
        metadata["swot"][key] = value
    storage_service.save_metadata(client_id, metadata)
    return metadata["swot"]

@app.put("/clients/{client_id}/strategy")
async def put_strategy(client_id: str, update: ResearchUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["strategy"] = update.content
    storage_service.save_metadata(client_id, metadata)
    return {"strategy": metadata["strategy"]}

@app.put("/clients/{client_id}/objectives")
async def put_objectives(client_id: str, update: ResearchUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["objectives"] = update.content
    storage_service.save_metadata(client_id, metadata)
    return {"objectives": metadata["objectives"]}

@app.post("/clients/{client_id}/deep-analysis")
async def run_deep_analysis(client_id: str):
    """Re-generates SWOT, buyer personas, tone, objectives and strategy from existing research."""
    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    if not research_path.exists():
        raise HTTPException(status_code=400, detail="Perform research first before running deep analysis")
    
    with open(research_path, "r") as f:
        research_text = f.read()
    
    metadata = storage_service.get_metadata(client_id)
    
    try:
        target_vocab = metadata.get("preferences", {}).get("target_vocabulary", [])
        analysis = await ai_service.generate_deep_analysis(research_text, metadata, target_vocabulary=target_vocab)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI analysis failed: {str(e)}")
    
    # Update metadata with analysis results
    if "swot" in analysis:
        metadata["swot"] = analysis["swot"]
    
    if "buyer_personas" in analysis and isinstance(analysis["buyer_personas"], list):
        # MERGE LOGIC: Don't overwrite. Append new ones if the name doesn't exist.
        if "brand_identity" not in metadata: metadata["brand_identity"] = {}
        existing = metadata["brand_identity"].get("buyer_personas", [])
        
        # Track names to avoid exact duplicates
        existing_names = {p.get("name", "").lower() for p in existing}
        
        new_count = 0
        for p in analysis["buyer_personas"]:
            name = p.get("name", "")
            if name.lower() not in existing_names:
                existing.append(p)
                new_count += 1
        
        metadata["brand_identity"]["buyer_personas"] = existing
        print(f"Deep Analysis: Added {new_count} new personas, preserved {len(existing) - new_count} existing.")

    if "tone" in analysis:
        metadata["brand_identity"]["tone"] = analysis["tone"]
        metadata["preferences"]["tone"] = analysis["tone"]
    if "objectives" in analysis:
        metadata["objectives"] = analysis["objectives"]
    if "strategy" in analysis:
        metadata["strategy"] = analysis["strategy"]
    
    storage_service.save_metadata(client_id, metadata)
    return {"analysis": analysis, "metadata": metadata}

class PersonaSpecificaRequest(BaseModel):
    target_service: str

@app.post("/clients/{client_id}/personas-specifiche")
async def create_specific_personas(client_id: str, req: PersonaSpecificaRequest):
    """Genera buyer personas specifiche per un servizio mirato usando Perplexity."""
    try:
        # Load fresh metadata
        metadata = storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
    
    client_name = metadata.get("name", "")
    sector = metadata.get("brand_identity", {}).get("tone", "")
    
    prompt = f"""Sei un esperto di marketing strategico e buyer personas. Utilizza le tue capacità di ricerca web per analizzare il mercato attuale.
    
Il cliente è: {client_name}
Contesto/Settore: {sector}

COMPITO: Genera 2 buyer personas ESTREMAMENTE SPECIFICHE per il seguente servizio/problema:
"{req.target_service}"

Per OGNI persona, fornisci una struttura completa e approfondita:
- "name": Nome e Cognome fittizio realistico italiano
- "type": Etichetta del tipo di persona (es: "Il manager stressato", "La studentessa con acne tardiva")
- "profile": Descrizione dettagliata (situazione, stile di vita, perché cerca questo servizio ora)
- "buying_habits": Abitudini d'acquisto (come sceglie, cosa guarda prima, canali preferiti)
- "fears": Le paure più grandi, le obiezioni, i blocchi emotivi legati a "{req.target_service}"
- "desires": I desideri profondi e la trasformazione finale cercata
- "critical_info": Info indispensabile per la strategia (trigger, piattaforme, messaggi chiave)

IMPORTANTE: Sii CHIRURGICO. Sfrutta la tua capacità di ricerca web per trovare insight reali. Non dare profili generici.
Rispondi SOLO con un JSON array valido, senza markdown:
[
  {{
    "name": "...",
    "type": "...",
    "profile": "...",
    "buying_habits": "...",
    "fears": "...",
    "desires": "...",
    "critical_info": "..."
  }}
]"""

    try:
        response = await ai_service._call_ai(
            model="perplexity/sonar",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.4,
            max_tokens=4000
        )
        
        import json_repair
        # Clean response from potential markdown fences if perplexity adds them
        clean_response = response.strip()
        if "```" in clean_response:
             import re
             match = re.search(r'```(?:json)?\s*(.*?)\s*```', clean_response, re.DOTALL)
             if match:
                 clean_response = match.group(1)
        
        personas = json_repair.loads(clean_response)
        if not isinstance(personas, list):
            personas = [personas]
        
        # Tag each persona with the target service
        for p in personas:
            p["servizio_specifico"] = req.target_service
        
        # Robust Merging
        if "brand_identity" not in metadata or not isinstance(metadata["brand_identity"], dict):
            metadata["brand_identity"] = {}
        
        existing = metadata["brand_identity"].get("buyer_personas", [])
        if not isinstance(existing, list):
            existing = []
            
        # Add new personas
        existing.extend(personas)
        metadata["brand_identity"]["buyer_personas"] = existing
        
        storage_service.save_metadata(client_id, metadata)
        return metadata
        
    except Exception as e:
        print(f"Error generating specific personas: {e}")
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")
        
    except Exception as e:
        print(f"Error generating specific personas: {e}")
        raise HTTPException(status_code=500, detail=f"AI generation failed: {str(e)}")

@app.patch("/clients/{client_id}/links")
async def update_links(client_id: str, update: LinksUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["links"] = [l.dict() for l in update.links]
    storage_service.save_metadata(client_id, metadata)
    return metadata["links"]

class IndustryUpdate(BaseModel):
    industry: str

@app.patch("/clients/{client_id}/industry")
async def update_industry(client_id: str, update: IndustryUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["industry"] = update.industry
    storage_service.save_metadata(client_id, metadata)
    return {"industry": metadata["industry"]}

@app.patch("/clients/{client_id}/competitors")
async def update_competitors(client_id: str, update: CompetitorsUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["competitors"] = [c.dict() for c in update.competitors]
    storage_service.save_metadata(client_id, metadata)
    return metadata["competitors"]

class AdAccountUpdate(BaseModel):
    ad_account_id: str

@app.patch("/clients/{client_id}/ad-account")
async def update_ad_account(client_id: str, update: AdAccountUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["ad_account_id"] = update.ad_account_id.strip()
    storage_service.save_metadata(client_id, metadata)
    return {"ad_account_id": metadata["ad_account_id"]}

class MetaPixelUpdate(BaseModel):
    pixel_id: Optional[str] = ""
    meta_access_token: Optional[str] = ""

@app.patch("/clients/{client_id}/meta-pixel")
async def update_meta_pixel(client_id: str, update: MetaPixelUpdate):
    metadata = storage_service.get_metadata(client_id)
    metadata["pixel_id"] = update.pixel_id.strip() if update.pixel_id else ""
    if update.meta_access_token:
        metadata["meta_access_token"] = update.meta_access_token.strip()
    elif update.meta_access_token == "":
        metadata.pop("meta_access_token", None)
    storage_service.save_metadata(client_id, metadata)
    return {"pixel_id": metadata.get("pixel_id", ""), "has_token": bool(metadata.get("meta_access_token"))}

@app.get("/clients/{client_id}/meta-ads-insights")
async def fetch_meta_ads_insights(client_id: str, date_preset: str = "last_30d", since: str = "", until: str = ""):
    metadata = storage_service.get_metadata(client_id)
    ad_account_id = metadata.get("ad_account_id", "").strip()
    if not ad_account_id:
        raise HTTPException(status_code=400, detail="Ad Account ID non configurato per questo cliente. Aggiungilo nella sezione Sorgenti.")

    token = metadata.get("meta_access_token") or os.getenv("META_ACCESS_TOKEN")
    if not token:
        raise HTTPException(status_code=500, detail="META_ACCESS_TOKEN non trovato nel file .env")

    # Ensure act_ prefix
    if not ad_account_id.startswith("act_"):
        ad_account_id = f"act_{ad_account_id}"

    fields = "spend,cpm,ctr,cpc,cpp,impressions,reach,actions,cost_per_action_type"
    url = f"https://graph.facebook.com/v19.0/{ad_account_id}/insights"

    async with httpx.AsyncClient(timeout=30.0) as client:
        params: dict = {
            "fields": fields,
            "level": "account",
            "access_token": token,
        }
        if since and until:
            import json as _json
            params["time_range"] = _json.dumps({"since": since, "until": until})
        else:
            params["date_preset"] = date_preset
        resp = await client.get(url, params=params)

    if resp.status_code != 200:
        error_msg = resp.json().get("error", {}).get("message", resp.text[:200])
        raise HTTPException(status_code=resp.status_code, detail=f"Meta Ads API error: {error_msg}")

    data = resp.json().get("data", [])
    if not data:
        raise HTTPException(status_code=404, detail="Nessun dato trovato per questo periodo. Verifica che l'Ad Account abbia campagne attive.")

    row = data[0]

    # Extract purchase conversions from actions list
    actions = row.get("actions", [])
    purchase_action = next((a for a in actions if a.get("action_type") in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
    conversioni = purchase_action["value"] if purchase_action else ""

    # Extract CPA from cost_per_action_type
    cpa_actions = row.get("cost_per_action_type", [])
    cpa_action = next((a for a in cpa_actions if a.get("action_type") in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
    cpa = round(float(cpa_action["value"]), 2) if cpa_action else ""

    return {
        "budget_speso": round(float(row.get("spend", 0)), 2) if row.get("spend") else "",
        "cpm": round(float(row.get("cpm", 0)), 2) if row.get("cpm") else "",
        "ctr": round(float(row.get("ctr", 0)), 2) if row.get("ctr") else "",
        "cpc": round(float(row.get("cpc", 0)), 2) if row.get("cpc") else "",
        "cpa": str(cpa) if cpa else "",
        "impressions": row.get("impressions", ""),
        "reach": row.get("reach", ""),
        "conversioni": conversioni,
        "date_preset": date_preset,
    }


# ══════════════════════════════════════════════════════════
#  LIVE ADS — BM overview + campaign breakdown + AI analysis
# ══════════════════════════════════════════════════════════

def _meta_extract_conversions(actions: list) -> int:
    purchase = next((a for a in (actions or []) if a.get("action_type") in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
    return int(float(purchase["value"])) if purchase else 0

def _meta_extract_cpa(cpa_list: list) -> Optional[float]:
    item = next((a for a in (cpa_list or []) if a.get("action_type") in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
    return round(float(item["value"]), 2) if item else None

ANDROMEDA_EXPERT_KNOWLEDGE = """
## COME FUNZIONA ANDROMEDA (fatti tecnici reali — deploy globale ottobre 2024)
- Sistema di retrieval basato su embedding: mappa utenti, contesti e ads in uno spazio vettoriale ad alta dimensione.
- 10.000x più complesso dei sistemi precedenti; usa chip NVIDIA Grace Hopper e MTIA di Meta.
- Processo a 3 stadi: (1) Retrieval via embedding similarity → shortlist di migliaia di ads; (2) Ranking: predizione di azione utente; (3) Asta finale.
- Se il tuo ad non supera il retrieval (embedding similarity bassa), non entra MAI nell'asta — indipendentemente dal bid o dal targeting.
- Segnali comportamentali: quanto l'utente si ferma su un post, quali video rivede, chi messaggia dopo aver visto qualcosa, pattern di acquisto, interazioni cross-platform (FB + IG + WhatsApp).
- IMPLICAZIONE CRITICA: La rilevanza della creatività (non il targeting demografico) determina la visibilità. La stessa audience vede ads completamente diversi in base al profilo comportamentale.

## STRUTTURA CAMPAGNA OTTIMALE PER ANDROMEDA (2025)
- 1 campagna per obiettivo (1 Sales, 1 Lead Gen — non di più).
- 1-3 ad set massimo per campagna (non 5-10).
- 10-40 creatività per ad set — la diversità creativa è il nuovo targeting.
- CBO (Advantage Campaign Budget): Meta distribuisce il budget in real-time tra gli ad set.
- Broad targeting: nessuna restrizione di interesse, età, genere — lascia che Andromeda trovi il pubblico via embedding.
- Evidenza: 1 ad set con 25 creatività outperforma 5 ad set con 5 creatività ciascuno di +17% conversioni e -16% costi.

## FASE DI APPRENDIMENTO (aggiornato 2025)
- Soglia ufficiale: 50 eventi di ottimizzazione per ad set in 7 giorni per uscire dalla learning phase.
- Budget minimo reale = (CPA obiettivo × 50) ÷ 7 giorni. Es: CPA €30 → serve €214/giorno minimo.
- NOVITÀ 2025: Aggiungere nuove creatività a campagne già in apprendimento NON riavvia più la learning phase completamente.
- Aumenti di budget fino a €179/giorno non riavviano il learning; oltre quella soglia rischi un reset parziale.
- Aspetta 3-4 giorni tra un aumento e il successivo; aumenta al massimo del 15-20% alla volta.
- ERRORE COMUNE: Uccidere campagne a 2-3 giorni durante la volatilità normale della learning phase.

## PERCHÉ LE CAMPAGNE SI "ROMPONO" DOPO 2-3 GIORNI
1. Volatilità learning phase (normale): CPA/ROAS oscillano molto nei primi 7 giorni — non toccare nulla.
2. Creative fatigue (causa più comune): Con Andromeda la fatigue arriva 35% prima rispetto ai sistemi precedenti. Fatigue reale: 10-21 giorni per campagne standard, 5-7 giorni per retargeting. NON si vede sempre dalla frequenza — molte ads fatigate hanno frequenza 2.0-2.5 apparentemente "ok".
3. Budget frammentato: Budget diviso su molti ad set → ogni ad set troppo debole per imparare.
4. Modifiche frequenti: Ogni cambio nei primi 7 giorni resetta l'apprendimento.
5. Saturazione audience: Con targeting stretto, esaurisci il pool disponibile in pochi giorni.
Soluzioni: Non pausare/duplicare (crea un nuovo ciclo di learning da zero). Se fatigue: aggiungi 2-3 creatività strutturalmente diverse. Se fragmentation: consolida ad set. Se learning volatility: aspetta 7 giorni. Se saturazione: passa a broad targeting.

## ADVANTAGE+ E ASC IN 2025
- Tutte le campagne di vendita sono essenzialmente Advantage+/ASC ora — è diventata la struttura default di Meta.
- Advantage+ Sales supporta fino a 50 ads per ad set, più ad set per campagna, basic audience preferences.
- Meta sta progressivamente eliminando le campagne manuali per la maggior parte degli obiettivi.
- Dati interni Meta aprile 2025: campagne con full Advantage+ automation → 12% costi per acquisto più bassi, 17% ROAS più alto.

## CREATIVITÀ: IL VERO DRIVER
- 4+ formati creativi unici in un ad set → +32% performance.
- Diversità strutturale necessaria: video (varie lunghezze), statico, carousel, UGC, demo prodotto, testimonial.
- Meta considera due ads "simili" se i primi 3 secondi sono identici — hook diversi sono obbligatori.
- Rotazione proattiva ogni 7-10 giorni (non aspettare i segnali di fatigue).
- UGC e product demo sono i format più efficaci nel 2025.

## TARGETING BROAD vs STRETTO
- Targeting broad = paese intero, nessun interesse, nessuna restrizione età/genere — è ora la best practice ufficiale Meta.
- Andromeda trova il pubblico giusto via embedding similarity alla creatività, non via targeting impostato.
- Targeting stretto: satura l'audience, rallenta il learning, va contro il funzionamento di Andromeda.

## CPM E ASTA (2025)
- CPM medio Q1 2025: €10.88 (+19.2% YoY).
- Alta embedding similarity → CPM più basso; bassa similarity → CPM più alto.
- Creatività migliore = costi inferiori oltre che performance migliori.

## CBO vs ABO
- CBO: consigliato per campagne con ads vincenti da scalare.
- ABO: utile solo nella fase di test/validazione di nuove creatività (7-14 giorni), poi migra i vincitori su CBO.
"""

@app.get("/live-ads/overview")
async def get_live_ads_overview(date_preset: str = "last_30d", since: str = "", until: str = ""):
    token = os.getenv("META_ACCESS_TOKEN")
    if not token:
        raise HTTPException(status_code=500, detail="META_ACCESS_TOKEN non trovato nel file .env")

    all_clients = storage_service.list_clients()
    clients_with_ads = [(c, storage_service.get_metadata(c["id"])) for c in all_clients]
    clients_with_ads = [(c, m) for c, m in clients_with_ads if m.get("ad_account_id", "").strip()]

    if not clients_with_ads:
        return {"clients": [], "totals": {}, "period": date_preset}

    async def fetch_account_summary(client: dict, metadata: dict) -> dict:
        base = {"client_id": client["id"], "name": client["name"]}
        ad_id = metadata["ad_account_id"].strip()
        if not ad_id.startswith("act_"):
            ad_id = f"act_{ad_id}"
        try:
            overview_params: dict = {
                "fields": "spend,impressions,reach,clicks,ctr,cpc,cpm,actions,cost_per_action_type",
                "access_token": token,
            }
            if since and until:
                import json as _json
                overview_params["time_range"] = _json.dumps({"since": since, "until": until})
            else:
                overview_params["date_preset"] = date_preset
            async with httpx.AsyncClient(timeout=30.0) as hc:
                resp = await hc.get(
                    f"https://graph.facebook.com/v19.0/{ad_id}/insights",
                    params=overview_params
                )
            data = resp.json()
            if "error" in data:
                return {**base, "error": data["error"].get("message", "API error")}
            rows = data.get("data", [])
            if not rows:
                return {**base, "error": "Nessun dato per questo periodo"}
            row = rows[0]
            actions = row.get("actions", [])
            purchase = next((a for a in actions if a["action_type"] in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
            cpa_list = row.get("cost_per_action_type", [])
            cpa_item = next((a for a in cpa_list if a["action_type"] in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
            return {
                **base,
                "ad_account_id": ad_id,
                "spend": round(float(row.get("spend") or 0), 2),
                "impressions": int(row.get("impressions") or 0),
                "reach": int(row.get("reach") or 0),
                "clicks": int(row.get("clicks") or 0),
                "ctr": round(float(row.get("ctr") or 0), 2),
                "cpc": round(float(row.get("cpc") or 0), 2),
                "cpm": round(float(row.get("cpm") or 0), 2),
                "conversioni": int(float(purchase["value"])) if purchase else 0,
                "cpa": round(float(cpa_item["value"]), 2) if cpa_item else None,
            }
        except Exception as e:
            return {**base, "error": str(e)}

    results = await asyncio.gather(*[fetch_account_summary(c, m) for c, m in clients_with_ads])
    valid = [r for r in results if "error" not in r]
    totals = {
        "spend": round(sum(r["spend"] for r in valid), 2),
        "impressions": sum(r["impressions"] for r in valid),
        "reach": sum(r["reach"] for r in valid),
        "clicks": sum(r["clicks"] for r in valid),
        "conversioni": sum(r["conversioni"] for r in valid),
        "ctr": round(sum(r["ctr"] for r in valid) / len(valid), 2) if valid else 0,
        "cpm": round(sum(r["cpm"] for r in valid) / len(valid), 2) if valid else 0,
    }
    return {"clients": list(results), "totals": totals, "period": date_preset}


@app.get("/live-ads/campaigns/{client_id}")
async def get_client_campaigns(client_id: str, date_preset: str = "last_30d", since: str = "", until: str = ""):
    metadata = storage_service.get_metadata(client_id)
    ad_id = metadata.get("ad_account_id", "").strip()
    if not ad_id:
        raise HTTPException(status_code=400, detail="Ad Account ID non configurato per questo cliente")
    token = metadata.get("meta_access_token") or os.getenv("META_ACCESS_TOKEN")
    if not token:
        raise HTTPException(status_code=500, detail="META_ACCESS_TOKEN non trovato nel file .env")
    if not ad_id.startswith("act_"):
        ad_id = f"act_{ad_id}"

    camp_params: dict = {
        "fields": "campaign_name,campaign_id,spend,impressions,reach,clicks,ctr,cpc,cpm,actions,cost_per_action_type",
        "level": "campaign",
        "filtering": json.dumps([{"field": "campaign.effective_status", "operator": "IN", "value": ["ACTIVE"]}]),
        "access_token": token,
    }
    if since and until:
        import json as _json
        camp_params["time_range"] = _json.dumps({"since": since, "until": until})
    else:
        camp_params["date_preset"] = date_preset

    async with httpx.AsyncClient(timeout=30.0) as hc:
        resp = await hc.get(
            f"https://graph.facebook.com/v19.0/{ad_id}/insights",
            params=camp_params
        )
    data = resp.json()
    if "error" in data:
        raise HTTPException(status_code=400, detail=data["error"].get("message", "Meta API error"))

    campaigns = []
    for row in data.get("data", []):
        actions = row.get("actions", [])
        purchase = next((a for a in actions if a["action_type"] in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
        cpa_list = row.get("cost_per_action_type", [])
        cpa_item = next((a for a in cpa_list if a["action_type"] in ("purchase", "offsite_conversion.fb_pixel_purchase", "omni_purchase")), None)
        campaigns.append({
            "campaign_id": row.get("campaign_id"),
            "campaign_name": row.get("campaign_name", "—"),
            "spend": round(float(row.get("spend") or 0), 2),
            "impressions": int(row.get("impressions") or 0),
            "reach": int(row.get("reach") or 0),
            "clicks": int(row.get("clicks") or 0),
            "ctr": round(float(row.get("ctr") or 0), 2),
            "cpc": round(float(row.get("cpc") or 0), 2),
            "cpm": round(float(row.get("cpm") or 0), 2),
            "conversioni": int(float(purchase["value"])) if purchase else 0,
            "cpa": round(float(cpa_item["value"]), 2) if cpa_item else None,
        })
    campaigns.sort(key=lambda x: x["spend"], reverse=True)
    return {"campaigns": campaigns, "period": date_preset, "client_id": client_id}


class LiveAdsAnalysisRequest(BaseModel):
    date_preset: str = "last_30d"
    campaigns: List[dict] = []

class LiveAdsChatRequest(BaseModel):
    messages: List[dict]
    campaigns: List[dict] = []
    date_preset: str = "last_30d"

@app.post("/live-ads/analyze/{client_id}")
async def analyze_live_ads(client_id: str, request: LiveAdsAnalysisRequest):
    metadata = storage_service.get_metadata(client_id)
    if not request.campaigns:
        raise HTTPException(status_code=400, detail="Nessuna campagna da analizzare")

    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    research_context = ""
    if research_path.exists():
        with open(research_path, "r") as f:
            research_context = f.read()[:2500]

    # Load creative intelligence if available
    intel_path = CLIENTS_DIR / client_id / "creative_intelligence.json"
    creative_intel_block = ""
    if intel_path.exists():
        with open(intel_path, "r") as f:
            intel_data = json.load(f)
        creative_intel = intel_data.get("analysis", "")[:3000]
        if creative_intel:
            creative_intel_block = f"\n\nINTELLIGENCE STORICA CREATIVITÀ (analisi di {intel_data.get('ads_count', '?')} ads — periodo {intel_data.get('period', '?')}):\n{creative_intel}"

    campaigns_text = json.dumps(request.campaigns, ensure_ascii=False, indent=2)
    research_block = f"RICERCA DI MERCATO (contesto audience):\n{research_context}" if research_context else ""
    prompt = f"""Sei un esperto senior di Meta Ads. Usa questa base di conoscenza aggiornata su Andromeda per l'analisi:
{ANDROMEDA_EXPERT_KNOWLEDGE}

CLIENTE: {metadata.get('name', '')} | Settore: {metadata.get('industry', '')}
PERIODO: {request.date_preset}

DATI CAMPAGNE ATTIVE:
{campaigns_text}

{research_block}{creative_intel_block}

Analizza e rispondi SOLO con questa struttura markdown:

## 🏆 Campagne Vincenti
Per ogni campagna che performa bene: nome, metriche chiave, perché funziona secondo Andromeda.

## ⚠️ Campagne da Ottimizzare
Per ogni campagna che brucia budget: nome, problema specifico (learning phase? fatigue? fragmentation?), azione correttiva immediata.

## 🤖 Raccomandazioni Andromeda
3-5 azioni concrete basate su come funziona l'algoritmo Meta oggi. Sii specifico per questi numeri — niente consigli generici.

## 🎯 Angoli e Hook che Funzionano
In base ai dati e alla storia delle creatività (se disponibile), che tipo di creatività/messaggio sta vincendo? Suggerisci 3 variazioni di angolo per scalare.

## ✅ Prossimi 7 Giorni — Piano d'Azione
Lista prioritizzata di 5 azioni concrete (con budget suggerito se pertinente).

Sii diretto, pratico, specifico per questi numeri."""

    analysis = await ai_service._call_ai(
        "anthropic/claude-sonnet-4-5",
        [{"role": "user", "content": prompt}]
    )
    return {"analysis": analysis, "client_id": client_id, "period": request.date_preset}


@app.post("/live-ads/chat/{client_id}")
async def chat_live_ads(client_id: str, request: LiveAdsChatRequest):
    metadata = storage_service.get_metadata(client_id)
    if not request.messages:
        raise HTTPException(status_code=400, detail="Nessun messaggio nella conversazione")

    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    research_context = ""
    if research_path.exists():
        with open(research_path, "r") as f:
            research_context = f.read()[:2500]

    # Load creative intelligence if available
    intel_path = CLIENTS_DIR / client_id / "creative_intelligence.json"
    creative_intel_block = ""
    if intel_path.exists():
        with open(intel_path, "r") as f:
            intel_data = json.load(f)
        creative_intel = intel_data.get("analysis", "")[:3000]
        if creative_intel:
            creative_intel_block = f"\n\nINTELLIGENCE STORICA CREATIVITÀ (analisi di {intel_data.get('ads_count', '?')} ads — periodo {intel_data.get('period', '?')}):\n{creative_intel}"

    campaigns_text = json.dumps(request.campaigns, ensure_ascii=False, indent=2) if request.campaigns else "Nessun dato campagna disponibile."
    research_block = f"\nRICERCA DI MERCATO (contesto audience):\n{research_context}" if research_context else ""

    system_prompt = f"""Sei un esperto senior di Meta Ads. Usa questa base di conoscenza aggiornata su Andromeda:
{ANDROMEDA_EXPERT_KNOWLEDGE}

CLIENTE: {metadata.get('name', '')} | Settore: {metadata.get('industry', '')}
PERIODO ANALIZZATO: {request.date_preset}

DATI CAMPAGNE ATTIVE:
{campaigns_text}
{research_block}{creative_intel_block}

Sei in una conversazione diretta con il marketer/imprenditore che gestisce questo cliente. Hai accesso alla storia completa delle creatività e delle campagne. Rispondi in modo conciso e pratico — applicando la conoscenza di Andromeda ai dati reali di questo account. Se ti forniscono aggiornamenti (campagna disattivata, cambio budget, nuove creatività, ecc.), adatta immediatamente i tuoi consigli. Usa emoji con parsimonia."""

    messages = [{"role": "system", "content": system_prompt}] + request.messages
    reply = await ai_service._call_ai("anthropic/claude-sonnet-4-5", messages)
    return {"message": reply, "client_id": client_id}


@app.post("/clients/{client_id}/angles")
async def get_angles(client_id: str, request: AngleRequest = AngleRequest()):
    metadata = storage_service.get_metadata(client_id)

    # 🔥 NUOVO: Carica TUTTA l'Analisi Strategica (14 sezioni)
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id,
        metadata=metadata,
        supabase_client=supabase,
        focus_areas=["brand_identity", "customer_personas", "psychographic_analysis", "reasons_to_buy", "objections", "battlecards"]
    )

    # Enrich with creative intelligence from real ad performance if available
    intel_path = CLIENTS_DIR / client_id / "creative_intelligence.json"
    if intel_path.exists():
        with open(intel_path, "r") as f:
            intel_data = json.load(f)
        creative_intel = intel_data.get("analysis", "")[:2000]
        if creative_intel:
            strategic_context += f"\n\n### INTELLIGENCE DALLE ADS REALI ({intel_data.get('period', 'storico')}, {intel_data.get('ads_count', '?')} ads analizzate):\nQuesti dati provengono dall'analisi delle inserzioni realmente mandate in pubblicazione. Usali come FONTE PRIMARIA per capire cosa ha già funzionato e cosa no. Gli angoli vincitori reali devono guidare i nuovi angoli suggeriti:\n{creative_intel}"

    # DEBUG: Log del contesto strategico per verificare cosa vede l'AI
    print(f"\n{'='*80}\n[ANGLES DEBUG] Strategic context per {client_id} ({len(strategic_context)} chars):\n{strategic_context[:3000]}\n{'='*80}\n")

    angles_data = await ai_service.generate_communication_angles(strategic_context, request.user_prompt, request.funnel_stage)

    # Save angles
    angles_path = CLIENTS_DIR / client_id / "angles.json"
    with open(angles_path, "w") as f:
        json.dump(angles_data, f, ensure_ascii=False, indent=2)
    storage_service.sync_angles(client_id, angles_data)

    return angles_data

@app.get("/clients/{client_id}/angles")
async def fetch_existing_angles(client_id: str):
    angles_path = CLIENTS_DIR / client_id / "angles.json"
    if not angles_path.exists():
        return []
    with open(angles_path, "r") as f:
        return json.load(f)


# ══════════════════════════════════════════════════════════
#  VOICE OF CUSTOMER — Review Mining automatico da URL
# ══════════════════════════════════════════════════════════

class VoCRequest(BaseModel):
    instagram_url: Optional[str] = ""      # https://www.instagram.com/handle/
    google_reviews_url: Optional[str] = "" # https://maps.google.com/... oppure https://g.page/...
    include_competitors: Optional[bool] = True  # se True cerca anche sui competitor del cliente


async def _fetch_instagram_comments(ig_handle: str, token: str, ig_user_id: str = "") -> list[str]:
    """Recupera commenti recenti dal profilo Instagram tramite Meta Graph API."""
    comments = []
    try:
        # Se abbiamo già l'IG Business Account ID del cliente, usiamo Business Discovery
        # per leggere i post del profilo target (che può essere il brand stesso o un competitor)
        async with httpx.AsyncClient(timeout=30.0) as hc:
            # Step 1: trova l'IG Business Account ID del cliente (per usare Business Discovery)
            if not ig_user_id:
                me_resp = await hc.get(
                    "https://graph.facebook.com/v19.0/me/accounts",
                    params={"access_token": token, "fields": "id,name,instagram_business_account"}
                )
                pages = me_resp.json().get("data", [])
                for page in pages:
                    iba = page.get("instagram_business_account", {})
                    if iba.get("id"):
                        ig_user_id = iba["id"]
                        break

            if not ig_user_id:
                return comments

            # Step 2: Business Discovery sul profilo target
            disc_resp = await hc.get(
                f"https://graph.facebook.com/v19.0/{ig_user_id}",
                params={
                    "fields": f"business_discovery.fields(username,media_count,media{{caption,comments_count,comments{{text,timestamp}}}})",
                    "username": ig_handle,
                    "access_token": token,
                }
            )
            disc_data = disc_resp.json()
            media_list = disc_data.get("business_discovery", {}).get("media", {}).get("data", [])
            for post in media_list[:15]:  # ultimi 15 post
                caption = post.get("caption", "")
                if caption:
                    comments.append(f"[POST] {caption[:300]}")
                for c in post.get("comments", {}).get("data", []):
                    txt = c.get("text", "").strip()
                    if txt and len(txt) > 10:
                        comments.append(txt)
    except Exception as e:
        print(f"IG fetch error for {ig_handle}: {e}")
    return comments


async def _fetch_google_reviews_text(url: str) -> str:
    """Tenta di scaricare il testo della pagina Google Maps/Reviews."""
    try:
        headers = {
            "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36",
            "Accept-Language": "it-IT,it;q=0.9",
        }
        async with httpx.AsyncClient(timeout=20.0, follow_redirects=True) as hc:
            resp = await hc.get(url, headers=headers)
            text = resp.text
            # Estrai solo il testo visibile (rimuovi HTML)
            text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.DOTALL)
            text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
            text = re.sub(r'<[^>]+>', ' ', text)
            text = re.sub(r'\s+', ' ', text).strip()
            return text[:6000]
    except Exception as e:
        print(f"Google reviews fetch error: {e}")
        return ""


@app.post("/clients/{client_id}/voc/analyze")
async def analyze_voc(client_id: str, request: VoCRequest):
    """
    Raccoglie automaticamente recensioni e commenti da Instagram e Google Reviews,
    poi usa l'AI per estrarre VoC intelligence: Golden Hooks, Pain Points, ICP signals, ecc.
    Se i dati del brand sono scarsi, include anche i competitor.
    """
    metadata = storage_service.get_metadata(client_id)
    client_name = metadata.get("name", client_id)
    token = metadata.get("meta_access_token") or os.getenv("META_ACCESS_TOKEN", "")

    collected_texts: list[str] = []
    sources_used: list[str] = []

    # ── 1. Instagram del brand ────────────────────────────────────────────────
    ig_handle = ""
    if request.instagram_url:
        ig_handle = extract_ig_handle(request.instagram_url)

    # Prova anche dai link salvati nel profilo se non fornito esplicitamente
    if not ig_handle:
        for link in metadata.get("links", []):
            url_val = link.get("url", "") if isinstance(link, dict) else str(link)
            if "instagram.com" in url_val:
                ig_handle = extract_ig_handle(url_val)
                break

    if ig_handle and token:
        ig_comments = await _fetch_instagram_comments(ig_handle, token)
        if ig_comments:
            collected_texts.extend(ig_comments)
            sources_used.append(f"Instagram @{ig_handle} ({len(ig_comments)} commenti/caption)")

    # ── 2. Google Reviews ──────────────────────────────────────────────────────
    if request.google_reviews_url:
        google_text = await _fetch_google_reviews_text(request.google_reviews_url)
        if google_text and len(google_text) > 200:
            collected_texts.append(f"[GOOGLE REVIEWS]\n{google_text}")
            sources_used.append("Google Reviews")

    # ── 3. Fallback competitor se dati scarsi ─────────────────────────────────
    if request.include_competitors and len(collected_texts) < 5:
        competitors = metadata.get("competitors", [])
        for comp in competitors[:3]:
            comp_links = comp.get("links", []) if isinstance(comp, dict) else []
            for cl in comp_links:
                cl_url = cl.get("url", "") if isinstance(cl, dict) else str(cl)
                if "instagram.com" in cl_url and token:
                    comp_handle = extract_ig_handle(cl_url)
                    if comp_handle:
                        comp_comments = await _fetch_instagram_comments(comp_handle, token)
                        if comp_comments:
                            collected_texts.extend([f"[COMPETITOR @{comp_handle}] {c}" for c in comp_comments[:20]])
                            sources_used.append(f"Instagram competitor @{comp_handle}")

    if not collected_texts:
        raise HTTPException(
            status_code=400,
            detail="Nessun dato raccolto. Assicurati di avere: 1) Un link Instagram valido, 2) Un Meta Access Token configurato in Sorgenti, 3) Oppure un link Google Reviews raggiungibile."
        )

    raw_corpus = "\n".join(collected_texts)[:9000]

    # ── 4. AI Analysis ────────────────────────────────────────────────────────
    system_prompt = f"""Sei un esperto di Voice of Customer (VoC) e copywriting per Meta Ads con 15 anni di esperienza.
Il tuo compito è analizzare commenti e recensioni reali raccolti automaticamente e trasformarli in intelligence strategica.

BRAND: {client_name}
FONTI ANALIZZATE: {', '.join(sources_used)}

NOTA IMPORTANTE: Alcuni testi potrebbero venire da competitor (marcati [COMPETITOR]).
Usali per capire cosa vogliono i clienti nel settore, ma distinguili dal brand principale.
I testi [POST] sono caption Instagram del brand — contengono il linguaggio del brand stesso.
I commenti non marcati sono del brand principale.

METODOLOGIA:
1. GOLDEN HOOKS: Frasi exact-match usate dai clienti reali che esprimono trasformazione/risultato. Da usare letteralmente nel copy.
2. PAIN POINTS: Il "prima" — frustrazioni, problemi, dolori prima del prodotto.
3. DESIDERI & OUTCOME: Il "dopo" — trasformazioni e risultati ottenuti/cercati.
4. OBIEZIONI: Dubbi, esitazioni ("inizialmente ero scettico", "avevo paura che", "non credevo").
5. TRIGGER PSICOGRAFICI: Motivazioni profonde, valori, identità del cliente ideale.
6. SOCIAL PROOF ANGLES: Storie specifiche o risultati quantificabili reali.

OUTPUT (JSON VALIDO):
{{
  "golden_hooks": ["frase esatta 1", ...],
  "pain_points": ["pain point 1", ...],
  "desires_outcomes": ["desiderio/risultato 1", ...],
  "objections": ["obiezione 1", ...],
  "psychographic_triggers": ["trigger 1", ...],
  "social_proof_angles": ["storia/dato specifico 1", ...],
  "top_copy_phrases": ["frase pronta per copy 1", ...],
  "icp_summary": "Chi è il cliente ideale in 2-3 frasi basate sui dati reali",
  "strategic_insights": "Insights su come posizionare il brand nel copy basati sui dati"
}}

Rispondi SOLO con JSON valido."""

    try:
        raw = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"DATI RACCOLTI:\n\n{raw_corpus}"}
            ],
            temperature=0.4,
            max_tokens=2000
        )
        voc_data = json_repair.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore AI: {str(e)}")

    result = {
        "data": voc_data,
        "sources": sources_used,
        "texts_count": len(collected_texts),
        "generated_at": datetime.now().isoformat(),
        "instagram_url": request.instagram_url or "",
        "google_reviews_url": request.google_reviews_url or "",
    }

    voc_path = CLIENTS_DIR / client_id / "voc_analysis.json"
    voc_path.parent.mkdir(parents=True, exist_ok=True)
    with open(voc_path, "w") as f:
        json.dump(result, f, ensure_ascii=False, indent=2)

    metadata["voc_analysis"] = result
    storage_service.save_metadata(client_id, metadata)

    return result


@app.get("/clients/{client_id}/voc")
async def get_voc(client_id: str):
    voc_path = CLIENTS_DIR / client_id / "voc_analysis.json"
    if voc_path.exists():
        with open(voc_path, "r") as f:
            return json.load(f)
    metadata = storage_service.get_metadata(client_id)
    if "voc_analysis" in metadata:
        return metadata["voc_analysis"]
    return {"data": None, "generated_at": None}


# ══════════════════════════════════════════════════════════
#  COPY GENERATOR — Framework-based Meta Ads copy
# ══════════════════════════════════════════════════════════

class CopyRequest(BaseModel):
    copy_type: str = "caption"       # caption | meta_ads | email | headline | bio
    framework: str = ""              # opzionale — PAS | AIDA | BAB | ecc.
    awareness_level: str = ""        # opzionale — unaware | problem_aware | solution_aware | product_aware | most_aware
    angle_title: str = ""
    angle_description: str = ""
    user_instructions: str = ""
    variations: int = 1

@app.post("/clients/{client_id}/copy/generate")
async def generate_copy(client_id: str, request: CopyRequest):
    """Genera copy in base al tipo selezionato, framework e livello di consapevolezza."""
    metadata = storage_service.get_metadata(client_id)

    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id,
        metadata=metadata,
        supabase_client=supabase,
        focus_areas=["brand_identity", "brand_voice", "customer_personas", "reviews_voc", "objections", "reasons_to_buy", "psychographic_analysis"]
    )

    from .knowledge_loader import (
        get_writing_knowledge, get_awareness_context,
        get_single_framework_knowledge, get_inspirational_stair,
        get_copy_type_knowledge
    )

    # ═══════════════════════════════════════════════════════
    # LAYER 1 — Conoscenza base (SEMPRE presente)
    # Regole di scrittura Vignali + Modelli narrativi Dosio
    # L'AI deve conoscere queste regole PRIMA di scrivere qualsiasi cosa
    # ═══════════════════════════════════════════════════════
    writing_knowledge = get_writing_knowledge()

    # ═══════════════════════════════════════════════════════
    # LAYER 2 — Livello di consapevolezza (se selezionato)
    # Dice all'AI COME parlare al target in base a dove si trova
    # ═══════════════════════════════════════════════════════
    awareness_section = ""
    if request.awareness_level:
        awareness_section = get_awareness_context(request.awareness_level)

    # ═══════════════════════════════════════════════════════
    # LAYER 3 — Framework (se selezionato)
    # Detta la STRUTTURA del copy — l'AI lo segue alla lettera
    # ═══════════════════════════════════════════════════════
    framework_section = ""
    if request.framework:
        if request.framework == "INSPIRATIONAL_STAIR":
            framework_section = get_inspirational_stair()
        else:
            fw = get_single_framework_knowledge(request.framework)
            if fw:
                framework_section = fw

    # ═══════════════════════════════════════════════════════
    # LAYER 4 — Tipo di copy (canale + vincoli tecnici)
    # NON detta la struttura — dice solo DOVE verrà pubblicato
    # e i vincoli tecnici di quel canale
    # ═══════════════════════════════════════════════════════
    type_knowledge = get_copy_type_knowledge(request.copy_type)

    # ── Contesto angolo ──
    angle_ctx = ""
    if request.angle_title:
        angle_ctx = f"ANGOLO COMUNICATIVO: {request.angle_title}"
        if request.angle_description:
            angle_ctx += f"\n{request.angle_description}"

    # ── Istruzioni utente ──
    user_instr = request.user_instructions or ""

    # ── Label leggibili per il tipo di copy ──
    TYPE_LABELS = {
        "caption": "CAPTION INSTAGRAM",
        "meta_ads": "META ADS COPY",
        "email": "EMAIL MARKETING",
        "headline": "HEADLINE / HOOK",
        "bio": "BIO PROFILO",
    }
    type_label = TYPE_LABELS.get(request.copy_type, request.copy_type.upper())

    # ── Limiti di lunghezza per tipo ──
    WORD_LIMITS = {
        "caption": 150,
        "meta_ads": 125,
        "email": 200,
        "headline": 0,  # 10 varianti, non un testo
        "bio": 0,       # limiti in caratteri, non parole
    }
    word_limit = WORD_LIMITS.get(request.copy_type, 200)

    # ── Fasi operative per ogni framework ──
    # Ogni framework ha istruzioni PRECISE su cosa scrivere in ogni fase.
    # L'AI deve seguire esattamente questa sequenza — ogni fase deve essere
    # riconoscibile nel testo finale (senza etichette visibili).
    FRAMEWORK_PHASES = {
        "AIDA": (
            "Scrivi il copy in ESATTAMENTE 4 blocchi, in quest'ordine:\n"
            "1. ATTENTION (1-2 righe): Domanda provocatoria, dato sorprendente o affermazione contro-intuitiva che ferma lo scroll.\n"
            "2. INTEREST (2-3 righe): Fatti concreti e specifici. Perché questo è rilevante per il target ORA.\n"
            "3. DESIRE (2-3 righe): Social proof, risultati reali, identità aspirazionale. Il target deve VOLERE quello che offri.\n"
            "4. ACTION (1 riga): CTA specifica e diretta. Non 'Scopri di più' — scrivi cosa deve fare concretamente."
        ),
        "PAS": (
            "Scrivi il copy in ESATTAMENTE 3 blocchi, in quest'ordine:\n"
            "1. PROBLEM (1-2 righe): Descrivi il problema con il linguaggio esatto del target. Specifico, non generico.\n"
            "2. AGITATE (2-3 righe): Amplifica le conseguenze reali. Costi nascosti, frustrazioni quotidiane, cosa perde ogni giorno. Tensione crescente.\n"
            "3. SOLVE (2-3 righe): La soluzione arriva come sollievo. Chiara, diretta, collegata al problema. CTA finale."
        ),
        "BAB": (
            "Scrivi il copy in ESATTAMENTE 3 blocchi, in quest'ordine:\n"
            "1. BEFORE (2-3 righe): Descrivi la situazione dolorosa attuale del target. Dettagli specifici e riconoscibili.\n"
            "2. AFTER (2-3 righe): Dipingi la vita DOPO la soluzione. Dettagli vividi e desiderabili. Come si sente? Cosa fa di diverso?\n"
            "3. BRIDGE (1-2 righe): Il prodotto/servizio è il ponte naturale tra le due realtà. CTA."
        ),
        "FAB": (
            "Scrivi il copy in ESATTAMENTE 3 blocchi, in quest'ordine:\n"
            "1. FEATURES (1-2 righe): Caratteristiche chiave del prodotto/servizio — quelle che differenziano.\n"
            "2. ADVANTAGES (1-2 righe): Vantaggi concreti e misurabili che derivano da quelle feature.\n"
            "3. BENEFITS (2-3 righe): Benefici EMOTIVI nella vita reale del cliente. Non tecnici — identitari, esperienziali. CTA."
        ),
        "ACCA": (
            "Scrivi il copy in ESATTAMENTE 4 blocchi, in quest'ordine:\n"
            "1. AWARENESS (1-2 righe): Presenta il problema con uno scenario riconoscibile. Non assumere che il target lo conosca già.\n"
            "2. COMPREHENSION (2-3 righe): Perché questo problema è importante per LUI specificamente. Dati, conseguenze concrete.\n"
            "3. CONVICTION (2-3 righe): Prove che la soluzione funziona. Risultati numerici, casi studio, testimonial. Non opinioni — fatti.\n"
            "4. ACTION (1 riga): CTA a bassa frizione — il minor commitment possibile (call, demo, prova)."
        ),
        "SSS": (
            "Scrivi il copy come una STORIA in ESATTAMENTE 3 blocchi:\n"
            "1. STAR (1-2 righe): Presenta il protagonista — un cliente reale o archetipo credibile. Nome, età, situazione. Deve essere 'uno come me' per il target.\n"
            "2. STORY (3-4 righe): Racconta la sua storia. Il dolore, le frustrazioni, i tentativi falliti. Dettagli specifici che creano immedesimazione.\n"
            "3. SOLUTION (1-2 righe): La scoperta del prodotto/servizio come svolta naturale. Risultato credibile. CTA."
        ),
        "4C": (
            "Scrivi un copy che rispetta TUTTE e 4 queste qualità contemporaneamente:\n"
            "- CHIARO: messaggio comprensibile immediatamente, zero ambiguità\n"
            "- CONCISO: ogni parola ha un motivo, zero grasso\n"
            "- CREDIBILE: ogni claim supportato da prove (dati, testimonial, garanzie)\n"
            "- COMPELLING (irresistibile): il lettore non può ignorare l'offerta\n"
            "Il copy deve essere breve, diretto e potente. Ogni frase supera il test di tutte e 4 le C."
        ),
        "4U": (
            "Scrivi un copy che rispetta TUTTE e 4 queste qualità contemporaneamente:\n"
            "- UTILE: beneficio tangibile e immediatamente percepito\n"
            "- URGENTE: ragione reale per agire ORA (scarsità vera, costo dell'inazione)\n"
            "- UNICO: angolo o claim che nessun competitor sta usando\n"
            "- ULTRA-SPECIFICO: numeri precisi, dati concreti. 'Risparmi 3 ore' non 'Risparmi tempo'\n"
            "Il copy deve essere breve, tagliente, con dati specifici. Ogni frase supera il test di tutte e 4 le U."
        ),
        "5_OBIEZIONI": (
            "Identifica le 2-3 obiezioni più forti per QUESTO target e QUESTO prodotto tra:\n"
            "- 'Non ho tempo' / 'Non ho soldi' / 'Non funzionerà per me' / 'Non ti credo' / 'Non ne ho bisogno'\n"
            "Smontale in modo indiretto e conversazionale (non difensivo). Usa prove concrete: testimonial, dati, garanzie.\n"
            "Struttura: Hook → Obiezione 1 smontata → Obiezione 2 smontata → Obiezione 3 smontata → CTA."
        ),
        "3_MOTIVI": (
            "Scrivi il copy rispondendo a ESATTAMENTE 3 domande, in quest'ordine:\n"
            "1. PERCHÉ SEI IL MIGLIORE? Differenziatori reali e verificabili — non aggettivi vuoti.\n"
            "2. PERCHÉ DOVREI CREDERTI? Social proof forte: numeri precisi, nomi reali, risultati misurabili.\n"
            "3. PERCHÉ DOVREI COMPRARE ADESSO? Scarsità reale o costo concreto dell'inazione.\n"
            "Ogni risposta è un blocco di 2-3 righe. CTA finale."
        ),
        "E_QUINDI": (
            "Parti dall'affermazione principale del brand/prodotto e applica il processo 'E quindi?' 3-5 volte\n"
            "fino ad arrivare al beneficio emotivo finale. USA L'ULTIMA RISPOSTA come corpo del copy.\n"
            "Non usare la versione superficiale — scava fino all'identità e al desiderio profondo del target.\n"
            "Il copy finale è breve e colpisce il beneficio emotivo reale, non quello funzionale."
        ),
        "HOOK_BODY_CTA": (
            "Scrivi il copy in ESATTAMENTE 3 blocchi:\n"
            "1. HOOK (1 riga): Frase devastante che ferma lo scroll in 0.3 secondi. Domanda, dato shock, provocazione.\n"
            "2. BODY (3-5 righe): Sviluppo con prove, storia breve o benefici. Supporta l'hook senza ripeterlo.\n"
            "3. CTA (1 riga): Azione specifica. Non 'Scopri di più' — scrivi cosa deve fare e cosa ottiene."
        ),
        "INSPIRATIONAL_STAIR": (
            "Segui le 5 fasi neurochimiche di Borzacchiello nell'ordine esatto:\n"
            "1. CORTISOLO (paura/urgenza): Attiva la minaccia — cosa rischia se non agisce?\n"
            "2. DOPAMINA (ricompensa/visione): Mostra il futuro desiderabile — cosa ottiene?\n"
            "3. OSSITOCINA (connessione/fiducia): Crea vicinanza — 'ti capisco, ci sono passato'\n"
            "4. ENDORFINA (sollievo/piacere): Il momento di liberazione — la soluzione concreta\n"
            "5. SEROTONINA (sicurezza/azione): Conferma finale — CTA con sicurezza e urgenza positiva"
        ),
    }

    # ═══════════════════════════════════════════════════════
    # RAG DA NOTION — Esempi gold standard (copy a 5 stelle)
    # Costo: zero (è una chiamata API Notion, non AI)
    # ═══════════════════════════════════════════════════════
    swipe_context = ""
    if _notion_available and notion_service:
        try:
            swipe_file_examples = await notion_service.get_vault_examples(NOTION_COPY_VAULT_DB_ID)
            if swipe_file_examples:
                swipe_context = (
                    f"\n{'='*60}\nESEMPI GOLD STANDARD (Copy valutati 5 stelle)\n{'='*60}\n"
                    "Prendi ispirazione dalla qualità e stile di questi copy vincenti, "
                    "ma NON copiarli — adattali al cliente e all'angolo richiesto:\n"
                    f"{swipe_file_examples}\n"
                )
        except Exception as e:
            print(f"⚠️ Notion vault copy non disponibile: {e}")

    # ══════════════════════════════════════════════════════════
    # SYSTEM PROMPT — Background knowledge (contesto)
    # ══════════════════════════════════════════════════════════
    sep = "=" * 60
    sys_parts = []

    sys_parts.append(
        "Sei un copywriter. Scrivi copy breve, diretto, che converte. Odi i muri di testo.\n\n"
        "⚠️ PRIMA DI SCRIVERE — LEGGI I DATI:\n"
        "Leggi CUSTOMER PERSONAS nei dati del cliente: quelle sono le persone a cui parli.\n"
        "Leggi BRAND IDENTITY: quella è l'azienda/professionista che vende.\n"
        "Il copy deve parlare al target descritto nelle personas, usando i loro dolori, desideri e linguaggio."
    )

    # Contesto strategico del cliente
    sys_parts.append(f"{sep}\nDATI DEL CLIENTE\n{sep}\n{strategic_context}")

    # Regole di scrittura (sempre)
    sys_parts.append(f"{sep}\nREGOLE DI SCRITTURA\n{sep}\n{writing_knowledge}")

    # Awareness (se selezionato)
    if awareness_section:
        sys_parts.append(f"{sep}\nLIVELLO DI CONSAPEVOLEZZA DEL TARGET\n{sep}\n{awareness_section}")

    # Framework knowledge di approfondimento (se selezionato)
    if framework_section:
        sys_parts.append(f"{sep}\nAPPROFONDIMENTO FRAMEWORK\n{sep}\n{framework_section}")

    # Vincoli tecnici del canale
    if type_knowledge:
        sys_parts.append(f"{sep}\nVINCOLI CANALE: {type_label}\n{sep}\n{type_knowledge}")

    # Esempi gold standard dal vault Notion (se disponibili)
    if swipe_context:
        sys_parts.append(swipe_context)

    system_prompt = "\n\n".join(sys_parts)

    # ══════════════════════════════════════════════════════════
    # USER MESSAGE — Istruzioni operative (l'ordine vero)
    # Questo è ciò che l'AI legge per ULTIMO prima di generare.
    # Deve essere chiaro, diretto, specifico.
    # ══════════════════════════════════════════════════════════
    user_parts = []
    user_parts.append(f"Scrivi un {type_label}.")

    if word_limit > 0:
        user_parts.append(f"MASSIMO {word_limit} PAROLE. Se ne scrivi di più, il copy è inutilizzabile. Sii conciso.")

    if request.framework:
        phases = FRAMEWORK_PHASES.get(request.framework, "")
        if phases:
            user_parts.append(
                f"⚠️ STRUTTURA OBBLIGATORIA — FRAMEWORK {request.framework}:\n"
                f"{phases}\n\n"
                f"REGOLE FRAMEWORK:\n"
                f"- Segui ESATTAMENTE questa sequenza di fasi. Non saltarne nessuna.\n"
                f"- Ogni fase deve essere riconoscibile nel testo (ma senza etichette visibili come [ATTENTION] o [PROBLEM]).\n"
                f"- NON mescolare le fasi. NON aggiungere fasi extra.\n"
                f"- Il framework è la STRUTTURA del copy. Non inventare una struttura diversa."
            )

    if request.awareness_level:
        user_parts.append(f"TONO: adattato al livello {request.awareness_level.replace('_', ' ').upper()} (vedi knowledge). Questo influenza il tono, NON la struttura.")

    if request.angle_title:
        user_parts.append(f"ANGOLO: {request.angle_title}")
        if request.angle_description:
            user_parts.append(f"DESCRIZIONE ANGOLO: {request.angle_description}")

    if user_instr:
        user_parts.append(f"ISTRUZIONI EXTRA: {user_instr}")

    user_parts.append("Rispondi SOLO con il copy. Zero premesse, zero commenti, zero spiegazioni.")

    user_msg = "\n\n".join(user_parts)

    try:
        raw = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_msg}
            ],
            temperature=0.6,
            max_tokens=1500
        )
        return {"copy_text": raw.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore AI: {str(e)}")


# ══════════════════════════════════════════════════════════
#  HELPER: assembla tutto il contesto del cliente
# ══════════════════════════════════════════════════════════

def _load_full_client_context(client_id: str, metadata: dict) -> str:
    """Assembla tutta l'intelligence disponibile per un cliente in un unico testo di contesto."""
    parts = []

    # Brand identity
    bi = metadata.get("brand_identity", {})
    if bi:
        parts.append(f"BRAND IDENTITY:\nTono: {bi.get('tone','')}\nVisual: {bi.get('visuals','')}\nColori: {', '.join(bi.get('colors',[]))}")

    # Research
    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    if research_path.exists():
        parts.append("RICERCA DI MERCATO:\n" + research_path.read_text()[:2000])

    # VoC
    voc_path = CLIENTS_DIR / client_id / "voc_analysis.json"
    if voc_path.exists():
        voc = json.loads(voc_path.read_text()).get("data", {})
        if voc:
            hooks = voc.get("golden_hooks", [])[:4]
            pains = voc.get("pain_points", [])[:4]
            desires = voc.get("desires_outcomes", [])[:4]
            objections = voc.get("objections", [])[:3]
            parts.append(f"VOICE OF CUSTOMER (VoC):\nGolden Hooks: {'; '.join(hooks)}\nPain Points: {'; '.join(pains)}\nDesideri: {'; '.join(desires)}\nObiezioni: {'; '.join(objections)}")

    # Psychographic
    psych_path = CLIENTS_DIR / client_id / "psychographic.json"
    if psych_path.exists():
        psych = json.loads(psych_path.read_text()).get("data", {})
        if psych:
            p1 = psych.get("level_1_primary", {})
            parts.append(f"ANALISI PSICOGRAFICA:\nDesideri primari: {p1.get('desires','')}\nPaure primarie: {p1.get('fears','')}\nTrigger inconsci: {psych.get('level_3_unconscious',{}).get('archetypes','')}")

    # Battlecards
    bc_path = CLIENTS_DIR / client_id / "battlecards.json"
    if bc_path.exists():
        bc = json.loads(bc_path.read_text()).get("data", [])
        if bc:
            summary = "; ".join([f"{b.get('competitor_name','')}: {b.get('our_advantage','')}" for b in bc[:3]])
            parts.append(f"VANTAGGI VS COMPETITOR: {summary}")

    # Creative intelligence
    intel_path = CLIENTS_DIR / client_id / "creative_intelligence.json"
    if intel_path.exists():
        intel = json.loads(intel_path.read_text()).get("analysis", "")
        if intel:
            parts.append(f"INTELLIGENCE ADS REALI:\n{intel[:1000]}")

    return "\n\n".join(parts)


# ══════════════════════════════════════════════════════════
#  COMPETITOR BATTLECARDS
# ══════════════════════════════════════════════════════════

@app.post("/clients/{client_id}/battlecards")
async def generate_battlecards(client_id: str):
    """Genera schede competitive strutturate per ogni competitor del cliente."""
    metadata = storage_service.get_metadata(client_id)
    client_name = metadata.get("name", client_id)
    competitors = metadata.get("competitors", [])

    if not competitors:
        raise HTTPException(status_code=400, detail="Nessun competitor salvato nel profilo. Aggiungili nella sezione Sorgenti.")

    # Carica contesto strategico completo dal Supabase/file locale
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id, metadata=metadata, supabase_client=supabase,
        focus_areas=["brand_identity", "brand_values", "product_portfolio", "reasons_to_buy", "customer_personas"]
    )

    research_text = ""
    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    if research_path.exists():
        research_text = research_path.read_text()[:3000]

    comp_list = []
    for c in competitors:
        name = c.get("name", "") if isinstance(c, dict) else str(c)
        links = c.get("links", []) if isinstance(c, dict) else []
        urls = [l.get("url","") if isinstance(l,dict) else str(l) for l in links]
        comp_list.append(f"- {name}: {', '.join(urls)}")

    system_prompt = f"""Sei un analista strategico esperto di marketing competitivo e Meta Ads.
Analizza i competitor di {client_name} e crea Battlecard strutturate per ogni concorrente.

DATI DEL CLIENTE:
{strategic_context}

BRAND CLIENTE: {client_name}
COMPETITOR DA ANALIZZARE:
{chr(10).join(comp_list)}

RICERCA DI MERCATO DISPONIBILE:
{research_text[:2000]}

Per ogni competitor genera una battlecard con:
1. POSIZIONAMENTO: Come si posiziona nel mercato
2. PUNTI DI FORZA: Cosa fanno bene
3. PUNTI DEBOLI: Vulnerabilità da sfruttare
4. ANGOLI CHE USANO: Tipo di messaggi/angoli che usano nelle loro comunicazioni
5. NOSTRO VANTAGGIO: Come {client_name} si differenzia e batte questo competitor
6. COME RISPONDERE: Angoli e messaggi da usare per sottrarre clienti a questo competitor

OUTPUT JSON:
{{
  "battlecards": [
    {{
      "competitor_name": "nome",
      "positioning": "come si posiziona",
      "strengths": ["punto 1", "punto 2"],
      "weaknesses": ["debolezza 1", "debolezza 2"],
      "their_angles": ["angolo che usano 1", "angolo 2"],
      "our_advantage": "vantaggio chiave di {client_name} vs questo competitor",
      "counter_strategy": "come batterlo con il nostro marketing",
      "steal_customers_hooks": ["hook per sottrarre i loro clienti 1", "hook 2"]
    }}
  ],
  "overall_competitive_position": "sintesi del posizionamento complessivo di {client_name} nel mercato"
}}

Rispondi SOLO con JSON valido."""

    try:
        raw = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": f"Genera le battlecard per {client_name}"}],
            temperature=0.5,
            max_tokens=3000
        )
        data = json_repair.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore AI: {str(e)}")

    result = {"data": data.get("battlecards", data), "overall": data.get("overall_competitive_position",""), "generated_at": datetime.now().isoformat()}
    bc_path = CLIENTS_DIR / client_id / "battlecards.json"
    bc_path.parent.mkdir(parents=True, exist_ok=True)
    bc_path.write_text(json.dumps(result, ensure_ascii=False, indent=2))
    metadata["battlecards"] = result
    storage_service.save_metadata(client_id, metadata)
    return result

@app.get("/clients/{client_id}/battlecards")
async def get_battlecards(client_id: str):
    bc_path = CLIENTS_DIR / client_id / "battlecards.json"
    if bc_path.exists():
        return json.loads(bc_path.read_text())
    metadata = storage_service.get_metadata(client_id)
    return metadata.get("battlecards", {"data": None, "generated_at": None})


# ══════════════════════════════════════════════════════════
#  PSYCHOGRAPHIC ANALYSIS — 3 livelli di profondità
# ══════════════════════════════════════════════════════════

@app.post("/clients/{client_id}/psychographic")
async def generate_psychographic(client_id: str):
    """Analisi psicografica a 3 livelli di profondità del cliente ideale."""
    metadata = storage_service.get_metadata(client_id)
    client_name = metadata.get("name", client_id)

    # Carica contesto strategico completo dal Supabase/file locale
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id, metadata=metadata, supabase_client=supabase,
        focus_areas=["brand_identity", "customer_personas", "reviews_voc", "reasons_to_buy", "objections", "brand_voice"]
    )

    system_prompt = f"""Sei uno psicologo del consumatore e stratega di marketing con 20 anni di esperienza.
Il tuo compito è creare un'analisi psicografica profonda a 3 livelli del cliente ideale di {client_name}.

DATI DEL CLIENTE:
{strategic_context}

STRUTTURA ANALISI:

LIVELLO 1 — PSICOGRAFIA PRIMARIA (consapevole, dichiarata):
- Cosa dice di volere
- Obiettivi espliciti
- Desideri consapevoli
- Pain points dichiarati

LIVELLO 2 — PSICOGRAFIA SECONDARIA (emotiva, identitaria):
- Come vuole essere visto dagli altri
- Identità aspirazionale
- Valori e credenze core
- Tribù di appartenenza (chi sono "loro")
- Paure sociali (giudizio, fallimento, esclusione)

LIVELLO 3 — PSICOGRAFIA TERZIARIA (inconscia, archetipica):
- Archetipi psicologici attivi (Eroe, Ribelle, Amante, Saggio, ecc.)
- Trigger inconsci che guidano l'acquisto
- La vera ragione per cui compra (spesso diversa da quella dichiarata)
- Narrative interiori ("sono il tipo di persona che...")
- Come il prodotto risolve un conflitto identitario profondo

OUTPUT JSON:
{{
  "level_1_primary": {{
    "desires": "cosa vuole consapevolmente",
    "explicit_goals": ["goal 1", "goal 2"],
    "declared_pain_points": ["pain 1", "pain 2"],
    "what_they_say": "come descrivono il loro problema"
  }},
  "level_2_secondary": {{
    "aspirational_identity": "chi vuole essere/sembrare",
    "core_values": ["valore 1", "valore 2"],
    "tribe": "a quale gruppo vuole appartenere",
    "social_fears": ["paura sociale 1", "paura 2"],
    "identity_statement": "Sono il tipo di persona che..."
  }},
  "level_3_unconscious": {{
    "archetypes": "archetipi psicologici dominanti",
    "real_purchase_reason": "la vera ragione inconscia per cui compra",
    "unconscious_triggers": ["trigger 1", "trigger 2"],
    "identity_conflict": "quale conflitto interiore risolve il prodotto",
    "deepest_fear": "la paura più profonda da non nominare mai esplicitamente"
  }},
  "copywriting_implications": {{
    "words_that_activate": ["parola/frase che risuona 1", "parola 2"],
    "words_to_avoid": ["parola da evitare 1", "parola 2"],
    "best_hook_types": ["tipo di hook 1", "tipo 2"],
    "narrative_arc": "la storia che risuona di più con questa psicografia"
  }}
}}

Rispondi SOLO con JSON valido."""

    try:
        raw = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": f"Genera analisi psicografica per il cliente ideale di {client_name}"}],
            temperature=0.5,
            max_tokens=2500
        )
        data = json_repair.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore AI: {str(e)}")

    result = {"data": data, "generated_at": datetime.now().isoformat()}
    psych_path = CLIENTS_DIR / client_id / "psychographic.json"
    psych_path.parent.mkdir(parents=True, exist_ok=True)
    psych_path.write_text(json.dumps(result, ensure_ascii=False, indent=2))
    metadata["psychographic"] = result
    storage_service.save_metadata(client_id, metadata)
    return result

@app.get("/clients/{client_id}/psychographic")
async def get_psychographic(client_id: str):
    psych_path = CLIENTS_DIR / client_id / "psychographic.json"
    if psych_path.exists():
        return json.loads(psych_path.read_text())
    metadata = storage_service.get_metadata(client_id)
    return metadata.get("psychographic", {"data": None, "generated_at": None})


# ══════════════════════════════════════════════════════════
#  VISUAL BRIEF — Brief per designer/videomaker
# ══════════════════════════════════════════════════════════

@app.post("/clients/{client_id}/visual-brief")
async def generate_visual_brief(client_id: str):
    """Genera un brief visivo completo per designer e videomaker."""
    metadata = storage_service.get_metadata(client_id)
    client_name = metadata.get("name", client_id)

    # Carica contesto strategico completo dal Supabase/file locale
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id, metadata=metadata, supabase_client=supabase,
        focus_areas=["brand_identity", "brand_voice", "brand_values", "customer_personas", "product_portfolio", "content_matrix"]
    )

    system_prompt = f"""Sei un Creative Director con esperienza in brand identity e advertising per Meta/TikTok.
Genera un Visual Brief professionale per {client_name} che possa essere consegnato direttamente a designer e videomaker.

DATI DEL CLIENTE:
{strategic_context}

Il brief deve coprire:
1. MOOD & AESTHETIC: Atmosfera visiva del brand
2. COLOR PALETTE: Colori principali e come usarli
3. TYPOGRAPHY DIRECTION: Stile tipografico
4. CREATIVE FORMATS: Specifiche per ogni formato Meta Ads
5. DO's: Cosa includere sempre nelle creatività
6. DON'Ts: Cosa non fare mai
7. TARGET VISUAL CUES: Elementi visivi che risuonano con il target
8. HOOK VISIVI: Prime 3 secondi del video — cosa deve succedere a schermo
9. REFERENCE AESTHETIC: Tipo di stile/mood di riferimento (es: "clean minimalista", "energico urban", ecc.)
10. VIDEO BRIEF: Struttura video consigliata (timing preciso)

OUTPUT JSON:
{{
  "mood_aesthetic": "descrizione dell'atmosfera visiva",
  "color_palette": {{
    "primary": ["#colore1"],
    "secondary": ["#colore2"],
    "usage_notes": "come usare i colori nelle ads"
  }},
  "typography": "direzione tipografica",
  "dos": ["cosa fare 1", "cosa fare 2"],
  "donts": ["cosa NON fare 1", "cosa NON fare 2"],
  "target_visual_cues": ["elemento visivo che risuona 1", "elemento 2"],
  "visual_hooks_3sec": ["hook visivo per i primi 3 secondi 1", "hook 2"],
  "reference_aesthetic": "stile di riferimento",
  "formats": {{
    "stories_9x16": "brief specifico per Stories",
    "feed_4x5": "brief specifico per Feed",
    "reels": "brief specifico per Reels"
  }},
  "video_structure": {{
    "0_3s": "cosa succede nei primi 3 secondi",
    "3_15s": "sviluppo 3-15 secondi",
    "15_30s": "chiusura e CTA 15-30 secondi"
  }}
}}

Rispondi SOLO con JSON valido."""

    try:
        raw = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": f"Genera il visual brief per {client_name}"}],
            temperature=0.6,
            max_tokens=2000
        )
        data = json_repair.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore AI: {str(e)}")

    result = {"data": data, "generated_at": datetime.now().isoformat()}
    vb_path = CLIENTS_DIR / client_id / "visual_brief.json"
    vb_path.parent.mkdir(parents=True, exist_ok=True)
    vb_path.write_text(json.dumps(result, ensure_ascii=False, indent=2))
    metadata["visual_brief"] = result
    storage_service.save_metadata(client_id, metadata)
    return result

@app.get("/clients/{client_id}/visual-brief")
async def get_visual_brief(client_id: str):
    vb_path = CLIENTS_DIR / client_id / "visual_brief.json"
    if vb_path.exists():
        return json.loads(vb_path.read_text())
    metadata = storage_service.get_metadata(client_id)
    return metadata.get("visual_brief", {"data": None, "generated_at": None})


# ══════════════════════════════════════════════════════════
#  SEASONALITY ROADMAP — Calendario angoli e offerte
# ══════════════════════════════════════════════════════════

@app.post("/clients/{client_id}/seasonality")
async def generate_seasonality(client_id: str):
    """Genera una roadmap stagionale con angoli e offerte per ogni periodo dell'anno."""
    metadata = storage_service.get_metadata(client_id)
    client_name = metadata.get("name", client_id)
    industry = metadata.get("industry", "")

    # Carica contesto strategico completo dal Supabase/file locale
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id, metadata=metadata, supabase_client=supabase,
        focus_areas=["brand_identity", "product_portfolio", "customer_personas", "seasonal_roadmap", "content_matrix"]
    )

    system_prompt = f"""Sei un esperto di marketing stagionale e pianificazione campagne Meta Ads.
Crea una Seasonality Roadmap completa per {client_name} nel settore {industry or "analizzato"}.

DATI DEL CLIENTE:
{strategic_context}

Per ogni mese/periodo identifica:
- Evento/stagione rilevante per il settore
- Angolo di comunicazione consigliato
- Tipo di offerta/promozione da fare
- Urgency trigger da usare nel copy
- Formato creativo consigliato

OUTPUT JSON:
{{
  "year_overview": "sintesi strategica dell'anno",
  "months": [
    {{
      "month": "Gennaio",
      "season_context": "contesto stagionale",
      "key_events": ["evento 1", "evento 2"],
      "recommended_angle": "angolo di comunicazione consigliato",
      "offer_type": "tipo di offerta/promozione",
      "urgency_trigger": "leva di urgenza da usare",
      "creative_format": "formato creativo consigliato",
      "budget_priority": "alta/media/bassa"
    }}
  ],
  "peak_periods": ["periodo di picco 1", "periodo 2"],
  "dead_periods": ["periodo morto da evitare o usare per brand awareness"],
  "annual_strategy": "strategia annuale in 3 punti chiave"
}}

Rispondi SOLO con JSON valido con tutti i 12 mesi."""

    try:
        raw = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": f"Genera la seasonality roadmap per {client_name}"}],
            temperature=0.5,
            max_tokens=3000
        )
        data = json_repair.loads(raw)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Errore AI: {str(e)}")

    result = {"data": data, "generated_at": datetime.now().isoformat()}
    sea_path = CLIENTS_DIR / client_id / "seasonality.json"
    sea_path.parent.mkdir(parents=True, exist_ok=True)
    sea_path.write_text(json.dumps(result, ensure_ascii=False, indent=2))
    metadata["seasonality"] = result
    storage_service.save_metadata(client_id, metadata)
    return result

@app.get("/clients/{client_id}/seasonality")
async def get_seasonality(client_id: str):
    sea_path = CLIENTS_DIR / client_id / "seasonality.json"
    if sea_path.exists():
        return json.loads(sea_path.read_text())
    metadata = storage_service.get_metadata(client_id)
    return metadata.get("seasonality", {"data": None, "generated_at": None})


# ══════════════════════════════════════════════════════════
#  EXPORT — Report completo formattato per il cliente
# ══════════════════════════════════════════════════════════

@app.get("/clients/{client_id}/export")
async def export_client_report(client_id: str):
    """Genera un report HTML completo formattato, pronto da stampare come PDF e presentare al cliente."""
    from fastapi.responses import HTMLResponse
    import html as html_mod
    metadata = storage_service.get_metadata(client_id)
    client_name = metadata.get("name", client_id)
    industry = metadata.get("industry", "")
    today = datetime.now().strftime("%d %B %Y")

    # ── Carica l'analisi completa da Supabase ──
    analysis = storage_service.get_complete_analysis(client_id) or {}

    # ── Helper generici per renderizzare qualsiasi struttura dati ──
    def esc(text):
        """Escape HTML."""
        if not text:
            return "—"
        return html_mod.escape(str(text)).replace("\n", "<br>")

    def render_value(value, depth=0):
        """Renderizza qualsiasi valore in HTML leggibile."""
        if value is None:
            return "<span style='color:#999'>—</span>"
        if isinstance(value, str):
            return f"<p style='font-size:13px;color:#333;margin-bottom:6px'>{esc(value)}</p>"
        if isinstance(value, (int, float, bool)):
            return f"<span style='font-weight:600'>{value}</span>"
        if isinstance(value, list):
            if not value:
                return "<span style='color:#999'>—</span>"
            # Lista di stringhe semplici
            if all(isinstance(i, str) for i in value):
                items = "".join(
                    f"<div style='display:flex;gap:8px;align-items:flex-start;margin-bottom:6px;font-size:13px'>"
                    f"<span style='width:7px;height:7px;border-radius:50%;background:#ff9e1c;flex-shrink:0;margin-top:5px'></span>"
                    f"<span>{esc(i)}</span></div>" for i in value if i
                )
                return items or "<span style='color:#999'>—</span>"
            # Lista di oggetti
            parts = []
            for i, item in enumerate(value):
                if isinstance(item, dict):
                    # Prova a usare un campo titolo se esiste
                    title_key = next((k for k in ["name", "title", "nome", "titolo", "competitor_name", "persona_name", "month"] if k in item), None)
                    if title_key:
                        parts.append(f"<div style='background:#f8f9fa;border-radius:8px;padding:14px;margin-bottom:10px'>")
                        parts.append(f"<h4 style='font-size:14px;font-weight:700;color:#003366;margin-bottom:10px'>{esc(item[title_key])}</h4>")
                        for k, v in item.items():
                            if k == title_key:
                                continue
                            label = k.replace("_", " ").title()
                            parts.append(f"<div style='margin-bottom:8px'><strong style='font-size:11px;text-transform:uppercase;letter-spacing:.05em;color:#888'>{esc(label)}</strong>{render_value(v, depth+1)}</div>")
                        parts.append("</div>")
                    else:
                        parts.append(f"<div style='background:#f8f9fa;border-radius:8px;padding:14px;margin-bottom:10px'>")
                        for k, v in item.items():
                            label = k.replace("_", " ").title()
                            parts.append(f"<div style='margin-bottom:8px'><strong style='font-size:11px;text-transform:uppercase;letter-spacing:.05em;color:#888'>{esc(label)}</strong>{render_value(v, depth+1)}</div>")
                        parts.append("</div>")
                else:
                    parts.append(render_value(item, depth+1))
            return "".join(parts)
        if isinstance(value, dict):
            if not value:
                return "<span style='color:#999'>—</span>"
            parts = []
            for k, v in value.items():
                label = k.replace("_", " ").title()
                pad = 10 if depth > 0 else 0
                border = "3px solid #e0e0e0" if depth > 0 else "none"
                parts.append(
                    f"<div style='padding-left:{pad}px;border-left:{border};margin-bottom:12px;padding-bottom:4px'>"
                    f"<strong style='font-size:11px;font-weight:700;text-transform:uppercase;letter-spacing:.06em;color:#888;display:block;margin-bottom:4px'>{esc(label)}</strong>"
                    f"{render_value(v, depth+1)}</div>"
                )
            return "".join(parts)
        return f"<span>{esc(str(value))}</span>"

    def section_html(title, icon, data, color="#003366"):
        content = render_value(data)
        return (
            f'<div class="section"><div class="section-header" style="border-left:4px solid {color}">'
            f'<span class="section-icon">{icon}</span><h2>{title}</h2></div>'
            f'<div class="section-body">{content}</div></div>'
        )

    # ── Mappa sezioni → label, icona, colore ──
    SECTION_CONFIG = {
        "brand_identity":        ("Brand Identity & Posizionamento", "🏷️", "#003366"),
        "brand_values":          ("Valori del Brand", "💎", "#6366f1"),
        "product_portfolio":     ("Portafoglio Prodotti", "📦", "#10b981"),
        "service_vertical":      ("Analisi Verticale Servizi", "🔧", "#0ea5e9"),
        "product_vertical":      ("Analisi Verticale Prodotti", "📊", "#8b5cf6"),
        "customer_personas":     ("Customer Personas", "👥", "#f59e0b"),
        "reasons_to_buy":        ("Reasons to Buy", "🎯", "#ef4444"),
        "objections":            ("Gestione Obiezioni", "🛡️", "#6366f1"),
        "reviews_voc":           ("Voice of Customer", "🗣️", "#f59e0b"),
        "brand_voice":           ("Brand Voice & Guidelines", "✍️", "#ec4899"),
        "content_matrix":        ("Matrice Contenuti", "📐", "#ff9e1c"),
        "seasonal_roadmap":      ("Roadmap Stagionale", "📅", "#10b981"),
        "battlecards":           ("Competitor Battlecards", "⚔️", "#ef4444"),
        "psychographic_analysis":("Analisi Psicografica", "🧠", "#8b5cf6"),
        "visual_brief":          ("Visual Brief", "🎨", "#ec4899"),
    }

    # ── Genera le sezioni HTML ──
    sections_out = ""
    for key, (label, icon, color) in SECTION_CONFIG.items():
        data = analysis.get(key)
        if data and (isinstance(data, str) or (isinstance(data, (dict, list)) and len(data) > 0)):
            sections_out += section_html(label, icon, data, color)

    html = f"""<!DOCTYPE html>
<html lang="it">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>Analisi Strategica — {client_name}</title>
<style>
  @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800;900&display=swap');
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: 'Inter', sans-serif; background: #f8f9fa; color: #1a1a2e; font-size: 14px; line-height: 1.6; }}
  .cover {{ background: linear-gradient(135deg, #003366 0%, #001a33 100%); color: white; padding: 80px 60px; min-height: 280px; display: flex; flex-direction: column; justify-content: space-between; }}
  .cover-top {{ display: flex; justify-content: space-between; align-items: flex-start; }}
  .cover-brand {{ font-size: 11px; font-weight: 700; letter-spacing: .15em; text-transform: uppercase; color: rgba(255,255,255,.5); margin-bottom: 8px; }}
  .cover h1 {{ font-size: 42px; font-weight: 900; letter-spacing: -.02em; margin-bottom: 8px; }}
  .cover-subtitle {{ font-size: 16px; color: rgba(255,255,255,.65); }}
  .cover-meta {{ font-size: 12px; color: rgba(255,255,255,.4); margin-top: 40px; }}
  .orange-bar {{ background: #ff9e1c; height: 5px; }}
  .container {{ max-width: 960px; margin: 0 auto; padding: 40px 24px; }}
  .section {{ background: white; border-radius: 12px; margin-bottom: 28px; overflow: hidden; box-shadow: 0 1px 4px rgba(0,0,0,.06); }}
  .section-header {{ display: flex; align-items: center; gap: 12px; padding: 20px 24px; background: #fafafa; border-bottom: 1px solid #f0f0f0; }}
  .section-icon {{ font-size: 20px; }}
  .section-header h2 {{ font-size: 16px; font-weight: 700; color: #003366; letter-spacing: -.01em; }}
  .section-body {{ padding: 24px; }}
  .grid-2 {{ display: grid; grid-template-columns: 1fr 1fr; gap: 16px; margin-bottom: 16px; }}
  .card-inner {{ background: #f8f9fa; border-radius: 8px; padding: 16px; }}
  .card-inner h4 {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: .08em; color: #888; margin-bottom: 8px; }}
  .card-inner p {{ font-size: 13px; color: #333; }}
  .highlight-box {{ border-radius: 8px; padding: 16px 20px; margin-bottom: 16px; }}
  .highlight-box.lime {{ background: rgba(199,239,0,.08); border: 1px solid rgba(199,239,0,.3); }}
  .highlight-box.orange {{ background: rgba(255,158,28,.06); border: 1px solid rgba(255,158,28,.25); }}
  .highlight-box.blue {{ background: rgba(59,130,246,.06); border: 1px solid rgba(59,130,246,.2); }}
  .highlight-box h4 {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: .08em; color: #888; margin-bottom: 6px; }}
  .list-item {{ display: flex; gap: 8px; align-items: flex-start; margin-bottom: 6px; font-size: 13px; }}
  .dot {{ width: 7px; height: 7px; border-radius: 50%; flex-shrink: 0; margin-top: 5px; }}
  .tag-list {{ display: flex; gap: 8px; flex-wrap: wrap; margin: 12px 0; }}
  .tag {{ font-size: 12px; font-weight: 700; color: white; padding: 4px 12px; border-radius: 20px; }}
  .research-text {{ font-size: 12.5px; line-height: 1.8; color: #444; white-space: pre-wrap; }}
  .psych-levels {{ display: flex; flex-direction: column; gap: 16px; }}
  .psych-level {{ border-radius: 10px; padding: 20px; }}
  .level1 {{ background: rgba(59,130,246,.05); border: 1px solid rgba(59,130,246,.15); }}
  .level2 {{ background: rgba(139,92,246,.05); border: 1px solid rgba(139,92,246,.15); }}
  .level3 {{ background: rgba(236,72,153,.05); border: 1px solid rgba(236,72,153,.15); }}
  .level-badge {{ font-size: 10px; font-weight: 800; text-transform: uppercase; letter-spacing: .1em; color: #888; margin-bottom: 8px; }}
  .psych-level h4 {{ font-size: 13px; font-weight: 700; margin-bottom: 12px; }}
  .psych-level p {{ font-size: 13px; margin-bottom: 6px; }}
  .identity-stmt {{ font-style: italic; font-weight: 600; color: #6366f1; }}
  .battlecard {{ background: #f8f9fa; border-radius: 10px; padding: 20px; margin-bottom: 16px; }}
  .battlecard-header h4 {{ font-size: 15px; font-weight: 700; color: #003366; margin-bottom: 14px; }}
  .battlecard h5 {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: .08em; color: #888; margin-bottom: 8px; }}
  .advantage-box {{ background: rgba(199,239,0,.1); border: 1px solid rgba(199,239,0,.3); border-radius: 8px; padding: 14px; margin: 14px 0; }}
  .advantage-box h5 {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: .08em; color: #888; margin-bottom: 6px; }}
  .angle-card {{ display: flex; gap: 16px; align-items: flex-start; padding: 16px; background: #f8f9fa; border-radius: 10px; margin-bottom: 10px; }}
  .angle-num {{ background: #003366; color: white; width: 28px; height: 28px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-size: 12px; font-weight: 800; flex-shrink: 0; }}
  .angle-card h4 {{ font-size: 14px; font-weight: 700; margin-bottom: 4px; }}
  .angle-card p {{ font-size: 13px; color: #555; }}
  .funnel-badge {{ display: inline-block; font-size: 10px; font-weight: 700; text-transform: uppercase; letter-spacing: .07em; color: #ff9e1c; background: rgba(255,158,28,.1); padding: 2px 8px; border-radius: 4px; margin-bottom: 6px; }}
  .video-timeline {{ margin-top: 16px; }}
  .video-timeline h4 {{ font-size: 13px; font-weight: 700; margin-bottom: 12px; }}
  .timeline-row {{ display: flex; gap: 12px; align-items: flex-start; margin-bottom: 10px; }}
  .time-badge {{ background: #003366; color: white; font-size: 11px; font-weight: 700; padding: 3px 10px; border-radius: 4px; flex-shrink: 0; white-space: nowrap; }}
  .months-grid {{ display: grid; grid-template-columns: repeat(4, 1fr); gap: 12px; margin-bottom: 16px; }}
  .month-card {{ background: #f8f9fa; border-radius: 8px; padding: 14px; font-size: 12px; }}
  .month-name {{ font-weight: 800; font-size: 13px; color: #003366; margin-bottom: 6px; }}
  .month-angle {{ font-weight: 600; margin-bottom: 4px; color: #333; }}
  .month-offer {{ color: #555; margin-bottom: 3px; }}
  .month-urgency {{ color: #555; margin-bottom: 4px; }}
  .month-priority {{ font-size: 11px; font-weight: 700; text-transform: uppercase; letter-spacing: .06em; }}
  .footer {{ text-align: center; padding: 40px; color: #bbb; font-size: 12px; }}
  @media print {{
    body {{ background: white; }}
    .section {{ box-shadow: none; border: 1px solid #eee; page-break-inside: avoid; }}
    .cover {{ print-color-adjust: exact; -webkit-print-color-adjust: exact; }}
    .months-grid {{ grid-template-columns: repeat(3,1fr); }}
  }}
</style>
</head>
<body>
<div class="cover">
  <div class="cover-top">
    <div>
      <div class="cover-brand">Antigravity Operative Manager</div>
      <h1>{client_name}</h1>
      <div class="cover-subtitle">Analisi Strategica Avanzata — Meta Ads & Brand Intelligence</div>
    </div>
  </div>
  <div class="cover-meta">Generato il {today} · Riservato e confidenziale</div>
</div>
<div class="orange-bar"></div>
<div class="container">
{sections_out}
</div>
<div class="footer">Analisi generata da Antigravity Operative Manager · {today}</div>
</body>
</html>"""

    return HTMLResponse(content=html, media_type="text/html")


# ══════════════════════════════════════════════════════════
#  CREATIVE INTELLIGENCE — fetch ad creatives + AI analysis
# ══════════════════════════════════════════════════════════

@app.get("/live-ads/creative-intelligence/{client_id}")
async def get_creative_intelligence(client_id: str):
    intel_path = CLIENTS_DIR / client_id / "creative_intelligence.json"
    if not intel_path.exists():
        return {"analysis": None, "period": None, "generated_at": None, "ads_count": 0}
    with open(intel_path, "r") as f:
        return json.load(f)


@app.get("/live-ads/creatives/{client_id}")
async def get_ad_creatives(client_id: str, date_preset: str = "last_90d", since: str = "", until: str = ""):
    metadata = storage_service.get_metadata(client_id)
    ad_id = metadata.get("ad_account_id", "").strip()
    if not ad_id:
        raise HTTPException(status_code=400, detail="Ad Account ID non configurato per questo cliente")
    token = metadata.get("meta_access_token") or os.getenv("META_ACCESS_TOKEN")
    if not token:
        raise HTTPException(status_code=500, detail="META_ACCESS_TOKEN non trovato nel file .env")
    if not ad_id.startswith("act_"):
        ad_id = f"act_{ad_id}"

    insights_params: dict = {
        "fields": "ad_id,ad_name,adset_name,campaign_name,spend,impressions,reach,clicks,ctr,cpc,cpm,actions,cost_per_action_type",
        "level": "ad",
        "access_token": token,
        "limit": 100,
    }
    if since and until:
        insights_params["time_range"] = json.dumps({"since": since, "until": until})
    else:
        insights_params["date_preset"] = date_preset

    async with httpx.AsyncClient(timeout=60.0) as hc:
        insights_resp = await hc.get(
            f"https://graph.facebook.com/v19.0/{ad_id}/insights",
            params=insights_params
        )
    insights_data = insights_resp.json()
    if "error" in insights_data:
        raise HTTPException(status_code=400, detail=insights_data["error"].get("message", "Meta API error"))

    ads_insights = insights_data.get("data", [])
    if not ads_insights:
        return {"ads": [], "period": date_preset, "client_id": client_id, "total_ads": 0}

    ads_insights.sort(key=lambda x: float(x.get("spend", 0) or 0), reverse=True)
    top_ads = ads_insights[:40]

    async def fetch_creative(ad_insight: dict) -> dict:
        ad_insight_id = ad_insight.get("ad_id")
        base = {
            "ad_id": ad_insight_id,
            "ad_name": ad_insight.get("ad_name", "—"),
            "adset_name": ad_insight.get("adset_name", "—"),
            "campaign_name": ad_insight.get("campaign_name", "—"),
            "spend": round(float(ad_insight.get("spend", 0) or 0), 2),
            "impressions": int(ad_insight.get("impressions", 0) or 0),
            "reach": int(ad_insight.get("reach", 0) or 0),
            "clicks": int(ad_insight.get("clicks", 0) or 0),
            "ctr": round(float(ad_insight.get("ctr", 0) or 0), 3),
            "cpc": round(float(ad_insight.get("cpc", 0) or 0), 2),
            "cpm": round(float(ad_insight.get("cpm", 0) or 0), 2),
            "conversioni": _meta_extract_conversions(ad_insight.get("actions", [])),
            "cpa": _meta_extract_cpa(ad_insight.get("cost_per_action_type", [])),
        }
        try:
            async with httpx.AsyncClient(timeout=15.0) as hc:
                creative_resp = await hc.get(
                    f"https://graph.facebook.com/v19.0/{ad_insight_id}",
                    params={
                        "fields": "creative{id,body,title,description,image_url,thumbnail_url,object_story_spec}",
                        "access_token": token,
                    }
                )
            creative_data = creative_resp.json()
            creative = creative_data.get("creative", {})
            story_spec = creative.get("object_story_spec", {})
            link_data = story_spec.get("link_data", {})
            video_data = story_spec.get("video_data", {})
            return {
                **base,
                "creative_id": creative.get("id", ""),
                "body": (creative.get("body") or link_data.get("message") or video_data.get("message", "")).strip(),
                "title": (creative.get("title") or link_data.get("name") or video_data.get("title", "")).strip(),
                "description": (creative.get("description") or link_data.get("description", "")).strip(),
                "image_url": creative.get("image_url") or link_data.get("picture", ""),
                "thumbnail_url": creative.get("thumbnail_url") or video_data.get("image_url", ""),
            }
        except Exception:
            return base

    ads_with_creatives = await asyncio.gather(*[fetch_creative(ad) for ad in top_ads])
    return {
        "ads": list(ads_with_creatives),
        "period": date_preset,
        "client_id": client_id,
        "total_ads": len(ads_insights),
    }


class CreativeAnalysisRequest(BaseModel):
    ads: List[dict] = []
    date_preset: str = "last_90d"


@app.post("/live-ads/analyze-creatives/{client_id}")
async def analyze_ad_creatives(client_id: str, request: CreativeAnalysisRequest):
    metadata = storage_service.get_metadata(client_id)
    if not request.ads:
        raise HTTPException(status_code=400, detail="Nessuna ad da analizzare")

    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    research_context = ""
    if research_path.exists():
        with open(research_path, "r") as f:
            research_context = f.read()[:2000]

    ads_sorted = sorted(request.ads, key=lambda x: float(x.get("ctr", 0) or 0), reverse=True)

    ads_text_parts = []
    for i, ad in enumerate(ads_sorted[:35], 1):
        body = (ad.get("body") or "").strip()
        title = (ad.get("title") or "").strip()
        desc = (ad.get("description") or "").strip()
        cpa_str = f"€{ad.get('cpa')}" if ad.get("cpa") else "N/D"
        lines = [
            f"AD #{i} | Campagna: {ad.get('campaign_name', '?')} | AdSet: {ad.get('adset_name', '?')}",
            f"Nome: {ad.get('ad_name', '?')}",
            f"Spend: €{ad.get('spend', 0)} | CTR: {ad.get('ctr', 0)}% | CPC: €{ad.get('cpc', 0)} | Conv: {ad.get('conversioni', 0)} | CPA: {cpa_str}",
        ]
        if title:
            lines.append(f"Headline: {title}")
        if body:
            lines.append(f"Copy: {body}")
        if desc:
            lines.append(f"Descrizione: {desc}")
        ads_text_parts.append("\n".join(lines))

    ads_text = "\n\n---\n\n".join(ads_text_parts)
    research_block = f"CONTESTO DI MERCATO:\n{research_context}\n\n" if research_context else ""

    prompt_text = f"""Sei un Senior Creative Strategist specializzato in Meta Ads, con expertise in copywriting, psicologia persuasiva e analisi delle performance creative.

CLIENTE: {metadata.get('name', '')} | Settore: {metadata.get('industry', '')}
PERIODO ANALIZZATO: {request.date_preset}

{research_block}━━━ INSERZIONI ANALIZZATE ({len(request.ads)} ads totali, ordinate per CTR) ━━━

{ads_text}

━━━ ANALISI STRATEGICA RICHIESTA ━━━
Produci un'analisi approfondita con questa struttura markdown:

## 🏆 Creatività Vincenti
Le ads con CTR/conversioni/CPA migliori. Per ognuna: angolo comunicativo usato, tipo di hook, tono, promessa principale, perché funziona psicologicamente.

## 📊 Pattern delle Ads che Funzionano
Caratteristiche comuni tra le ads top performer: stile copy (lungo/corto, formale/informale), tipo di apertura, struttura del messaggio, call-to-action usate.

## 🎯 Angoli Comunicativi Identificati
Lista degli angoli unici trovati nelle ads reali. Per ognuno: nome angolo, come viene usato, performance media (CTR/CPA). Distingui tra angoli provati e vincenti vs angoli provati ma deboli.

## ❌ Cosa Non Funziona
Pattern delle ads con CTR basso o CPA alto: perché non risuonano, quali errori comunicativi emergono.

## 💡 Raccomandazioni per Nuove Creatività
Basandoti SOLO sui dati reali:
- I 3 hook/aperture più efficaci da replicare e variare
- 3 angoli nuovi non ancora testati che potrebbero funzionare per questo cliente
- Formati da testare (video vs statico, lunghezza copy, struttura)

## 🔑 Vocabolario che Converte
Parole, frasi, strutture linguistiche trovate nelle ads vincenti che risuonano col target. Questo è il linguaggio da replicare.

## 📌 Intelligence per il Team
3-5 insight pratici e immediati che il team creativo deve conoscere per produrre nuove creatività efficaci per questo cliente.

Sii specifico. Cita le ads per nome/numero. Usa i dati reali forniti."""

    content: list = [{"type": "text", "text": prompt_text}]

    # Add thumbnails of top performing ads that have images
    ads_with_images = [ad for ad in ads_sorted[:10] if ad.get("thumbnail_url") or ad.get("image_url")][:5]
    for ad in ads_with_images:
        img_url = ad.get("thumbnail_url") or ad.get("image_url")
        if img_url:
            content.append({"type": "text", "text": f"\n[Creatività visiva — '{ad.get('ad_name', '?')}' CTR:{ad.get('ctr', 0)}%]"})
            content.append({"type": "image_url", "image_url": {"url": img_url}})

    messages = [{"role": "user", "content": content}]
    analysis = await ai_service._call_ai("anthropic/claude-3.7-sonnet", messages)

    intel_data = {
        "analysis": analysis,
        "period": request.date_preset,
        "ads_count": len(request.ads),
        "generated_at": datetime.now().isoformat(),
        "client_id": client_id,
    }
    intel_path = CLIENTS_DIR / client_id / "creative_intelligence.json"
    with open(intel_path, "w") as f:
        json.dump(intel_data, f, ensure_ascii=False, indent=2)
    storage_service.sync_creative_intelligence(client_id, intel_data)

    return intel_data


# ══════════════════════════════════════════════════════════
#  SHOPIFY — OAuth connection + data endpoints
# ══════════════════════════════════════════════════════════

SHOPIFY_CLIENT_ID = os.getenv("SHOPIFY_CLIENT_ID", "")
SHOPIFY_CLIENT_SECRET = os.getenv("SHOPIFY_CLIENT_SECRET", "")
SHOPIFY_SCOPES = "read_orders,read_all_orders,read_customers,read_analytics,read_checkouts,read_customer_events,read_reports"
SHOPIFY_REDIRECT_URI = os.getenv("SHOPIFY_REDIRECT_URI", "http://127.0.0.1:8001/shopify/callback")


@app.get("/shopify/install")
async def shopify_install(client_id: str, shop: str):
    """Avvia il flusso OAuth Shopify per un cliente specifico."""
    if not SHOPIFY_CLIENT_ID:
        raise HTTPException(status_code=500, detail="SHOPIFY_CLIENT_ID non configurato nel .env")
    # Normalize shop domain
    shop = shop.replace("https://", "").replace("http://", "").rstrip("/")
    if not shop.endswith(".myshopify.com"):
        shop = f"{shop}.myshopify.com"
    # Save shop domain to client metadata immediately
    metadata = storage_service.get_metadata(client_id)
    metadata["shopify_domain"] = shop
    storage_service.save_metadata(client_id, metadata)
    # Build OAuth URL
    from urllib.parse import urlencode
    params = urlencode({
        "client_id": SHOPIFY_CLIENT_ID,
        "scope": SHOPIFY_SCOPES,
        "redirect_uri": SHOPIFY_REDIRECT_URI,
        "state": client_id,
    })
    auth_url = f"https://{shop}/admin/oauth/authorize?{params}"
    return RedirectResponse(url=auth_url)


@app.get("/shopify/callback")
async def shopify_callback(code: str, shop: str, state: str, hmac: str = ""):
    """Riceve il callback OAuth da Shopify e salva l'access token."""
    if not SHOPIFY_CLIENT_SECRET:
        raise HTTPException(status_code=500, detail="SHOPIFY_CLIENT_SECRET non configurato nel .env")
    client_id = state
    # Exchange code for access token
    async with httpx.AsyncClient(timeout=30.0) as hc:
        resp = await hc.post(
            f"https://{shop}/admin/oauth/access_token",
            json={
                "client_id": SHOPIFY_CLIENT_ID,
                "client_secret": SHOPIFY_CLIENT_SECRET,
                "code": code,
            }
        )
    data = resp.json()
    access_token = data.get("access_token")
    if not access_token:
        raise HTTPException(status_code=400, detail=f"Token non ricevuto: {data}")
    # Save token to client metadata
    metadata = storage_service.get_metadata(client_id)
    metadata["shopify_domain"] = shop
    metadata["shopify_token"] = access_token
    storage_service.save_metadata(client_id, metadata)
    # Redirect back to client page
    from fastapi.responses import RedirectResponse
    frontend_url = os.getenv("FRONTEND_URL", "http://localhost:3001")
    return RedirectResponse(url=f"{frontend_url}/client/{client_id}?shopify=connected")


@app.get("/clients/{client_id}/shopify/orders")
async def get_shopify_orders(client_id: str, limit: int = 50, days: int = 90):
    """Recupera gli ordini recenti da Shopify per il cliente."""
    metadata = storage_service.get_metadata(client_id)
    shop = metadata.get("shopify_domain", "").strip()
    token = metadata.get("shopify_token", "").strip()
    if not shop or not token:
        raise HTTPException(status_code=400, detail="Shopify non collegato. Usa il bottone 'Connetti Shopify' nella sezione Sorgenti.")
    from datetime import timedelta
    since = (datetime.now() - timedelta(days=days)).strftime("%Y-%m-%dT%H:%M:%S")
    async with httpx.AsyncClient(timeout=30.0) as hc:
        resp = await hc.get(
            f"https://{shop}/admin/api/2024-01/orders.json",
            headers={"X-Shopify-Access-Token": token},
            params={
                "status": "any",
                "limit": limit,
                "created_at_min": since,
                "fields": "id,created_at,total_price,subtotal_price,currency,financial_status,customer,line_items,source_name,landing_site,referring_site,utm_parameters,note_attributes",
            }
        )
    data = resp.json()
    if "errors" in data:
        raise HTTPException(status_code=400, detail=str(data["errors"]))
    orders = data.get("orders", [])
    # Aggregate summary
    total_revenue = sum(float(o.get("total_price", 0)) for o in orders)
    new_customers = sum(1 for o in orders if o.get("customer", {}).get("orders_count", 1) == 1)
    returning = len(orders) - new_customers
    return {
        "orders": orders,
        "summary": {
            "total_orders": len(orders),
            "total_revenue": round(total_revenue, 2),
            "new_customers": new_customers,
            "returning_customers": returning,
            "avg_order_value": round(total_revenue / len(orders), 2) if orders else 0,
            "period_days": days,
        },
        "shop": shop,
        "client_id": client_id,
    }


@app.get("/clients/{client_id}/shopify/customers")
async def get_shopify_customers(client_id: str, limit: int = 50):
    """Recupera i clienti Shopify con dati per CAPI."""
    metadata = storage_service.get_metadata(client_id)
    shop = metadata.get("shopify_domain", "").strip()
    token = metadata.get("shopify_token", "").strip()
    if not shop or not token:
        raise HTTPException(status_code=400, detail="Shopify non collegato.")
    async with httpx.AsyncClient(timeout=30.0) as hc:
        resp = await hc.get(
            f"https://{shop}/admin/api/2024-01/customers.json",
            headers={"X-Shopify-Access-Token": token},
            params={
                "limit": limit,
                "fields": "id,email,phone,first_name,last_name,city,province,country,zip,orders_count,total_spent,created_at",
            }
        )
    data = resp.json()
    if "errors" in data:
        raise HTTPException(status_code=400, detail=str(data["errors"]))
    return {"customers": data.get("customers", []), "shop": shop}


@app.delete("/clients/{client_id}/shopify")
async def disconnect_shopify(client_id: str):
    """Rimuove la connessione Shopify dal cliente."""
    metadata = storage_service.get_metadata(client_id)
    metadata.pop("shopify_token", None)
    metadata.pop("shopify_domain", None)
    storage_service.save_metadata(client_id, metadata)
    return {"connected": False}


class ShopifyManualTokenRequest(BaseModel):
    shop: str
    access_token: str

@app.put("/clients/{client_id}/shopify/token")
async def save_shopify_token_manual(client_id: str, body: ShopifyManualTokenRequest):
    """Salva manualmente un access token Shopify (da custom app o collaborator access)."""
    shop = body.shop.replace("https://", "").replace("http://", "").rstrip("/")
    if not shop.endswith(".myshopify.com"):
        shop = shop + ".myshopify.com"
    token = body.access_token.strip()
    if not token:
        raise HTTPException(status_code=400, detail="access_token vuoto")
    # Verify the token works before saving
    try:
        async with httpx.AsyncClient(timeout=10.0) as hc:
            resp = await hc.get(
                f"https://{shop}/admin/api/2024-01/shop.json",
                headers={"X-Shopify-Access-Token": token},
            )
        if resp.status_code != 200:
            raise HTTPException(status_code=400, detail=f"Token non valido (HTTP {resp.status_code})")
        shop_data = resp.json().get("shop", {})
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Impossibile connettersi allo store: {e}")
    metadata = storage_service.get_metadata(client_id)
    metadata["shopify_domain"] = shop
    metadata["shopify_token"] = token
    storage_service.save_metadata(client_id, metadata)
    return {
        "connected": True,
        "shop": shop,
        "shop_name": shop_data.get("name", shop),
        "currency": shop_data.get("currency", ""),
    }


@app.get("/clients/{client_id}/shopify/status")
async def get_shopify_status(client_id: str):
    """Controlla se Shopify è collegato e restituisce lo stato."""
    metadata = storage_service.get_metadata(client_id)
    shop = metadata.get("shopify_domain", "").strip()
    token = metadata.get("shopify_token", "").strip()
    if not shop or not token:
        return {"connected": False, "shop": None}
    # Quick check
    try:
        async with httpx.AsyncClient(timeout=10.0) as hc:
            resp = await hc.get(
                f"https://{shop}/admin/api/2024-01/shop.json",
                headers={"X-Shopify-Access-Token": token},
            )
        shop_data = resp.json().get("shop", {})
        return {
            "connected": True,
            "shop": shop,
            "shop_name": shop_data.get("name", shop),
            "currency": shop_data.get("currency", ""),
            "plan": shop_data.get("plan_name", ""),
        }
    except Exception:
        return {"connected": False, "shop": shop, "error": "Connessione fallita"}


# ══════════════════════════════════════════════════════════════════════
#  GOOGLE TASKS INTEGRATION (attività su Google Calendar)
# ══════════════════════════════════════════════════════════════════════

GOOGLE_CAL_CLIENT_ID = os.getenv("GOOGLE_CALENDAR_CLIENT_ID", "")
GOOGLE_CAL_CLIENT_SECRET = os.getenv("GOOGLE_CALENDAR_CLIENT_SECRET", "")
GOOGLE_CAL_REDIRECT_URI = os.getenv("GOOGLE_CALENDAR_REDIRECT_URI", "https://antigravity-backend-production-41ee.up.railway.app/google-calendar/callback")
_GCAL_TOKEN_FILE = Path(__file__).parent.parent / "google_calendar_token.json"


def _load_google_token() -> dict | None:
    """Carica il token Google dal file JSON. Ritorna il dict raw."""
    if not _GCAL_TOKEN_FILE.exists():
        return None
    try:
        with open(_GCAL_TOKEN_FILE, "r") as f:
            return json.load(f)
    except Exception:
        return None


def _save_google_token(token_data: dict):
    """Salva il token Google come JSON."""
    with open(_GCAL_TOKEN_FILE, "w") as f:
        json.dump(token_data, f, indent=2)


async def _get_valid_access_token() -> str | None:
    """Ritorna un access_token valido, facendo refresh se scaduto."""
    token = _load_google_token()
    if not token or "access_token" not in token:
        return None

    # Se abbiamo un refresh_token, facciamo sempre refresh per sicurezza
    refresh_token = token.get("refresh_token")
    if not refresh_token:
        return token["access_token"]

    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post("https://oauth2.googleapis.com/token", data={
                "client_id": GOOGLE_CAL_CLIENT_ID,
                "client_secret": GOOGLE_CAL_CLIENT_SECRET,
                "refresh_token": refresh_token,
                "grant_type": "refresh_token",
            })
            data = resp.json()
        if "access_token" in data:
            token["access_token"] = data["access_token"]
            _save_google_token(token)
            return data["access_token"]
    except Exception as e:
        print(f"⚠️ Refresh token fallito: {e}")

    return token["access_token"]


@app.get("/google-calendar/status")
async def google_calendar_status():
    """Controlla se Google è connesso (ha un token salvato)."""
    token = _load_google_token()
    if token and token.get("access_token"):
        return {"connected": True}
    return {"connected": False}


@app.get("/google-calendar/install")
async def google_calendar_install():
    """Avvia OAuth Google per permesso Tasks."""
    if not GOOGLE_CAL_CLIENT_ID or not GOOGLE_CAL_CLIENT_SECRET:
        raise HTTPException(status_code=500, detail="Google non configurato sul server")

    from urllib.parse import urlencode
    params = {
        "client_id": GOOGLE_CAL_CLIENT_ID,
        "redirect_uri": GOOGLE_CAL_REDIRECT_URI,
        "response_type": "code",
        "scope": "https://www.googleapis.com/auth/tasks",
        "access_type": "offline",
        "prompt": "consent",
    }
    return RedirectResponse(f"https://accounts.google.com/o/oauth2/auth?{urlencode(params)}")


@app.get("/google-calendar/callback")
async def google_calendar_callback(code: str = "", error: str = ""):
    """Callback OAuth — scambia il code per access+refresh token."""
    if error:
        raise HTTPException(status_code=400, detail=f"Errore Google: {error}")
    if not code:
        raise HTTPException(status_code=400, detail="Codice mancante")

    try:
        async with httpx.AsyncClient() as client:
            resp = await client.post("https://oauth2.googleapis.com/token", data={
                "code": code,
                "client_id": GOOGLE_CAL_CLIENT_ID,
                "client_secret": GOOGLE_CAL_CLIENT_SECRET,
                "redirect_uri": GOOGLE_CAL_REDIRECT_URI,
                "grant_type": "authorization_code",
            })
            token_data = resp.json()

        if "error" in token_data:
            print(f"❌ Token exchange error: {token_data}")
            raise HTTPException(status_code=400, detail=f"Google error: {token_data.get('error_description', token_data['error'])}")

        _save_google_token(token_data)
        print(f"✅ Google Tasks connesso! Token salvato.")

        frontend_url = os.getenv("FRONTEND_URL", "https://operative.alessioferlizzo.com")
        return RedirectResponse(f"{frontend_url}?gcal=connected")
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ Callback error: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.delete("/google-calendar")
async def google_calendar_disconnect():
    """Disconnetti Google."""
    if _GCAL_TOKEN_FILE.exists():
        _GCAL_TOKEN_FILE.unlink()
    return {"connected": False}


async def _call_tasks_api(method: str, path: str, body: dict | None = None) -> dict:
    """Chiama Google Tasks API direttamente via HTTP."""
    access_token = await _get_valid_access_token()
    if not access_token:
        raise HTTPException(status_code=401, detail="Google non connesso")

    url = f"https://tasks.googleapis.com/tasks/v1/{path}"
    headers = {"Authorization": f"Bearer {access_token}", "Content-Type": "application/json"}

    async with httpx.AsyncClient() as client:
        if method == "POST":
            resp = await client.post(url, headers=headers, json=body)
        elif method == "PUT":
            resp = await client.put(url, headers=headers, json=body)
        elif method == "DELETE":
            resp = await client.delete(url, headers=headers)
            return {}
        else:
            resp = await client.get(url, headers=headers)

    if resp.status_code == 401:
        raise HTTPException(status_code=401, detail="Token scaduto")
    if resp.status_code >= 400:
        raise HTTPException(status_code=resp.status_code, detail=resp.text)

    return resp.json() if resp.text else {}


@app.post("/tasks/{task_id}/calendar")
async def sync_task_to_calendar(task_id: str):
    """Crea un'attività su Google Tasks."""
    tasks = storage_service.get_tasks()
    task = next((t for t in tasks if t["id"] == task_id), None)
    if not task:
        raise HTTPException(status_code=404, detail="Task non trovata")

    title = task.get("title", "Task")
    if task.get("client_name"):
        title = f"[{task['client_name']}] {title}"

    notes_parts = []
    if task.get("notes"):
        notes_parts.append(task["notes"])
    if task.get("priority"):
        notes_parts.append(f"Priorità: {task['priority']}")
    if task.get("estimated_time"):
        notes_parts.append(f"Tempo stimato: {task['estimated_time']}")
    if task.get("due_time"):
        notes_parts.append(f"Orario: {task['due_time']}")

    gtask: dict = {"title": title, "notes": "\n".join(notes_parts)}

    # Data scadenza con orario
    due_date = task.get("due_date", "")
    due_time = task.get("due_time", "")
    if due_date and due_time:
        gtask["due"] = f"{due_date}T{due_time}:00+02:00"
    elif due_date:
        gtask["due"] = f"{due_date}T00:00:00Z"

    if task.get("status") == "done":
        gtask["status"] = "completed"
    else:
        gtask["status"] = "needsAction"

    existing_id = task.get("gcal_event_id")
    if existing_id:
        result = await _call_tasks_api("PUT", f"lists/@default/tasks/{existing_id}", gtask)
    else:
        result = await _call_tasks_api("POST", "lists/@default/tasks", gtask)

    task["gcal_event_id"] = result.get("id", existing_id)
    storage_service.save_tasks(tasks)
    return {"synced": True, "event_id": task["gcal_event_id"]}


@app.delete("/tasks/{task_id}/calendar")
async def unsync_task_from_calendar(task_id: str):
    """Rimuove un'attività da Google Tasks."""
    tasks = storage_service.get_tasks()
    task = next((t for t in tasks if t["id"] == task_id), None)
    if not task:
        raise HTTPException(status_code=404, detail="Task non trovata")

    gtask_id = task.get("gcal_event_id")
    if not gtask_id:
        return {"synced": False}

    try:
        await _call_tasks_api("DELETE", f"lists/@default/tasks/{gtask_id}")
    except Exception as e:
        print(f"⚠️ Errore rimozione: {e}")

    task.pop("gcal_event_id", None)
    storage_service.save_tasks(tasks)
    return {"synced": False}


# ── VAULT ENDPOINTS ──
class VaultSaveRequest(BaseModel):
    title: str
    text: str = ""
    funnel_stage: Optional[str] = ""
    format: Optional[str] = ""
    img_link: Optional[str] = ""

@app.post("/clients/{client_id}/vault/{vault_type}")
async def save_to_vault_endpoint(client_id: str, vault_type: str, request: VaultSaveRequest):
    if vault_type not in ["copy", "angle", "graphic"]:
        raise HTTPException(status_code=400, detail="Invalid vault type. Use 'copy', 'angle' or 'graphic'.")
        
    try:
        metadata = storage_service.get_metadata(client_id)
        from .notion_service import notion_service
        
        # Prepare basic item data
        item_data = {
            "title": request.title,
            "text": request.text,
            "client_name": metadata.get("name", ""),
            "sector": metadata.get("industry", ""),
            "funnel_stage": request.funnel_stage,
            "format": request.format,
            "img_link": request.img_link
        }
        
        if vault_type == "graphic" and request.img_link:
            # L'URL locale viene passato a notion_service.py che si occuperà
            # di caricarlo su host pubblico permanente (es. Catbox.moe)
            pass
        
        success = await notion_service.save_to_vault(vault_type, item_data)
        if success:
            return {"status": "success", "message": f"Saved to {vault_type} vault"}
        else:
            raise HTTPException(status_code=500, detail="Failed to save to Notion Vault")
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")

class ScriptRequest(BaseModel):
    title: str
    description: str = ""
    emotion: str = ""
    script_instructions: str = ""
    count: int = 1

@app.post("/clients/{client_id}/scripts")
async def generate_script_endpoint(client_id: str, request: ScriptRequest):
    metadata = storage_service.get_metadata(client_id)

    # 🔥 NUOVO: Carica TUTTA l'Analisi Strategica (14 sezioni)
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id,
        metadata=metadata,
        supabase_client=supabase,
        focus_areas=["customer_personas", "brand_voice", "visual_brief", "product_vertical", "psychographic_analysis"]
    )
    
    # Knowledge (Borzacchiello, Vignali, Dosio, Schwartz) is now injected
    # directly in ai_service.generate_script via knowledge_loader
    rules = ""

    angle = {"title": request.title, "description": request.description}
    count = min(max(request.count, 1), 5)

    scripts = []
    for i in range(count):
        script = await ai_service.generate_script(
            angle, strategic_context, rules, metadata.get("preferences", {}),
            script_instructions=request.script_instructions,
            variation_index=i,
            total_variations=count
        )
        
        scripts_dir = CLIENTS_DIR / client_id / "scripts"
        scripts_dir.mkdir(parents=True, exist_ok=True)
        script_id = f"script_{len(os.listdir(scripts_dir)) + 1}"
        script_path = scripts_dir / f"{script_id}.md"
        with open(script_path, "w") as f:
            f.write(script)

        scripts.append({"script_id": script_id, "content": script})

    storage_service.sync_scripts(client_id)
    return scripts

@app.post("/clients/{client_id}/feedback")
async def submit_feedback(client_id: str, request: FeedbackRequest):
    metadata = storage_service.get_metadata(client_id)
    
    # Record feedback to improve future scripts
    if "preferences" not in metadata:
        metadata["preferences"] = {"feedback_history": [], "avoid_words": []}
    
    metadata["preferences"].setdefault("feedback_history", []).append(request.feedback)
    
    # Extract keywords/rules from feedback (simple logic for now, could be AI-driven)
    if "evita" in request.feedback.lower():
        words = request.feedback.lower().split("evita")[1].replace("la parola", "").strip().split(",")
        metadata["preferences"].setdefault("avoid_words", []).extend([w.strip() for w in words])
    
    storage_service.save_metadata(client_id, metadata)
    
    # Regenerate script with new preferences
    research_path = CLIENTS_DIR / client_id / "research" / "market_research.md"
    research = ""
    if research_path.exists():
        with open(research_path, "r") as f:
            research = f.read()
    
    # Knowledge now injected via knowledge_loader in ai_service
    rules = ""
        
    # Load original script content for contextual refinement
    original_script_content = ""
    scripts_dir = CLIENTS_DIR / client_id / "scripts"
    if request.script_id:
        script_path = scripts_dir / f"{request.script_id}.md"
        if script_path.exists():
            with open(script_path, "r") as f:
                original_script_content = f.read()

    # Load strategic context (replaces the stale research file)
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id,
        metadata=metadata,
        supabase_client=supabase,
        focus_areas=["customer_personas", "brand_voice", "visual_brief", "product_vertical", "service_vertical", "reviews_voc"]
    )

    angle = {"title": request.angle_title or "Script", "description": "Regenerated based on feedback"}
    new_script = await ai_service.generate_script(
        angle, strategic_context, rules, metadata["preferences"], 
        script_instructions=request.feedback,
        original_script=original_script_content
    )
    
    # Save as new version
    base_id = request.script_id or "script_1"
    scripts_dir = CLIENTS_DIR / client_id / "scripts"
    scripts_dir.mkdir(parents=True, exist_ok=True)
    
    existing_versions = [f for f in os.listdir(scripts_dir) if f.startswith(base_id.split('_v')[0])]
    version = len(existing_versions) + 1
    script_id = f"{base_id.split('_v')[0]}_v{version}"
    
    script_path = scripts_dir / f"{script_id}.md"
    with open(script_path, "w") as f:
        f.write(new_script)

    storage_service.sync_scripts(client_id)
    return {"script_id": script_id, "content": new_script}


# ═══════════════════════════════════════════════
#  LISTS (global)
# ═══════════════════════════════════════════════

class ListCreate(BaseModel):
    title: str
    color: Optional[str] = "#007aff"

@app.get("/lists")
async def list_custom_lists():
    return storage_service.get_lists()

@app.post("/lists")
async def create_custom_list(lst: ListCreate):
    lists = storage_service.get_lists()
    new_lst = {"id": str(uuid.uuid4()), "title": lst.title, "color": lst.color}
    lists.append(new_lst)
    storage_service.save_lists(lists)
    return new_lst

@app.delete("/lists/{list_id}")
async def delete_custom_list(list_id: str):
    lists = storage_service.get_lists()
    new_lists = [l for l in lists if l["id"] != list_id]
    storage_service.save_lists(new_lists)
    return {"message": "List deleted"}


# ═══════════════════════════════════════════════
#  TASKS (global)
# ═══════════════════════════════════════════════

class TaskCreate(BaseModel):
    title: str
    client_id: Optional[str] = ""
    client_name: Optional[str] = ""
    priority: Optional[str] = ""
    due_date: Optional[str] = ""
    due_time: Optional[str] = ""           # HH:MM time string
    notes: Optional[str] = ""
    estimated_time: Optional[str] = ""
    parent_id: Optional[str] = None
    task_type: Optional[str] = ""          # ads, report, contenuto, chiamata, admin
    subtasks: Optional[List[Dict]] = []    # [{id, text, done}]
    recurring: Optional[bool] = False
    recurring_frequency: Optional[str] = ""  # daily, weekly, monthly
    reminder_at: Optional[str] = ""           # ISO datetime string or short code
    list_id: Optional[str] = ""               # Custom list ID
    flagged: Optional[bool] = False            # Bandierina/contrassegna

class TaskUpdate(BaseModel):
    title: Optional[str] = None
    status: Optional[str] = None
    priority: Optional[str] = None
    due_date: Optional[str] = None
    due_time: Optional[str] = None          # HH:MM time string
    notes: Optional[str] = None
    client_id: Optional[str] = None
    client_name: Optional[str] = None
    estimated_time: Optional[str] = None
    parent_id: Optional[str] = None
    task_type: Optional[str] = None
    subtasks: Optional[List[Dict]] = None
    recurring: Optional[bool] = None
    recurring_frequency: Optional[str] = None
    reminder_at: Optional[str] = None
    list_id: Optional[str] = None
    flagged: Optional[bool] = None          # Bandierina/contrassegna
    completed_at: Optional[str] = None

@app.get("/tasks")
async def list_tasks():
    return storage_service.get_tasks()

@app.post("/tasks")
async def create_task(task: TaskCreate):
    return storage_service.create_task(
        title=task.title,
        client_id=task.client_id,
        client_name=task.client_name,
        priority=task.priority or "",
        due_date=task.due_date,
        due_time=task.due_time,
        notes=task.notes,
        estimated_time=task.estimated_time,
        parent_id=task.parent_id,
        task_type=task.task_type,
        subtasks=task.subtasks,
        recurring=task.recurring,
        recurring_frequency=task.recurring_frequency,
        reminder_at=task.reminder_at,
        list_id=task.list_id,
        flagged=task.flagged or False,
    )

@app.patch("/tasks/{task_id}")
async def update_task(task_id: str, update: TaskUpdate):
    try:
    # Only update fields that were explicitly set (exclude_unset), but also allow empty strings
        updates = {k: v for k, v in update.dict(exclude_unset=True).items()}
        # Allow None only for subtasks, recurring, completed_at, flagged (boolean can be False)
        updates = {k: v for k, v in updates.items() if v is not None or k in ("subtasks", "recurring", "completed_at", "flagged", "due_time", "reminder_at", "recurring_frequency", "list_id", "priority")}
        return storage_service.update_task(task_id, updates)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Task not found")

@app.delete("/tasks/{task_id}")
async def delete_task(task_id: str):
    try:
        storage_service.delete_task(task_id)
        return {"message": "Task deleted"}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Task not found")

class TaskOrderRequest(BaseModel):
    task_ids: List[str]

@app.put("/tasks/order")
async def update_tasks_order(req: TaskOrderRequest):
    """Saves the manual order computed on the frontend"""
    return storage_service.reorder_tasks(req.task_ids)

@app.post("/tasks/sort")
async def ai_sort_tasks():
    """Use AI to suggest an optimized, prioritized order for open tasks."""
    tasks = storage_service.get_tasks()
    open_tasks = [t for t in tasks if t.get("status") != "done"]
    if not open_tasks:
        return {"order": [], "reasoning": "Nessuna task aperta da ordinare."}
    result = await ai_service.sort_tasks(open_tasks)
    return result


# ═══════════════════════════════════════════════
#  REPORTS (per-client)
# ═══════════════════════════════════════════════

class ReportCreate(BaseModel):
    period_label: Optional[str] = ""       # es: "Febbraio 2026"
    budget_speso: Optional[str] = ""
    roas: Optional[str] = ""
    ctr: Optional[str] = ""
    cpc: Optional[str] = ""
    cpm: Optional[str] = ""
    conversioni: Optional[str] = ""
    revenue: Optional[str] = ""
    reach: Optional[str] = ""
    impressions: Optional[str] = ""
    note: Optional[str] = ""
    best_angles: Optional[str] = ""
    best_creatives: Optional[str] = ""
    best_copy: Optional[str] = ""

@app.get("/clients/{client_id}/reports")
async def list_reports(client_id: str):
    try:
        storage_service.get_metadata(client_id)  # verify client exists
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")
    return storage_service.get_reports(client_id)

@app.post("/clients/{client_id}/reports")
async def create_report(client_id: str, report: ReportCreate):
    try:
        storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")
    data = report.dict(exclude_none=True)
    return storage_service.save_report(client_id, data)

@app.post("/clients/{client_id}/reports/{report_id}/generate")
async def generate_report_ai(client_id: str, report_id: str):
    try:
        metadata = storage_service.get_metadata(client_id)
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Client not found")
    reports = storage_service.get_reports(client_id)
    report = next((r for r in reports if r["id"] == report_id), None)
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")

    # Carica contesto strategico per report più intelligenti
    supabase = _get_sb()
    strategic_context = await get_strategic_context_for_generator(
        client_id=client_id, metadata=metadata, supabase_client=supabase,
        focus_areas=["brand_identity", "customer_personas", "reasons_to_buy"]
    )
    metadata["_strategic_context"] = strategic_context

    try:
        ai_text = await ai_service.generate_performance_report(metadata, report)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI error: {str(e)}")
    updated = storage_service.update_report(client_id, report_id, {"ai_report": ai_text})
    return updated

@app.delete("/clients/{client_id}/reports/{report_id}")
async def delete_report(client_id: str, report_id: str):
    try:
        storage_service.delete_report(client_id, report_id)
        return {"message": "Report deleted"}
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Report not found")


# ═══════════════════════════════════════════════
# ═══════════════════════════════════════════════
#  IMAGE GENERATION (Fal.ai)
# ═══════════════════════════════════════════════

FAL_KEY = os.getenv("FAL_KEY", "")


class ImageReference(BaseModel):
    type: str  # logo, persona, prodotto, contesto
    data: str  # base64 string
    mime: str  # image/png, image/jpeg etc.

class ImageGenerateRequest(BaseModel):
    prompt: str
    client_id: str
    format: str = "Feed IG (1:1)"
    references: List[ImageReference] = []
    use_rag: bool = False
    model_id: str = "fal-ai/flux-pro/v1.1-ultra"
    reference_filename: Optional[str] = None # For "Modifica" or "Aggiungi Formato" from existing gallery


def save_graphic_metadata(client_id: str, filename: str, prompt: str, enhanced_prompt: str, format: str, original_url: str = ""):
    """Helper to save metadata to graphics_meta.json uniformly."""
    graphics_dir = CLIENTS_DIR / client_id / "graphics"
    graphics_dir.mkdir(parents=True, exist_ok=True)
    meta_file = graphics_dir / "graphics_meta.json"
    meta_list = []
    if meta_file.exists():
        try:
            meta_list = json.loads(meta_file.read_text())
        except:
            meta_list = []
            
    meta_list.insert(0, {
        "filename": filename,
        "prompt": prompt,
        "enhanced_prompt": enhanced_prompt,
        "format": format,
        "created_at": datetime.now().isoformat(),
        "original_url": original_url
    })
    meta_file.write_text(json.dumps(meta_list, indent=2, ensure_ascii=False))
    storage_service.sync_graphics_meta(client_id, meta_list)

@app.post("/generate-image")
async def generate_image(request: ImageGenerateRequest):
    """Universal image generation endpoint. Claude writes the prompt, the model executes with direct file access."""
    GOOGLE_AI_API_KEY = os.getenv("GOOGLE_AI_API_KEY", "")
    is_gemini = request.model_id in ("gemini-image", "gemini-image-pro")
    is_fal = not is_gemini

    if is_gemini and not GOOGLE_AI_API_KEY:
        raise HTTPException(status_code=500, detail="GOOGLE_AI_API_KEY not configured")
    if is_fal and not FAL_KEY:
        raise HTTPException(status_code=500, detail="FAL_KEY not configured")

    # Load image from disk if reference_filename is provided
    if request.reference_filename:
        ref_path = CLIENTS_DIR / request.client_id / "graphics" / request.reference_filename
        if ref_path.exists():
            with open(ref_path, "rb") as f:
                ref_data = base64.b64encode(f.read()).decode("utf-8")
                # Append to references list so the rest of the logic sees it as an image ref
                request.references.insert(0, ImageReference(type="context", data=ref_data, mime="image/png"))

    # ── STEP 1: Load client context (shared for ALL models) ───────────────────
    metadata = {}
    client_context = ""
    try:
        metadata = storage_service.get_metadata(request.client_id)

        # 🔥 NUOVO: Carica TUTTA l'Analisi Strategica (14 sezioni)
        supabase = _get_sb()
        client_context = await get_strategic_context_for_generator(
            client_id=request.client_id,
            metadata=metadata,
            supabase_client=supabase,
            focus_areas=["visual_brief", "brand_voice", "brand_identity", "psychographic_analysis", "customer_personas", "reviews_voc", "product_vertical", "service_vertical"]
        )
    except Exception as e:
        print(f"Error loading strategic context for graphics: {e}")
        # Fallback to basic context if strategic context fails
        try:
            brand = metadata.get("brand_identity", {})
            client_context = f"""Cliente: {metadata.get('name', '')}
Settore: {metadata.get('industry', 'non specificato')}
Tono: {brand.get('tone', 'non specificato')}
Colori brand: {', '.join(brand.get('colors', []))}
Stile visivo: {brand.get('visuals', 'non specificato')}"""
        except Exception:
            pass

    # ── STEP 2: Fetch Notion RAG data if requested ────────────────────────────
    rag_context_text = ""
    if request.use_rag:
        try:
            bm = metadata.get("industry") or metadata.get("business_type")
            rag_creatives = await notion_service.get_winning_creatives(business_model=bm, limit=3)
            if rag_creatives:
                rag_context_text = "\n\n=== ADS AD ALTE PERFORMANCE NEL SETTORE ===\nPrendi ispirazione da questi layout e angoli testati:\n"
                for i, cr in enumerate(rag_creatives):
                    rag_context_text += f"\nEsempio {i+1}: {cr.get('headline', '')} — {cr.get('notes', '')}\n"
        except Exception as e:
            print(f"RAG fetch failed: {e}")

    # ── STEP 3: Claude writes the creative prompt ─────────────────────────────
    # Claude's ONLY job: turn the user's idea into the perfect prompt for the chosen model.
    # Claude NEVER describes the uploaded images — those go directly to the model as pixels.
    n_refs = len(request.references)
    ref_note = ""
    if n_refs == 1:
        ref_note = f"NOTA: L'utente ha caricato 1 immagine di riferimento (Immagine 1). Il modello la vedrà direttamente. Nel tuo prompt scrivi dove/come deve usarla (es. 'using the person in the reference as the model', 'place the product from the reference image on the bathroom counter')."
    elif n_refs >= 2:
        ref_note = f"NOTA: L'utente ha caricato {n_refs} immagini di riferimento (Immagine 1 = Persona, Immagine 2 = Prodotto). Il modello le vedrà direttamente come pixel. Nel tuo prompt scrivi cosa fare con loro (es. 'use Image 1 as the woman in the ad', 'place the exact product from Image 2 in her hand'). NON descrivere l'aspetto delle immagini — il modello le ha già."

    # Select model name for the prompt system message
    model_names = {
        "gemini-image": "Google Gemini 3.1 Flash Image",
        "gemini-image-pro": "Google Gemini 3 Pro Image",
        "fal-ai/flux-pro/v1.1-ultra": "Flux Pro 1.1 Ultra (Black Forest Labs)",
        "fal-ai/flux-realism": "Flux Realism (Black Forest Labs)",
    }
    model_label = model_names.get(request.model_id, request.model_id)

    enhanced_prompt = request.prompt
    try:
        system_msg = f"""Sei il Creative Director di un'agenzia pubblicitaria di lusso. Il tuo UNICO compito in questo momento è scrivere il prompt perfetto per il modello AI scelto ({model_label}) che genererà una creative pubblicitaria.

FLUSS DI LAVORO:
1. Leggi l'idea dell'utente e il contesto del cliente
2. Se l'utente chiede di usare headline o hook dall'analisi → ESTRAILE dal contesto (sezioni 【10】 VOICE OF CUSTOMER, 【7】 ANALISI VERTICALE PRODOTTI, 【7b】 ANALISI VERTICALE SERVIZI) e includile LETTERALMENTE nel prompt come testo overlay.
3. Consulta gli esempi di ads vincenti nel settore (se presenti)
4. Scrivi un prompt INGLESE ottimizzato per {model_label}

REGOLE FONDAMENTALI:
- Il modello AI vedrà le immagini di riferimento direttamente come pixel. NON descrivere l'aspetto fisico delle persone o dei prodotti — usa frasi come "the woman in Image 1", "the product in Image 2".
- Specifica la scena, la composizione, l'illuminazione, il mood, il formato e il testo overlay.
- HEADLINE REALI: Se l'utente dice 'usa una headline dall'analisi' o simili, seleziona il golden hook più rilevante dalla sezione 【10】 VOICE OF CUSTOMER del contesto e includilo come testo overlay in italiano.
- Il testo in italiano (Headline, bullet point, CTA) va incluso nel prompt tra virgolette come overlay tipografico.
- Formato richiesto: {request.format}. Specifica esplicitamente nel prompt: "in {request.format} vertical format, portrait orientation".
- Stile: commercial photography, 8k, photorealistic, NOT illustrated, NOT cartoon.

{ref_note}

Rispondi SOLO con il prompt inglese. Senza spiegazioni, senza prefissi."""

        user_msg_text = f"Contesto brand:\n{client_context}\n{get_awareness_context()}\n{rag_context_text}\n\nFormato richiesto: {request.format}\nModello: {model_label}\n\nIdea dell'utente:\n{request.prompt}"

        user_content = [
            {"type": "text", "text": user_msg_text}
        ]

        # Inject RAG Ad images so Claude can physically see the styles
        for cr in rag_creatives:
            if cr.get("thumbnail_b64"):
                user_content.append({
                    "type": "image_url",
                    "image_url": {
                        "url": f"data:{cr.get('thumbnail_mime', 'image/jpeg')};base64,{cr.get('thumbnail_b64')}"
                    }
                })

        # Inject direct reference images so Claude understands the composition request
        # NOTE: Claude does NOT describe these — it only writes how to position them in the scene
        for ref in request.references:
            user_content.append({
                "type": "image_url",
                "image_url": {"url": f"data:{ref.mime};base64,{ref.data}"}
            })

        enhanced_prompt = await ai_service._call_ai(
            model="anthropic/claude-sonnet-4-5",
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_content}
            ],
            temperature=0.7,
            max_tokens=800
        )
    except Exception as e:
        print(f"Claude enhancement failed, using original prompt: {e}")

    # ── STEP 4a: Execute on Google Gemini (if gemini model selected) ────────────
    if is_gemini:
        gemini_model = (
            "gemini-3-pro-image-preview" if request.model_id == "gemini-image-pro"
            else "gemini-3.1-flash-image-preview"
        )
        # Format instruction injected into the prompt text
        fmt_low = request.format.lower()
        if "4:5" in fmt_low:
            ar_note = "IMPORTANT: Generate this in 4:5 portrait format (tall vertical image, like Instagram feed). "
        elif "9:16" in fmt_low or "stories" in fmt_low:
            ar_note = "IMPORTANT: Generate this in 9:16 portrait format (phone Stories ratio). "
        elif "16:9" in fmt_low or "banner" in fmt_low:
            ar_note = "IMPORTANT: Generate this in 16:9 wide landscape format. "
        else:
            ar_note = ""

        # Build Gemini parts: inline images first, then Claude-generated prompt
        gemini_parts = []
        for ref in request.references:
            gemini_parts.append({"inline_data": {"mime_type": ref.mime, "data": ref.data}})
        gemini_parts.append({"text": ar_note + enhanced_prompt})

        gemini_payload = {
            "contents": [{"role": "user", "parts": gemini_parts}],
            "generationConfig": {"responseModalities": ["TEXT", "IMAGE"]}
        }
        gemini_url = (
            f"https://generativelanguage.googleapis.com/v1beta/models/{gemini_model}:generateContent"
            f"?key={GOOGLE_AI_API_KEY}"
        )
        print(f"Gemini {gemini_model}: Claude prompt ready, {len(request.references)} images as inline_data")
        async with httpx.AsyncClient(timeout=180.0) as gclient:
            try:
                resp = await gclient.post(gemini_url, headers={"Content-Type": "application/json"}, json=gemini_payload)
                resp.raise_for_status()
            except httpx.HTTPStatusError as e:
                raise HTTPException(status_code=502, detail=f"Gemini error: {e.response.text[:500]}")

        image_b64 = None
        for candidate in resp.json().get("candidates", []):
            for part in candidate.get("content", {}).get("parts", []):
                if "inlineData" in part:
                    image_b64 = part["inlineData"]["data"]
                    break
        if not image_b64:
            raise HTTPException(status_code=502, detail="Gemini did not return an image")

        client_dir = CLIENTS_DIR / request.client_id / "graphics"
        client_dir.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{timestamp}_{uuid.uuid4().hex[:6]}.png"
        img_bytes = base64.b64decode(image_b64)
        with open(client_dir / filename, "wb") as f:
            f.write(img_bytes)
        storage_service.save_graphic_to_supabase(request.client_id, filename, img_bytes)

        # Save metadata (was missing for Gemini!)
        save_graphic_metadata(
            request.client_id, filename, request.prompt, enhanced_prompt, request.format
        )

        return {"url": f"/clients/{request.client_id}/graphics/{filename}", "filename": filename, "enhanced_prompt": enhanced_prompt, "model_used": gemini_model}

    # ── STEP 4b: Execute on Fal.ai (Flux models only) ────────────────────────
    # Flux models: use Redux for identity+scene, standard for text-to-image
    fmt_low = request.format.lower()
    actual_model = request.model_id
    if request.references and "flux" in request.model_id.lower():
        actual_model = "fal-ai/flux-pro/v1.1-ultra/redux"

    fal_url = f"https://fal.run/{actual_model}"
    payload = {"prompt": enhanced_prompt, "output_format": "png"}

    # Pixel-perfect image sizes for Flux
    if "9:16" in fmt_low or "stories" in fmt_low:
        payload["image_size"] = {"width": 720, "height": 1280}
    elif "16:9" in fmt_low or "banner" in fmt_low:
        payload["image_size"] = {"width": 1280, "height": 720}
    elif "4:5" in fmt_low:
        payload["image_size"] = {"width": 1024, "height": 1280}
    elif "2:3" in fmt_low or "pinterest" in fmt_low:
        payload["image_size"] = {"width": 854, "height": 1280}
    else:
        payload["image_size"] = {"width": 1024, "height": 1024}

    if request.references:
        import random
        payload["image_url"] = f"data:{request.references[0].mime};base64,{request.references[0].data}"
        payload["strength"] = 0.60
        payload["seed"] = random.randint(0, 2**32 - 1)

    # 4. Call Fal.ai REST API - always use JSON
    fal_headers = {"Authorization": f"Key {FAL_KEY}", "Content-Type": "application/json"}

    async with httpx.AsyncClient(timeout=180.0) as client:
        try:
            resp = await client.post(fal_url, headers=fal_headers, json=payload)
            resp.raise_for_status()
        except httpx.HTTPStatusError as e:
            error_detail = e.response.text[:500] if e.response else str(e)
            raise HTTPException(status_code=502, detail=f"Fal.ai error: {error_detail}")
        except Exception as e:
            raise HTTPException(status_code=502, detail=f"Fal.ai network error: {str(e)}")

    result = resp.json()

    # 5. Extract generated image URL
    image_url = None
    if "images" in result and len(result["images"]) > 0:
        image_url = result["images"][0].get("url")
    elif "image" in result and isinstance(result["image"], dict):
        image_url = result["image"].get("url")
        
    if not image_url:
        raise HTTPException(status_code=500, detail=f"No image returned from Fal.ai. Response: {str(result)[:300]}")

    # 6. Download the image from Fal.ai url and save strictly to disk
    graphics_dir = CLIENTS_DIR / request.client_id / "graphics"
    graphics_dir.mkdir(parents=True, exist_ok=True)

    filename = f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}.png"
    filepath = graphics_dir / filename
    
    async with httpx.AsyncClient(timeout=60.0) as img_client:
        try:
            img_resp = await img_client.get(image_url)
            img_resp.raise_for_status()
            with open(filepath, "wb") as f:
                f.write(img_resp.content)
            storage_service.save_graphic_to_supabase(request.client_id, filename, img_resp.content)
        except Exception as e:
             raise HTTPException(status_code=500, detail=f"Failed to download generated image from Fal.ai: {str(e)}")

    save_graphic_metadata(
        request.client_id, filename, request.prompt, enhanced_prompt, request.format, image_url
    )

    return {
        "filename": filename,
        "prompt": request.prompt,
        "enhanced_prompt": enhanced_prompt,
        "response_text": "",
        "url": f"/clients/{request.client_id}/graphics/{filename}",
        "original_url": image_url
    }


@app.get("/clients/{client_id}/graphics")
def list_graphics(client_id: str):
    """List all generated graphics for a client."""
    graphics_dir = CLIENTS_DIR / client_id / "graphics"
    meta_file = graphics_dir / "graphics_meta.json"
    if meta_file.exists():
        return json.loads(meta_file.read_text())
    return []


@app.get("/clients/{client_id}/graphics/{filename}")
def get_graphic(client_id: str, filename: str):
    """Serve a generated graphic file."""
    filepath = CLIENTS_DIR / client_id / "graphics" / filename
    if not filepath.exists():
        storage_service.ensure_local_graphic(client_id, filename)
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="Graphic not found")
    return FileResponse(filepath, media_type="image/png")

@app.delete("/clients/{client_id}/graphics/{filename}")
def delete_graphic(client_id: str, filename: str):
    """Delete a graphic file and its metadata."""
    graphics_dir = CLIENTS_DIR / client_id / "graphics"
    filepath = graphics_dir / filename
    meta_file = graphics_dir / "graphics_meta.json"
    
    # 1. Delete actual image
    if filepath.exists():
        filepath.unlink()
    storage_service.delete_graphic_from_supabase(client_id, filename)

    # 2. Update metadata
    meta_list = []
    if meta_file.exists():
        try:
            meta_list = json.loads(meta_file.read_text())
            meta_list = [m for m in meta_list if m.get("filename") != filename]
            meta_file.write_text(json.dumps(meta_list, indent=2, ensure_ascii=False))
        except Exception as e:
            print(f"Error updating graphics meta: {e}")
    storage_service.sync_graphics_meta(client_id, meta_list)

    return {"message": "Graphic deleted"}

@app.post("/knowledge")
async def add_knowledge(req: KnowledgeRequest):
    """Save a rule or framework directly to Notion Knowledge Base DB."""
    success = await notion_service.save_to_knowledge_base(
        kb_type=req.kb_type,
        name=req.name,
        instructions=req.instructions,
        funnel_stage=req.funnel_stage
    )
    if not success:
        raise HTTPException(status_code=500, detail="Failed to save to Knowledge Base")
    return {"message": "Success"}

@app.get("/knowledge/creatives/{filename}")
def serve_creative_image(filename: str):
    """Serve a locally cached creative image from the knowledge base."""
    filepath = Path("backend/storage/knowledge/winning_creatives") / filename
    if not filepath.exists():
        # Try absolute path
        filepath = CLIENTS_DIR.parent / "knowledge" / "winning_creatives" / filename
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="File not found")
    
    # Determine media type
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "jpg"
    media_types = {"jpg": "image/jpeg", "jpeg": "image/jpeg", "png": "image/png", "gif": "image/gif", "webp": "image/webp", "mp4": "video/mp4"}
    return FileResponse(filepath, media_type=media_types.get(ext, "image/jpeg"))


# ══════════════════════════════════════════════════════════════════════
#  COMPLETE ANALYSIS — Analisi completa secondo metodologia guida
# ══════════════════════════════════════════════════════════════════════

@app.post("/clients/{client_id}/analysis/complete")
async def generate_complete_client_analysis(client_id: str):
    """
    🔥 Lancia l'analisi in background e risponde subito con job_id.
    Il frontend fa polling su /analysis/status/{job_id} per sapere quando è pronta.
    """
    job_id = str(uuid.uuid4())
    _analysis_jobs[job_id] = {"status": "running", "progress": "Avvio raccolta dati..."}

    async def _run_analysis():
        try:
            await _do_complete_analysis(client_id, job_id)
        except Exception as e:
            print(f"❌ Errore job {job_id}: {e}")
            _analysis_jobs[job_id] = {"status": "error", "error": str(e)}

    asyncio.create_task(_run_analysis())
    return {"status": "started", "job_id": job_id}


async def _do_complete_analysis(client_id: str, job_id: str):
    """Logica reale dell'analisi — gira in background."""
    def _progress(msg: str):
        if job_id in _analysis_jobs:
            _analysis_jobs[job_id]["progress"] = msg

    metadata = storage_service.get_metadata(client_id)
    client_info = {
        "id": client_id,
        "name": metadata.get("name", ""),
        "industry": metadata.get("industry", ""),
        "location": metadata.get("location", ""),
        "metadata": metadata
    }

    # 🔥 RACCOLTA DATI MASSIVA
    from .data_collection_service import DataCollectionService

    data_collector = DataCollectionService(ai_service)

    print(f"\n{'='*80}")
    print(f"🚀 ANALISI STRATEGICA COMPLETA: {client_info['name']}")
    print(f"{'='*80}\n")

    # Avvia cost tracking per questa operazione
    ai_service.start_cost_tracking("analisi_strategica")

    # Raccogli TUTTI i dati
    _progress("Raccolta dati: sito web, social, recensioni, competitor…")
    all_data = await data_collector.collect_all_data(client_id, metadata)
    _progress("Dati raccolti. Preparazione analisi strategica…")

    site_url = all_data.get("site_url", "")
    if not site_url:
        print("⚠️ Nessun sito web trovato nei link forniti. L'analisi proseguirà con i dati disponibili (recensioni, competitor, documenti).")
    
    # Estrai i dati raccolti
    site_content = all_data.get("site_content", {})
    google_reviews = all_data.get("google_reviews", {})
    instagram_data = all_data.get("instagram_data", {})
    meta_ads = all_data.get("meta_ads", {})
    competitor_data = all_data.get("competitor_data", {})

    # Raccogli documenti caricati (raw-data folder)
    raw_docs = ""
    scraped_products = all_data.get("products_txt") or ""
    products_csv = scraped_products if scraped_products else "Non disponibili"
    services_txt = all_data.get("services_txt") or "Non disponibili"

    raw_files = list((CLIENTS_DIR / client_id / "raw_data").glob("*")) if (CLIENTS_DIR / client_id / "raw_data").exists() else []
    for file_path in raw_files:
        try:
            file_name_lower = file_path.name.lower()

            # Cerca CSV (prodotti Shopify, listini, etc.) — qualsiasi CSV viene trattato come dati prodotto
            if file_name_lower.endswith('.csv'):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    csv_content = f.read()[:10000]  # Limite 10K caratteri per CSV
                    if products_csv == "Non disponibili":
                        products_csv = csv_content
                    else:
                        products_csv += f"\n\n--- Dati extra ({file_path.name}) ---\n{csv_content}"
                    print(f"✅ Trovato CSV prodotti: {file_path.name}")

            # Cerca TXT servizi
            elif file_name_lower.endswith('.txt') and ('serviz' in file_name_lower or 'service' in file_name_lower):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    services_txt = f.read()[:8000]  # Limite 8K caratteri per TXT
                    print(f"✅ Trovato TXT servizi: {file_path.name}")

            # Cerca PDF (Analisi manuali, Knowledge Base)
            elif file_name_lower.endswith('.pdf'):
                try:
                    reader = pypdf.PdfReader(file_path)
                    text = ""
                    for page in reader.pages[:20]: # Leggi prime 20 pagine
                        text += page.extract_text() + "\n"
                    raw_docs += f"\n\n--- PDF: {file_path.name} ---\n{text[:20000]}"
                    print(f"✅ Estratto testo da PDF: {file_path.name} ({len(text)} caratteri)")
                except Exception as e:
                    print(f"❌ Errore lettura PDF {file_path.name}: {e}")

            # Cerca DOC/DOCX (Dossier, documenti Word)
            elif file_name_lower.endswith(('.doc', '.docx')):
                try:
                    doc = docx.Document(file_path)
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    raw_docs += f"\n\n--- DOCX: {file_path.name} ---\n{text[:20000]}"
                    print(f"✅ Estratto testo da DOCX: {file_path.name} ({len(text)} caratteri)")
                except Exception as e:
                    print(f"❌ Errore lettura DOCX {file_path.name}: {e}")

            # Altri documenti generici (TXT)
            elif file_name_lower.endswith(('.txt', '.md', '.json')):
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read()[:10000]
                    raw_docs += f"\n\n--- {file_path.name} ---\n{content}"
                    print(f"✅ Letto documento: {file_path.name}")
        except Exception as e:
            print(f"⚠️ Errore processamento file {file_path.name}: {e}")

    # 🔥 ESTRAI CONTENUTO DEL SITO — Approccio universale senza keyword filtering
    # Tutte le pagine scrappate vengono passate all'AI che identifica autonomamente prodotti e servizi.
    if isinstance(site_content, dict):
        pages_data = site_content.get("pages", [])
        all_site_pages = []

        # Pagine tecniche da escludere sempre (non contengono info su prodotti/servizi)
        SKIP_PATHS = ["privacy", "cookie", "legal", "terms", "login", "cart",
                      "checkout", "wp-", "xmlrpc", "feed", "sitemap", "robots",
                      "tag/", "author/", "page/", "404", "search"]

        for page in pages_data:
            if not isinstance(page, dict) or "url" not in page or "data" not in page:
                continue
            url = page["url"]
            path = url.replace(site_url, "").strip("/")

            # Salta pagine tecniche
            if any(skip in path.lower() for skip in SKIP_PATHS):
                continue

            if isinstance(page["data"], dict):
                raw_text = page["data"].get("raw_text", "")
                if raw_text and len(raw_text) > 80:
                    all_site_pages.append(f"--- PAGINA: {url} ---\n{raw_text}")

        # 🔥 PREPARAZIONE DATI SITO (CHI È IL CLIENTE - CONTESTO BRAND)
        # Queste pagine servono per l'analisi completa dell'identità, SWOT, ecc.
        full_site_dump = "\n\n".join(all_site_pages)[:25000] if all_site_pages else ""

        # 🔥 ANALISI SERVIZI (FISSAZIONE SULLA LANDING SPECIFICA)
        # Se 'services_txt' ha già contenuto (dai link marchiati 'Landing Servizio' dal menu a tendina), 
        # NON lo sovrascriviamo col dump del sito. Lo teniamo pulito e specifico come richiesto.
        if services_txt == "Non disponibili" and full_site_dump:
            services_txt = f"[CONTENUTO SITO WEB]\n\n{full_site_dump}"
            print("✅ Nessuna landing specifica: uso [CONTENUTO SITO WEB] come fallback per i servizi")
        elif services_txt != "Non disponibili":
            # Aggiungiamo il dump del sito solo in coda come CONTESTO aggiuntivo, ma non come fonte principale
            services_txt = f"{services_txt}\n\n--- CONTESTO BRAND AGGIUNTIVO ---\n{full_site_dump[:5000]}"
            print(f"✅ Landing specifica trovata: focus chirurgico su quella. Dump sito aggiunto solo come contesto.")

        # Per i prodotti, se non c'è CSV/Documento, usiamo il dump come fallback
        if products_csv == "Non disponibili" and full_site_dump:
            products_csv = full_site_dump
            print("✅ Prodotti: nessun documento trovato, uso il dump del sito come fallback")

    # 1. Flatten Instagram Comments for easier Review Mining
    flattened_ig = []
    if isinstance(instagram_data, dict) and "posts" in instagram_data:
        for post in instagram_data["posts"]:
            for comment in post.get("all_comments", []):
                text = comment.get("text", "").strip()
                if text:
                    flattened_ig.append(f"IG - {comment.get('username', 'User')}: {text}")
    
    instagram_comments_text = "\n".join(flattened_ig) if flattened_ig else "Nessun commento Instagram trovato."

    # 2. Flatten Google Reviews (Include ora Extra Reviews da Miner Multi-Fonte)
    flattened_google = []
    if isinstance(google_reviews, dict):
        # Handle merged structure (all_reviews) or original
        reviews_list = google_reviews.get("all_reviews", []) or google_reviews.get("reviews", [])
        
        if not reviews_list and "raw_text" in google_reviews:
            google_reviews_text = google_reviews["raw_text"]
        else:
            for rev in reviews_list:
                text = rev.get("text") or rev.get("content") or ""
                if text:
                    stars = rev.get("stars") or rev.get("rating") or "5"
                    author = rev.get("author") or rev.get("user") or "Anonimo"
                    flattened_google.append(f"({stars} stelle) - {author}: {text}")
            google_reviews_text = "\n".join(flattened_google) if flattened_google else "Nessuna recensione trovata nelle fonti fornite."

    # 3. Serialize remaining data
    import json
    site_content_text = json.dumps(site_content, ensure_ascii=False) if isinstance(site_content, dict) else str(site_content)
    ads_text = json.dumps(meta_ads, ensure_ascii=False) if isinstance(meta_ads, dict) else str(meta_ads)
    competitor_text = json.dumps(competitor_data, ensure_ascii=False) if isinstance(competitor_data, (dict, list)) else str(competitor_data)

    # 🔥 CHIAMA L'ORCHESTRATOR CON TUTTI I DATI RACCOLTI
    print(f"\n{'='*80}")
    print(f"🧠 GENERAZIONE ANALISI STRATEGICA")
    print(f"{'='*80}\n")

    # 🔥 SALVA SNAPSHOT DATI GREZZI (per rigenerazione veloce)
    # Rinominiamo alcune chiavi per coerenza nel contesto
    raw_data_snapshot = {
        "site_url": site_url,
        "site_content": site_content_text,
        "social_data": instagram_comments_text,
        "ads_data": ads_text,
        "raw_docs": raw_docs,
        "google_reviews": google_reviews_text,
        "instagram_comments": instagram_comments_text,
        "products_csv": products_csv,
        "services_txt": services_txt,
        "competitor_data": competitor_text,
        "timestamp": str(asyncio.get_event_loop().time())
    }
    metadata["raw_data_snapshot"] = raw_data_snapshot

    print(f"🧠 Avvio analisi completa...")
    _progress("Avvio generazione 18 sezioni strategiche…")
    complete_analysis = await ai_service.generate_complete_analysis(
        client_info=client_info,
        site_url=site_url,
        site_content=site_content_text,
        social_data=instagram_comments_text,
        ads_data=ads_text,
        raw_docs=raw_docs,
        google_reviews=google_reviews_text,
        instagram_comments=instagram_comments_text,
        products_csv=products_csv,
        services_txt=services_txt,
        competitor_data=competitor_text,
        progress_callback=_progress
    )

    # Backup dell'analisi completa originale per la rigenerazione
    metadata["analysis_completa_raw"] = complete_analysis 

    # Aggiungi dati competitor all'analisi per salvataggio
    complete_analysis["competitor_data"] = competitor_data

    # 🔥 AGGIORNA METADATA CON I NUOVI DATI (SWOT, OBIETTIVI, STRATEGIA, PERSONAS, TONE)
    if "brand_identity" not in metadata:
        metadata["brand_identity"] = {}
        
    if complete_analysis.get("swot"):
        metadata["swot"] = complete_analysis["swot"]
    if complete_analysis.get("objectives"):
        metadata["objectives"] = complete_analysis["objectives"]
    if complete_analysis.get("strategy"):
        metadata["strategy"] = complete_analysis["strategy"]
    if complete_analysis.get("customer_personas"):
        # Convert generated personas to exactly what the tab expects
        metadata["brand_identity"]["buyer_personas"] = complete_analysis["customer_personas"]
    
    if complete_analysis.get("brand_identity") and isinstance(complete_analysis["brand_identity"].get("tone_of_voice"), dict):
        t = complete_analysis["brand_identity"]["tone_of_voice"]
        style_text = f"**Stile**: {t.get('style', '')}\n**Pubblico Target**: {t.get('target_audience', '')}\n**Approccio Linguistico**: {t.get('linguistic_approach', '')}"
        if t.get("vocabulary"):
            style_text += f"\n**Vocabolario**: {', '.join(t['vocabulary']) if isinstance(t['vocabulary'], list) else t['vocabulary']}"
        metadata["brand_identity"]["tone"] = style_text

    # Salva metadata aggiornato
    storage_service.save_metadata(client_id, metadata)
    print("✅ Metadata aggiornato con SWOT, Obiettivi e Strategia")

    # Salva su file locale SEMPRE (backup sicuro — non si perde mai)
    analysis_file = CLIENTS_DIR / client_id / "complete_analysis.json"
    try:
        analysis_file.write_text(json.dumps(complete_analysis, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"✅ Analisi salvata su file locale: {analysis_file}")
    except Exception as e:
        print(f"❌ Errore salvataggio file locale: {e}")

    # Salva in Supabase — con fallback se alcune colonne non esistono ancora
    try:
        supabase = _get_sb()
        if not supabase:
            print("⚠️  Supabase non configurato - analisi NON salvata nel database")
        else:
            # IMPORTANTE: usa `or {}` / `or []` — .get() può restituire None
            # se la chiave esiste ma il valore è null, e Supabase rifiuta null
            # per colonne NOT NULL
            upsert_data = {
                "client_id": client_id,
                "brand_identity": complete_analysis.get("brand_identity") or {},
                "brand_values": complete_analysis.get("brand_values") or {},
                "product_portfolio": complete_analysis.get("product_portfolio") or {},
                "reasons_to_buy": complete_analysis.get("reasons_to_buy") or {},
                "customer_personas": complete_analysis.get("customer_personas") or [],
                "content_matrix": complete_analysis.get("content_matrix") or [],
                "product_vertical": complete_analysis.get("product_vertical") or [],
                "service_vertical": complete_analysis.get("service_vertical") or [],
                "brand_voice": complete_analysis.get("brand_voice") or {},
                "objections": complete_analysis.get("objections") or {},
                "reviews_voc": complete_analysis.get("reviews_voc") or {},
                "battlecards": complete_analysis.get("battlecards") or {},
                "seasonal_roadmap": complete_analysis.get("seasonal_roadmap") or {},
                "psychographic_analysis": complete_analysis.get("psychographic_analysis") or {},
                "visual_brief": complete_analysis.get("visual_brief") or {},
                "ad_copy_creation": complete_analysis.get("ad_copy_creation") or {},
                "video_scripts": complete_analysis.get("video_scripts") or {},
                "franzcopy_scaling": complete_analysis.get("franzcopy_scaling") or {},
                "swot": complete_analysis.get("swot") or {},
                "objectives": complete_analysis.get("objectives") or {},
                "strategy": complete_analysis.get("strategy") or "",
            }
            result = supabase.table("client_complete_analysis").upsert(upsert_data).execute()
            if result.data:
                print(f"✅ Analisi salvata in Supabase per {client_id} ({len(result.data)} rows)")
            else:
                print(f"⚠️ Supabase upsert non ha restituito dati per {client_id} — possibile errore silenzioso")
    except Exception as e:
        import traceback
        print(f"❌ Errore salvataggio Supabase per {client_id}: {e}")
        traceback.print_exc()

    # Stop cost tracking e salva
    cost_data = ai_service.stop_cost_tracking()
    if cost_data and cost_data.get("total_cost_usd", 0) > 0:
        print(f"💰 Costo totale analisi: ${cost_data['total_cost_usd']:.4f} ({cost_data['calls']} chiamate AI)")
        # Salva nel metadata del client
        if "ai_costs" not in metadata:
            metadata["ai_costs"] = []
        from datetime import datetime
        metadata["ai_costs"].append({
            "operation": cost_data["operation"],
            "cost_usd": round(cost_data["total_cost_usd"], 4),
            "prompt_tokens": cost_data["total_prompt_tokens"],
            "completion_tokens": cost_data["total_completion_tokens"],
            "calls": cost_data["calls"],
            "models": cost_data["models_used"],
            "date": datetime.now().isoformat()
        })
        storage_service.save_metadata(client_id, metadata)

    # Marca il job come completato nel registry
    _analysis_jobs[job_id] = {
        "status": "done",
        "analysis": complete_analysis,
        "progress": "Analisi completata!",
        "cost": cost_data
    }
    print(f"✅ Job {job_id} completato")


# ── Status endpoint per polling dal frontend ──────────────────────────────────
@app.get("/clients/{client_id}/analysis/status/{job_id}")
async def get_analysis_job_status(client_id: str, job_id: str):
    """Polling endpoint: controlla lo stato di un job di analisi in background."""
    job = _analysis_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job non trovato")
    return job


# ══════════════════════════════════════════════════════════════════════════════
#  AI COSTS — Costi AI per client e globali
# ══════════════════════════════════════════════════════════════════════════════

@app.get("/clients/{client_id}/ai-costs")
async def get_client_ai_costs(client_id: str):
    """Restituisce lo storico dei costi AI per un client."""
    metadata = storage_service.get_metadata(client_id)
    costs = metadata.get("ai_costs", [])
    total = sum(c.get("cost_usd", 0) for c in costs)
    return {"client_id": client_id, "costs": costs, "total_usd": round(total, 4)}


@app.get("/ai-costs/summary")
async def get_ai_costs_summary():
    """Restituisce il riepilogo costi AI di tutti i client."""
    all_clients = storage_service.list_clients()
    summary = []
    grand_total = 0.0
    for client in all_clients:
        cid = client.get("id", "")
        meta = storage_service.get_metadata(cid)
        costs = meta.get("ai_costs", [])
        client_total = sum(c.get("cost_usd", 0) for c in costs)
        if costs:
            summary.append({
                "client_id": cid,
                "client_name": meta.get("name", cid),
                "total_usd": round(client_total, 4),
                "operations": len(costs),
                "last_operation": costs[-1].get("date", "") if costs else ""
            })
        grand_total += client_total
    return {"clients": summary, "grand_total_usd": round(grand_total, 4)}


@app.get("/clients/{client_id}/analysis/complete")
async def get_complete_client_analysis(client_id: str):
    """
    Recupera l'analisi completa salvata per il cliente.
    Prova Supabase prima, poi fallback su file locale.
    """
    # 1. Prova Supabase
    try:
        supabase = _get_sb()
        if supabase:
            result = supabase.table("client_complete_analysis").select("*").eq("client_id", client_id).execute()
            if result.data and len(result.data) > 0:
                return result.data[0]
    except Exception as e:
        print(f"⚠️ Errore recupero da Supabase: {e}")

    # 2. Fallback: file locale
    analysis_file = CLIENTS_DIR / client_id / "complete_analysis.json"
    if analysis_file.exists():
        try:
            data = json.loads(analysis_file.read_text(encoding="utf-8"))
            print(f"📁 Analisi caricata da file locale per {client_id}")
            return data
        except Exception as e:
            print(f"❌ Errore lettura file locale analisi: {e}")

    return None


@app.delete("/clients/{client_id}/analysis/complete")
async def delete_complete_analysis(client_id: str):
    """
    🗑️ Reset TOTALE: elimina analisi, cache, snapshot, personas generate, brand identity generata.
    Dopo questo, una nuova analisi partirà completamente da zero.
    """
    try:
        # 1. Pulisci metadata (snapshot, raw cache, analisi generata)
        metadata = storage_service.get_metadata(client_id)
        keys_to_remove = [
            "raw_data_snapshot",
            "analysis_completa_raw",
            "swot",
            "objectives",
            "strategy",
        ]
        for key in keys_to_remove:
            metadata.pop(key, None)

        # Pulisci buyer personas generate dall'analisi (ma mantieni quelle manuali se presenti)
        if "brand_identity" in metadata:
            metadata["brand_identity"].pop("buyer_personas", None)
            metadata["brand_identity"].pop("tone", None)

        storage_service.save_metadata(client_id, metadata)
        print(f"🗑️ Metadata pulito per {client_id}: rimossi snapshot, cache, analisi")

        # 2. Elimina da Supabase
        supabase = _get_sb()
        if supabase:
            supabase.table("client_complete_analysis").delete().eq("client_id", client_id).execute()
            print(f"🗑️ Analisi Supabase eliminata per {client_id}")

        # 3. Elimina file locale
        analysis_file = CLIENTS_DIR / client_id / "complete_analysis.json"
        if analysis_file.exists():
            analysis_file.unlink()
            print(f"🗑️ File locale analisi eliminato per {client_id}")

        return {"success": True, "message": "Analisi, cache e dati generati eliminati. Pronto per una nuova analisi da zero."}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/clients/{client_id}/analysis/regenerate/{step_id}")
async def regenerate_analysis_section(client_id: str, step_id: str):
    """
    ♻️ RIGENERA UN SINGOLO BLOCCO - Usa lo snapshot esistente per velocità
    """
    metadata = storage_service.get_metadata(client_id)
    snapshot = metadata.get("raw_data_snapshot")
    
    if not snapshot:
        return {"error": "Nessun dato raccolto trovato. Esegui prima un'analisi completa."}

    client_info = {
        "id": client_id,
        "name": metadata.get("name", ""),
        "industry": metadata.get("industry", ""),
        "location": metadata.get("location", ""),
        "metadata": metadata
    }

    # Carica il Workflow e trova il task specifico
    from .ai_service_strategic_analysis import run_workflow_task
    workflow_path = Path(__file__).parent / "master_workflows" / "agostinis_meta_ads.json"
    import json
    with open(workflow_path, "r", encoding="utf-8") as f:
        workflow = json.load(f)
    
    # Mappatura frontend key → workflow step_id (dove differiscono)
    FRONTEND_TO_STEP = {"objections": "objections_management"}
    workflow_step_id = FRONTEND_TO_STEP.get(step_id, step_id)

    task = next((t for t in workflow["tasks"] if t["step_id"] == workflow_step_id), None)
    if not task:
        raise HTTPException(status_code=404, detail=f"Step {step_id} non trovato")

    # Mappa i risultati precedenti per le dipendenze
    # Usiamo 'analysis_completa_raw' che contiene l'output diretto dell'AI
    context = {
        "client_info": client_info,
        **snapshot,
        **(metadata.get("analysis_completa_raw", {}))
    }

    print(f"♻️ Rigenerazione {step_id} (workflow: {workflow_step_id})...")
    ai_service.start_cost_tracking(f"rigenera_{step_id}")
    new_result = await run_workflow_task(ai_service, task, context)
    cost_data = ai_service.stop_cost_tracking()

    # Salva costo
    if cost_data and cost_data.get("total_cost_usd", 0) > 0:
        print(f"💰 Costo rigenerazione {step_id}: ${cost_data['total_cost_usd']:.4f}")
        if "ai_costs" not in metadata:
            metadata["ai_costs"] = []
        from datetime import datetime
        metadata["ai_costs"].append({
            "operation": cost_data["operation"],
            "cost_usd": round(cost_data["total_cost_usd"], 4),
            "prompt_tokens": cost_data["total_prompt_tokens"],
            "completion_tokens": cost_data["total_completion_tokens"],
            "calls": cost_data["calls"],
            "models": cost_data["models_used"],
            "date": datetime.now().isoformat()
        })

    # Aggiorna il backup raw (usa workflow_step_id per il context chaining)
    if "analysis_completa_raw" not in metadata:
        metadata["analysis_completa_raw"] = {}
    metadata["analysis_completa_raw"][workflow_step_id] = new_result

    # Salva anche nell'analisi completa in Supabase (usa step_id = frontend key)
    storage_service.save_metadata(client_id, metadata)

    try:
        existing_analysis = storage_service.get_complete_analysis(client_id) or {}
        existing_analysis[step_id] = new_result
        storage_service.save_complete_analysis(client_id, existing_analysis)
    except Exception as e:
        print(f"Warning: Could not update Supabase analysis: {e}")

    # Restituisci in formato compatibile con il frontend
    return {
        "step_id": step_id,
        "new_data": new_result,
        "analysis_step": new_result,
        "cost": cost_data
    }


class SectionEditRequest(BaseModel):
    instructions: str


@app.post("/clients/{client_id}/analysis/edit/{step_id}")
async def edit_analysis_section(client_id: str, step_id: str, request: SectionEditRequest):
    """
    ✏️ MODIFICA SEZIONE CON AI — L'utente dà istruzioni testuali e l'AI modifica solo quella sezione.
    Es: "Aggiungi i cocktail alla spina tra i prodotti" oppure "Leggi il nuovo documento e aggiorna"
    """
    metadata = storage_service.get_metadata(client_id)
    snapshot = metadata.get("raw_data_snapshot")

    if not snapshot:
        raise HTTPException(status_code=400, detail="Nessun dato raccolto. Esegui prima un'analisi completa.")

    # Carica sezione attuale da Supabase
    existing_analysis = storage_service.get_complete_analysis(client_id) or {}
    current_section = existing_analysis.get(step_id)

    if not current_section:
        raise HTTPException(status_code=400, detail=f"Sezione {step_id} non ancora generata. Rigenerala prima.")

    # Carica documenti raw del cliente (per se l'utente dice "guarda il documento")
    raw_docs = ""
    raw_data_dir = CLIENTS_DIR / client_id / "raw_data"
    if raw_data_dir.exists():
        for file_path in raw_data_dir.glob("*"):
            try:
                fname = file_path.name.lower()
                if fname.endswith('.pdf'):
                    reader = pypdf.PdfReader(file_path)
                    text = ""
                    for page in reader.pages[:20]:
                        text += page.extract_text() + "\n"
                    raw_docs += f"\n--- PDF: {file_path.name} ---\n{text[:15000]}\n"
                elif fname.endswith(('.doc', '.docx')):
                    doc = docx.Document(file_path)
                    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                    raw_docs += f"\n--- DOCX: {file_path.name} ---\n{text[:15000]}\n"
                elif fname.endswith(('.txt', '.md', '.csv')):
                    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                        raw_docs += f"\n--- {file_path.name} ---\n{f.read()[:10000]}\n"
            except Exception as e:
                print(f"⚠️ Errore lettura {file_path.name} per edit: {e}")

    # Sezione attuale come JSON
    current_json = json.dumps(current_section, ensure_ascii=False, indent=2)

    # Prompt per l'AI
    client_name = metadata.get("name", "")
    industry = metadata.get("industry", "")

    SECTION_LABELS = {
        "brand_identity": "Brand Identity & Posizionamento",
        "brand_values": "Valori del Brand",
        "product_portfolio": "Portafoglio Prodotti",
        "reasons_to_buy": "Reasons to Buy",
        "customer_personas": "Customer Personas",
        "content_matrix": "Matrice Contenuti",
        "product_vertical": "Analisi Verticale Prodotti",
        "service_vertical": "Analisi Verticale Servizi",
        "brand_voice": "Brand Voice & Guidelines",
        "objections": "Gestione Obiezioni",
        "reviews_voc": "Voice of Customer",
        "battlecards": "Competitor Battlecards",
        "seasonal_roadmap": "Roadmap Stagionale",
        "psychographic_analysis": "Analisi Psicografica",
        "visual_brief": "Visual Brief",
    }
    section_label = SECTION_LABELS.get(step_id, step_id)

    messages = [
        {"role": "system", "content": (
            f"Sei un analista strategico. Devi MODIFICARE la sezione '{section_label}' dell'analisi strategica "
            f"del cliente '{client_name}' (settore: {industry}) seguendo le istruzioni dell'utente.\n\n"
            "REGOLE:\n"
            "1. Restituisci la sezione COMPLETA modificata in formato JSON — stessa struttura dell'originale.\n"
            "2. MANTIENI tutto il contenuto esistente che non viene toccato dalle istruzioni.\n"
            "3. AGGIUNGI o MODIFICA solo ciò che l'utente chiede.\n"
            "4. Se l'utente dice di guardare un documento, cerca le info nei DOCUMENTI CLIENTE forniti sotto.\n"
            "5. Restituisci SOLO il JSON della sezione, nessun testo prima o dopo."
        )},
        {"role": "user", "content": (
            f"ISTRUZIONI DELL'UTENTE:\n{request.instructions}\n\n"
            f"SEZIONE ATTUALE ({section_label}):\n{current_json[:12000]}\n\n"
            f"DOCUMENTI CLIENTE (per riferimento):\n{raw_docs[:20000] if raw_docs else '(Nessun documento caricato)'}"
        )}
    ]

    ai_service.start_cost_tracking(f"edit_{step_id}")

    try:
        response_str = await ai_service._call_ai(
            model="anthropic/claude-3.7-sonnet",
            messages=messages,
            max_tokens=6000
        )
        new_section = json_repair.loads(response_str)
    except Exception as e:
        ai_service.stop_cost_tracking()
        raise HTTPException(status_code=500, detail=f"Errore AI: {e}")

    cost_data = ai_service.stop_cost_tracking()

    # Salva costo
    if cost_data and cost_data.get("total_cost_usd", 0) > 0:
        print(f"💰 Costo edit {step_id}: ${cost_data['total_cost_usd']:.4f}")
        if "ai_costs" not in metadata:
            metadata["ai_costs"] = []
        from datetime import datetime
        metadata["ai_costs"].append({
            "operation": cost_data["operation"],
            "cost_usd": round(cost_data["total_cost_usd"], 4),
            "prompt_tokens": cost_data["total_prompt_tokens"],
            "completion_tokens": cost_data["total_completion_tokens"],
            "calls": cost_data["calls"],
            "models": cost_data["models_used"],
            "date": datetime.now().isoformat()
        })

    # Aggiorna raw backup
    FRONTEND_TO_STEP = {"objections": "objections_management"}
    workflow_step_id = FRONTEND_TO_STEP.get(step_id, step_id)
    if "analysis_completa_raw" not in metadata:
        metadata["analysis_completa_raw"] = {}
    metadata["analysis_completa_raw"][workflow_step_id] = new_section
    storage_service.save_metadata(client_id, metadata)

    # Salva in Supabase
    try:
        existing_analysis[step_id] = new_section
        storage_service.save_complete_analysis(client_id, existing_analysis)
    except Exception as e:
        print(f"Warning: Could not update Supabase after edit: {e}")

    print(f"✏️ Sezione {step_id} modificata per {client_id}")
    return {
        "step_id": step_id,
        "new_data": new_section,
        "cost": cost_data
    }


@app.delete("/clients/{client_id}/analysis/section/{step_id}")
async def delete_analysis_section(client_id: str, step_id: str):
    """🗑️ Cancella una singola sezione dell'analisi."""
    FRONTEND_TO_STEP = {"objections": "objections_management"}
    workflow_step_id = FRONTEND_TO_STEP.get(step_id, step_id)

    metadata = storage_service.get_metadata(client_id)

    # Rimuovi dal raw backup
    if "analysis_completa_raw" in metadata and workflow_step_id in metadata["analysis_completa_raw"]:
        del metadata["analysis_completa_raw"][workflow_step_id]

    storage_service.save_metadata(client_id, metadata)

    # Rimuovi da Supabase
    try:
        existing = storage_service.get_complete_analysis(client_id) or {}
        existing[step_id] = None  # Forza il campo a null su Supabase
        storage_service.save_complete_analysis(client_id, existing)
    except Exception as e:
        print(f"Warning: Could not update Supabase: {e}")

    print(f"🗑️ Sezione {step_id} cancellata per {client_id}")
    return {"success": True, "step_id": step_id}


@app.post("/clients/{client_id}/analysis/deepen/{step_id}")
async def deepen_analysis_section(client_id: str, step_id: str):
    """
    🔬 APPROFONDISCI — Ricerca online aggiuntiva + rigenerazione con più contesto.
    Usa Perplexity per ricerca specifica, poi rigenera con dati arricchiti.
    """
    metadata = storage_service.get_metadata(client_id)
    snapshot = metadata.get("raw_data_snapshot")

    if not snapshot:
        raise HTTPException(status_code=400, detail="Nessun dato raccolto. Esegui prima un'analisi completa.")

    client_info = {
        "id": client_id,
        "name": metadata.get("name", ""),
        "industry": metadata.get("industry", ""),
        "location": metadata.get("location", ""),
        "metadata": metadata
    }

    # Carica workflow task
    from .ai_service_strategic_analysis import run_workflow_task
    workflow_path = Path(__file__).parent / "master_workflows" / "agostinis_meta_ads.json"
    import json
    with open(workflow_path, "r", encoding="utf-8") as f:
        workflow = json.load(f)

    FRONTEND_TO_STEP = {"objections": "objections_management"}
    workflow_step_id = FRONTEND_TO_STEP.get(step_id, step_id)
    task = next((t for t in workflow["tasks"] if t["step_id"] == workflow_step_id), None)
    if not task:
        raise HTTPException(status_code=404, detail=f"Step {step_id} non trovato")

    # 🔬 RICERCA ONLINE APPROFONDITA con Perplexity
    ai_service.start_cost_tracking(f"approfondisci_{step_id}")

    client_name = metadata.get("name", "")
    industry = metadata.get("industry", "")
    location = metadata.get("location", "")
    site_url = snapshot.get("site_url", "")

    # Costruisci query di ricerca specifica per questa sezione
    SECTION_RESEARCH_QUERIES = {
        "brand_identity": f"Analisi approfondita brand \"{client_name}\" {location}: posizionamento, storia, mission, valori, target, competitor diretti. Sito: {site_url}",
        "brand_values": f"Valori e filosofia del brand \"{client_name}\" {industry} {location}: cosa li distingue, recensioni clienti, reputazione online",
        "product_portfolio": f"Tutti i prodotti e servizi offerti da \"{client_name}\" {location}: listino, prezzi, offerte speciali, servizi principali. Sito: {site_url}",
        "reasons_to_buy": f"Perché i clienti scelgono \"{client_name}\" {industry} {location}: recensioni, punti di forza, vantaggi rispetto ai competitor",
        "customer_personas": f"Target e clienti tipici di {industry} a {location}: demographics, comportamenti, esigenze, frustrazioni comuni nel settore",
        "brand_voice": f"Tono di voce e comunicazione di \"{client_name}\": come comunica sui social, sul sito, nelle ads. Stile linguistico del settore {industry}",
        "content_matrix": f"Migliori strategie di contenuto per {industry} su Meta/Instagram: format che funzionano, hook efficaci, trend attuali 2024-2025",
        "reviews_voc": f"Recensioni e opinioni clienti su \"{client_name}\" {location}: Google Reviews, TrustPilot, Facebook, Instagram commenti. Cosa dicono i clienti?",
        "battlecards": f"Competitor principali di \"{client_name}\" nel settore {industry} a {location}: chi sono, come si posizionano, punti di forza e debolezza, ads attive",
        "seasonal_roadmap": f"Stagionalità e trend nel settore {industry} in Italia: periodi di picco, eventi chiave, opportunità stagionali, calendario marketing",
        "product_vertical": f"Analisi dettagliata prodotti {industry}: tendenze, best seller del settore, margini, strategie di pricing, cross-sell",
        "service_vertical": f"Analisi dettagliata servizi {industry} a {location}: tendenze, prezzi medi, servizi più richiesti, innovazioni",
        "objections_management": f"Obiezioni comuni dei clienti nel settore {industry}: dubbi, paure, motivi per non acquistare, come superarle",
        "psychographic_analysis": f"Profilo psicografico dei clienti {industry}: valori, lifestyle, motivazioni d'acquisto, fattori emotivi",
        "visual_brief": f"Trend visual e design per {industry}: colori, stili grafici, tipografia, mood board, riferimenti visivi 2024-2025",
        "ad_copy_creation": f"Best practice copy ads per {industry}: headline che convertono, angoli di comunicazione, hook per Meta Ads, esempi di ads efficaci",
    }

    research_query = SECTION_RESEARCH_QUERIES.get(
        workflow_step_id,
        f"Ricerca approfondita su \"{client_name}\" {industry} {location}: tutto ciò che è rilevante per {task.get('title', step_id)}"
    )

    print(f"🔬 Approfondimento {step_id}: ricerca online...")
    extra_research = ""
    try:
        extra_research = await ai_service._call_ai(
            model="perplexity/sonar-pro",
            messages=[{"role": "user", "content": research_query}],
            temperature=0.1,
            max_tokens=8000
        )
        print(f"✅ Ricerca online completata: {len(extra_research)} caratteri")
    except Exception as e:
        print(f"⚠️ Ricerca online fallita (procedo con dati esistenti): {e}")

    # Arricchisci il contesto con i dati di ricerca extra
    context = {
        "client_info": client_info,
        **snapshot,
        **(metadata.get("analysis_completa_raw", {}))
    }

    # Aggiungi la ricerca extra ai dati
    if extra_research:
        existing_raw_docs = context.get("raw_docs", "")
        context["raw_docs"] = f"{existing_raw_docs}\n\n--- RICERCA ONLINE APPROFONDITA ({task.get('title', step_id)}) ---\n{extra_research}"

    # Modifica il prompt per chiedere più profondità
    original_prompt = task["system_prompt"]
    deepened_task = {**task}
    deepened_task["system_prompt"] = (
        f"{original_prompt}\n\n"
        f"⚠️ MODALITÀ APPROFONDIMENTO ATTIVA: Hai a disposizione dati aggiuntivi da una ricerca online dedicata "
        f"(troverai una sezione '--- RICERCA ONLINE APPROFONDITA ---' nei raw_docs). "
        f"Usa TUTTI i dati disponibili per generare un'analisi MOLTO PIÙ DETTAGLIATA e approfondita del normale. "
        f"Vai in profondità su ogni punto. Aggiungi insight, strategie specifiche, esempi concreti."
    )

    print(f"🧠 Rigenerazione approfondita {step_id}...")
    new_result = await run_workflow_task(ai_service, deepened_task, context)
    cost_data = ai_service.stop_cost_tracking()

    # Salva costo
    if cost_data and cost_data.get("total_cost_usd", 0) > 0:
        print(f"💰 Costo approfondimento {step_id}: ${cost_data['total_cost_usd']:.4f}")
        if "ai_costs" not in metadata:
            metadata["ai_costs"] = []
        from datetime import datetime
        metadata["ai_costs"].append({
            "operation": cost_data["operation"],
            "cost_usd": round(cost_data["total_cost_usd"], 4),
            "prompt_tokens": cost_data["total_prompt_tokens"],
            "completion_tokens": cost_data["total_completion_tokens"],
            "calls": cost_data["calls"],
            "models": cost_data["models_used"],
            "date": datetime.now().isoformat()
        })

    # Aggiorna raw e Supabase
    if "analysis_completa_raw" not in metadata:
        metadata["analysis_completa_raw"] = {}
    metadata["analysis_completa_raw"][workflow_step_id] = new_result
    storage_service.save_metadata(client_id, metadata)

    try:
        existing_analysis = storage_service.get_complete_analysis(client_id) or {}
        existing_analysis[step_id] = new_result
        storage_service.save_complete_analysis(client_id, existing_analysis)
    except Exception as e:
        print(f"Warning: Could not update Supabase: {e}")

    return {
        "step_id": step_id,
        "new_data": new_result,
        "analysis_step": new_result,
        "cost": cost_data
    }


# ═══════════════════════════════════════════════════════════════════
# SMART LISTS API (Apple Reminders style)
# ═══════════════════════════════════════════════════════════════════

@app.get("/smart-lists")
async def get_smart_lists():
    """Get all Smart Lists (system + custom)"""
    return smart_lists_service.get_all_smart_lists()


@app.get("/smart-lists/custom")
async def get_custom_smart_lists():
    """Get only custom Smart Lists (created by user)"""
    return smart_lists_service.get_custom_smart_lists()


class SmartListCreate(BaseModel):
    title: str
    color: str
    icon: str
    criteria: Dict[str, Any]


@app.post("/smart-lists")
async def create_smart_list(data: SmartListCreate):
    """Create a new custom Smart List"""
    return smart_lists_service.create_smart_list(
        title=data.title,
        color=data.color,
        icon=data.icon,
        criteria=data.criteria
    )


class SmartListUpdate(BaseModel):
    title: Optional[str] = None
    color: Optional[str] = None
    icon: Optional[str] = None
    criteria: Optional[Dict[str, Any]] = None


@app.patch("/smart-lists/{list_id}")
async def update_smart_list(list_id: str, data: SmartListUpdate):
    """Update a custom Smart List (system lists cannot be modified)"""
    updates = {k: v for k, v in data.dict().items() if v is not None}
    result = smart_lists_service.update_smart_list(list_id, updates)
    if result is None:
        raise HTTPException(status_code=404, detail="Smart List not found")
    return result


@app.delete("/smart-lists/{list_id}")
async def delete_smart_list(list_id: str):
    """Delete a custom Smart List (system lists cannot be deleted)"""
    success = smart_lists_service.delete_smart_list(list_id)
    if not success:
        raise HTTPException(status_code=404, detail="Smart List not found")
    return {"message": "Smart List deleted successfully"}


@app.post("/smart-lists/{list_id}/filter-tasks")
async def filter_tasks_by_smart_list(list_id: str):
    """Get tasks filtered by a specific Smart List criteria"""
    # Get the Smart List
    all_lists = smart_lists_service.get_all_smart_lists()
    smart_list = next((l for l in all_lists if l["id"] == list_id), None)

    if not smart_list:
        raise HTTPException(status_code=404, detail="Smart List not found")

    # Get all tasks
    tasks = storage_service.get_tasks()

    # Filter tasks using the Smart List criteria
    filtered_tasks = smart_lists_service.filter_tasks(tasks, smart_list["criteria"])

    return {
        "smart_list": smart_list,
        "tasks": filtered_tasks,
        "count": len(filtered_tasks)
    }


# ══════════════════════════════════════════════════════════════════
#  ARIA AGENT ENDPOINTS
# ══════════════════════════════════════════════════════════════════

# In-memory job registry for ARIA background tasks
_aria_jobs: Dict[str, Dict[str, Any]] = {}

class ARIATaskRequest(BaseModel):
    task: str
    client_id: str
    context: Optional[Dict[str, Any]] = None

class ARIAFeedbackRequest(BaseModel):
    client_id: str
    output_type: str   # "angle", "copy", "script", "analysis"
    output_content: str
    feedback: str
    kept: bool = False

@app.post("/aria/task")
async def aria_run_task(req: ARIATaskRequest):
    """Submit a task to ARIA. Returns a job_id for polling."""
    job_id = str(uuid.uuid4())
    _aria_jobs[job_id] = {"status": "running", "result": None, "error": None}

    async def _run():
        try:
            agent = get_aria_agent(ai_service, storage_service)
            result = await agent.run_task(
                task=req.task,
                client_id=req.client_id,
                context=req.context or {},
                job_id=job_id,
            )
            _aria_jobs[job_id] = {"status": "done", "result": result, "error": None}
        except Exception as e:
            print(f"❌ ARIA job {job_id} failed: {e}")
            _aria_jobs[job_id] = {"status": "error", "result": None, "error": str(e)}

    asyncio.create_task(_run())
    return {"job_id": job_id, "status": "running"}

@app.get("/aria/task/{job_id}")
async def aria_get_task(job_id: str):
    """Poll ARIA job status and result."""
    job = _aria_jobs.get(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job ARIA non trovato")
    return job

@app.post("/aria/feedback")
async def aria_save_feedback(req: ARIAFeedbackRequest):
    """Save explicit feedback on an ARIA-generated output so it can improve."""
    aria_memory.save_feedback(
        client_id=req.client_id,
        output_type=req.output_type,
        output_content=req.output_content,
        feedback=req.feedback,
        kept=req.kept,
    )
    return {"message": "Feedback salvato. ARIA imparerà da questo per i prossimi output."}

@app.get("/aria/memory/{client_id}")
async def aria_get_memory(client_id: str):
    """Returns what ARIA has learned about a specific client."""
    return aria_memory.get_client_summary(client_id)
